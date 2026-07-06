import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict, Counter

def vowel_clusters(word):
	"""
	Identify vowel clusters in a word with shared consonants between them.
	Returns a list of tuples (prev_consonant, vowel_cluster, next_consonant)
	where consonants are shared between adjacent clusters.
	"""
	clusters = re.findall(r"(^|[^aeiouy]+)([aeiouy]+)([^aeiouy]+$)?", word.lower())
	if not clusters:
		return []

	processed = []
	prev_end_consonant = ""

	for i, (start, vowels, end) in enumerate(clusters):
		if i == 0:
			prev_consonant = start
		else:
			prev_consonant = prev_end_consonant

		if i == len(clusters) - 1:
			next_consonant = end if end is not None else ""
		else:
			if clusters[i+1][0]:
				next_consonant = clusters[i+1][0]
				prev_end_consonant = clusters[i+1][0]
			else:
				next_consonant = ""
				prev_end_consonant = ""

		processed.append((prev_consonant, vowels, next_consonant))

	return processed

def count_sybs(word):
	return len(vowel_clusters(word))


exceptions = 0
total_adv = 0
# Initialize Word documents for exceptions
oxford_doc = docx.Document()
oxford_doc.add_heading('Oxford Verb Declension Exceptions', 0)

riverside_doc = docx.Document()
riverside_doc.add_heading('Riverside Verb Declension Exceptions', 0)

combined_doc = docx.Document()
combined_doc.add_heading('Combined Verb Declension Exceptions', 0)

# Configuration
base_csv_dir = 'dataset'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
					"hadden", "his", "her", "him", "hers", "hide", "hir",
					"hire", "hires", "hirs", "han"]

form_csv = 'verb_forms_simple.csv'
verbs_dict = {}
with open(form_csv, 'r', encoding='utf-8') as csvfile:
	reader = csv.DictReader(csvfile)
	for row in reader:
		if 'headword' in row and 'classification' in row:
			verbs_dict[row['headword'].lower()] = row['classification'].lower()

def is_strong(verb):
	return verbs_dict.get(verb, '') == 'strong'

def is_weak(verb):
	return verbs_dict.get(verb, '') == 'weak'

def is_elided(word, next_word):
	"""Check if word is elided before next word"""
	if next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS:
		return True
	return False

def clean_tag(tag):
	"""Remove digits before % in a tag (e.g., v2%pt_1 -> v%pt_1)."""
	if tag.startswith('ger') and '%' not in tag:
		return 'ger'
	return re.sub(r'\d+(?=%)', '', tag)

def parse_tagged_text(text, tagging, text_type):
	"""Extract words from text and tags from tagging"""
	if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
		return [], [], []

	words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()
	tag_tokens = tagging.strip().split()
	tags = []
	headwords = []

	for token in tag_tokens:
		if '@' in token:
			parts = token.split('@')
			headword = parts[0].lower()
			tag = parts[1] if len(parts) > 1 else ''
			headwords.append(headword)
			tags.append(tag)

	min_len = min(len(words), len(tags), len(headwords))
	return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(doc, exception_type, word, text, line_number, filename):
	"""Add an exception case to the Word document"""
	doc_para = doc.add_paragraph()
	doc_para.add_run(f"{exception_type}: {word}").bold = True
	doc_para.add_run(f"\n{line_number} ({filename})\n")
	for i, ox_word in enumerate(text.split(' ')):
		if ox_word.lower() == word.lower():
			doc_para.add_run(f"{ox_word}").italic = True
		else:
			doc_para.add_run(f"{ox_word}")
		if i < len(text.split(' ')) - 1:
			doc_para.add_run(" ")
	doc_para.add_run("\n\n")

def analyze_verbs(df, results, text_type, doc, current_filename):
	global exceptions
	global total_adv
	"""Analyze verb patterns in CSV data according to rules."""
	text_col = f'{text_type}_TEXT'
	tagging_col = f'{text_type}_TAGGING'
	filename_col = f'{text_type}_FILENAME'
	original_text_col = f'OG_{text_type}_TEXT'

	for idx, row in df.iterrows():
		text = row[text_col]
		tagging = row[tagging_col]
		line_number = row['LINE_NUMBER']
		filename = row[filename_col]
		original_text = row.get(original_text_col, text)

		words, headwords, tags = parse_tagged_text(text, tagging, text_type)

		for j in range(len(tags)):
			if j >= len(words) or j >= len(headwords):
				continue

			word = words[j].lower()
			headword = headwords[j]
			tag = clean_tag(tags[j])
			next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'

			# Skip if elided
			if is_elided(word, next_word):
				continue
			if not tag.startswith('adv'):
				continue

			if not (word.endswith('ly') or word.endswith('ful') or word.endswith('wel')) and count_sybs(word)>1:
				if not word.endswith('e'):
					exceptions+=1
					print(word)
				total_adv+=1
				#print('Exception! Adverb not ending in ly, ful, or wel, does not end with e.')
				#print(original_text)
				#print(word)
				#print(tag)
				#print(line_number)
			
						

	return 

def process_csv_directory(csv_dir, text_type, doc):
	results = {
		'exceptions': [],
		'ending_counts': defaultdict(Counter),
		'verb_tag_counts': defaultdict(lambda: defaultdict(Counter)),
		'ending_counts_by_class': defaultdict(Counter),
		'file_exceptions': defaultdict(list),
		'file_ending_counts': defaultdict(lambda: defaultdict(Counter)),
		'file_verb_tag_counts': defaultdict(lambda: defaultdict(lambda: defaultdict(Counter))),
		'file_ending_counts_by_class': defaultdict(lambda: defaultdict(Counter))
	}

	file_count = 0
	for root, dirs, files in os.walk(csv_dir):
		for file in files:
			if not file.endswith('_gui_complete.csv'):
				continue
			file_count += 1
			try:
				df = pd.read_csv(os.path.join(root, file), encoding='utf-8')
				required_columns = [f'{text_type}_TAGGING', f'{text_type}_TEXT', 'LINE_NUMBER', f'{text_type}_FILENAME']
				if not all(col in df.columns for col in required_columns):
					print(f"Skipping {file} for {text_type}: missing columns")
					continue
				results = analyze_verbs(df, results, text_type, doc, file)
			except Exception as e:
				print(f"Error processing {file} for {text_type}: {e}")
				continue
	print(f"Processed {file_count} files for {text_type}")
	return results

def write_results(results, output_dir, text_type):
	os.makedirs(output_dir, exist_ok=True)

	# Exceptions - aggregated
	with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type', 'source_file']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		for rec in results['exceptions']:
			writer.writerow(rec)

	# Exceptions - by file
	with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['source_file', 'exception_count', 'headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		for source_file, exceptions in results['file_exceptions'].items():
			for rec in exceptions:
				writer.writerow({
					'source_file': source_file,
					'exception_count': len(exceptions),
					**rec
				})

	# Distribution by tag - aggregated
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Tag', 'Ending', 'Count', 'Percent'])
		for tag, counts in results['ending_counts'].items():
			total = sum(counts.values())
			for ending, count in counts.items():
				percent = (count / total) * 100 if total > 0 else 0
				writer.writerow([tag, ending, count, f"{percent:.2f}%"])

	# Distribution by tag - by file
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Source_File', 'Tag', 'Ending', 'Count', 'Percent'])
		for source_file, tags in results['file_ending_counts'].items():
			for tag, counts in tags.items():
				total = sum(counts.values())
				for ending, count in counts.items():
					percent = (count / total) * 100 if total > 0 else 0
					writer.writerow([source_file, tag, ending, count, f"{percent:.2f}%"])

	# NEW: Distribution by tag and verb class - aggregated
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution_by_class.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Tag', 'Verb_Class', 'Ending', 'Count', 'Percent'])
		for (tag, verb_class), counts in sorted(results['ending_counts_by_class'].items()):
			total = sum(counts.values())
			for ending, count in counts.items():
				percent = (count / total) * 100 if total > 0 else 0
				writer.writerow([tag, verb_class, ending, count, f"{percent:.2f}%"])

	# NEW: Distribution by tag and verb class - by file
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution_by_class_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Source_File', 'Tag', 'Verb_Class', 'Ending', 'Count', 'Percent'])
		for source_file, tag_class_dict in results['file_ending_counts_by_class'].items():
			for (tag, verb_class), counts in sorted(tag_class_dict.items()):
				total = sum(counts.values())
				for ending, count in counts.items():
					percent = (count / total) * 100 if total > 0 else 0
					writer.writerow([source_file, tag, verb_class, ending, count, f"{percent:.2f}%"])

	# Distribution by verb+tag - aggregated
	with open(os.path.join(output_dir, f'{text_type.lower()}_verb_tag_distribution.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Headword', 'Tag', 'Ending', 'Count', 'Percent'])
		for verb, tags in results['verb_tag_counts'].items():
			for tag, counts in tags.items():
				total = sum(counts.values())
				for ending, count in counts.items():
					percent = (count / total) * 100 if total > 0 else 0
					writer.writerow([verb, tag, ending, count, f"{percent:.2f}%"])

	# Distribution by verb+tag - by file
	with open(os.path.join(output_dir, f'{text_type.lower()}_verb_tag_distribution_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Source_File', 'Headword', 'Tag', 'Ending', 'Count', 'Percent'])
		for source_file, verbs in results['file_verb_tag_counts'].items():
			for verb, tags in verbs.items():
				for tag, counts in tags.items():
					total = sum(counts.values())
					for ending, count in counts.items():
						percent = (count / total) * 100 if total > 0 else 0
						writer.writerow([source_file, verb, tag, ending, count, f"{percent:.2f}%"])

def write_combined_results(oxford_results, riverside_results, output_dir):
	"""Create combined analysis showing overlaps and differences"""
	os.makedirs(output_dir, exist_ok=True)

	# Create lookup keys for exceptions (headword + word + tag)
	oxford_exceptions = {}
	for exc in oxford_results['exceptions']:
		key = (exc['headword'], exc['word'], exc['tag'])
		oxford_exceptions[key] = exc

	riverside_exceptions = {}
	for exc in riverside_results['exceptions']:
		key = (exc['headword'], exc['word'], exc['tag'])
		riverside_exceptions[key] = exc

	all_keys = set(oxford_exceptions.keys()) | set(riverside_exceptions.keys())

	# Combined exceptions with TYPE column
	with open(os.path.join(output_dir, 'combined_exceptions.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['TYPE', 'headword', 'word', 'tag', 'oxford_count', 'riverside_count', 
					  'oxford_files', 'riverside_files', 'oxford_contexts', 'riverside_contexts']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()

		for key in sorted(all_keys):
			in_oxford = key in oxford_exceptions
			in_riverside = key in riverside_exceptions

			if in_oxford and in_riverside:
				exc_type = 'BOTH'
			elif in_oxford:
				exc_type = 'OXFORD'
			else:
				exc_type = 'RIVERSIDE'

			# Collect all instances
			oxford_contexts = []
			oxford_files = []
			riverside_contexts = []
			riverside_files = []

			if in_oxford:
				for exc in oxford_results['exceptions']:
					if (exc['headword'], exc['word'], exc['tag']) == key:
						oxford_contexts.append(f"{exc['line_number']}: {exc['context']}")
						oxford_files.append(exc['filename'])

			if in_riverside:
				for exc in riverside_results['exceptions']:
					if (exc['headword'], exc['word'], exc['tag']) == key:
						riverside_contexts.append(f"{exc['line_number']}: {exc['context']}")
						riverside_files.append(exc['filename'])

			writer.writerow({
				'TYPE': exc_type,
				'headword': key[0],
				'word': key[1],
				'tag': key[2],
				'oxford_count': len(oxford_contexts),
				'riverside_count': len(riverside_contexts),
				'oxford_files': '; '.join(set(oxford_files)),
				'riverside_files': '; '.join(set(riverside_files)),
				'oxford_contexts': ' | '.join(oxford_contexts[:3]),  # Limit to 3 examples
				'riverside_contexts': ' | '.join(riverside_contexts[:3])
			})

	# Summary statistics
	both_count = sum(1 for key in all_keys if key in oxford_exceptions and key in riverside_exceptions)
	oxford_only = sum(1 for key in all_keys if key in oxford_exceptions and key not in riverside_exceptions)
	riverside_only = sum(1 for key in all_keys if key in riverside_exceptions and key not in oxford_exceptions)

	with open(os.path.join(output_dir, 'combined_summary.txt'), 'w', encoding='utf-8') as f:
		f.write("COMBINED EXCEPTION ANALYSIS SUMMARY\n")
		f.write("=" * 50 + "\n\n")
		f.write(f"Total unique exceptions: {len(all_keys)}\n")
		f.write(f"Exceptions in BOTH: {both_count} ({both_count/len(all_keys)*100:.1f}%)\n")
		f.write(f"Exceptions OXFORD only: {oxford_only} ({oxford_only/len(all_keys)*100:.1f}%)\n")
		f.write(f"Exceptions RIVERSIDE only: {riverside_only} ({riverside_only/len(all_keys)*100:.1f}%)\n\n")
		f.write(f"Total Oxford exception instances: {len(oxford_results['exceptions'])}\n")
		f.write(f"Total Riverside exception instances: {len(riverside_results['exceptions'])}\n")

	# Add to combined doc
	combined_doc.add_heading('Summary Statistics', 1)
	combined_doc.add_paragraph(f"Total unique exceptions: {len(all_keys)}")
	combined_doc.add_paragraph(f"Exceptions in BOTH: {both_count}")
	combined_doc.add_paragraph(f"Exceptions OXFORD only: {oxford_only}")
	combined_doc.add_paragraph(f"Exceptions RIVERSIDE only: {riverside_only}")

	print(f"\nCombined analysis: {both_count} overlapping, {oxford_only} Oxford-only, {riverside_only} Riverside-only")

def write_summary_docx(results, output_dir, text_type):
	summary_doc = docx.Document()
	summary_doc.add_heading(f'{text_type} Verb Declension Analysis Summary', 0)

	summary_doc.add_heading('Exceptions', 1)
	summary_doc.add_paragraph(f"Total exceptions found: {len(results['exceptions'])}")
	
	# Per-file breakdown
	summary_doc.add_heading('Exceptions by File', 2)
	for source_file, exceptions in sorted(results['file_exceptions'].items()):
		summary_doc.add_paragraph(f"{source_file}: {len(exceptions)} exceptions")

	summary_doc.add_heading('Ending Distribution by Tag (Aggregated)', 1)
	for tag, counts in results['ending_counts'].items():
		para = summary_doc.add_paragraph()
		para.add_run(f"{tag}: ").bold = True
		total = sum(counts.values())
		for ending, count in counts.items():
			percent = (count / total) * 100 if total > 0 else 0
			para.add_run(f"{ending}={count} ({percent:.1f}%)  ")

	summary_doc.save(os.path.join(output_dir, f'{text_type.lower()}_analysis_summary.docx'))

# Main execution
if __name__ == "__main__":
	#print("Processing Oxford texts...")
	#oxford_results = process_csv_directory(base_csv_dir, 'OXFORD', oxford_doc)
	
	print("Processing Riverside texts...")
	riverside_results = process_csv_directory(base_csv_dir, 'RIVERSIDE', riverside_doc)
	print(exceptions/total_adv)
