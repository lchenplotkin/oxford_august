import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict, Counter
from datetime import date

# Initialize Word documents for exceptions
oxford_doc = docx.Document()
oxford_doc.add_heading('Oxford Verb Declension Exceptions', 0)

riverside_doc = docx.Document()
riverside_doc.add_heading('Riverside Verb Declension Exceptions', 0)

combined_doc = docx.Document()
combined_doc.add_heading('Combined Verb Declension Exceptions', 0)

# Configuration
base_csv_dir = 'dataset'
ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde", "hadden",
					 "his", "her", "him", "hers", "hide", "hir", "hire", "hires", "hirs", "han"]

form_csv = 'verb_forms_simple.csv'
verbs_dict = {}

with open(form_csv, 'r', encoding='utf-8') as csvfile:
	reader = csv.DictReader(csvfile)
	for row in reader:
		if 'headword' in row and 'classification' in row:
			verbs_dict[row['headword'].lower()] = row['classification'].lower()

def is_strong(verb):
	return verbs_dict.get(verb, '') == 'strong'

def is_weak(verb):
	return verbs_dict.get(verb, '') == 'weak'

def is_elided(word, next_word):
	"""Check if word is elided before next word"""
	if next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS:
		return True
	return False

def clean_tag(tag):
	"""Remove digits before % in a tag (e.g., v2%pt_1 -> v%pt_1)."""
	return re.sub(r'\d+(?=%)', '', tag)

def parse_tagged_text(text, tagging, text_type):
	"""Extract words from text and tags from tagging"""
	if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
		return [], [], []
	
	words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()
	tag_tokens = tagging.strip().split()
	
	tags = []
	headwords = []
	
	for token in tag_tokens:
		if '@' in token:
			parts = token.split('@')
			headword = parts[0].lower()
			tag = parts[1] if len(parts) > 1 else ''
			headwords.append(headword)
			tags.append(tag)
	
	min_len = min(len(words), len(tags), len(headwords))
	return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(doc, exception_type, word, text, line_number, filename):
	"""Add an exception case to the Word document"""
	doc_para = doc.add_paragraph()
	doc_para.add_run(f"{exception_type}: {word}").bold = True
	doc_para.add_run(f"\n{line_number} ({filename})\n")
	
	for i, ox_word in enumerate(text.split(' ')):
		if ox_word.lower() == word.lower():
			doc_para.add_run(f"{ox_word}").italic = True
		else:
			doc_para.add_run(f"{ox_word}")
		if i < len(text.split(' ')) - 1:
			doc_para.add_run(" ")
	
	doc_para.add_run("\n\n")

def classify_ending(word):
	"""Classify the word ending into categories."""
	if word.endswith('eth'):
		return '-eth'
	if word.endswith('en'):
		return '-en'
	"""
	if word.endswith('ede'):
		return '-ede'
	if word.endswith('de'):
		return '-de'
	if word.endswith('te'):
		return '-te'
	"""
	if word.endswith('e'):
		return '-e'
	if word.endswith('ed'):
		return '-ed'
	if word.endswith('d'):
		return '-d'
	if word.endswith('est'):
		return '-est'
	if word.endswith('et'):
		return '-et'
	if word.endswith('t'):
		return '-t'
	if word.endswith(('a','e','i','o','u')):
		return 'vowel'
	return 'other'

def classify_verb_by_morphology(patterns):
	"""
	Classify a verb as strong, weak, or irregular based on its observed patterns.
	
	Args:
		patterns: dict of {tag: Counter({ending: count})}
	
	Returns:
		str: 'strong', 'weak', 'unknown', or 'irregular'
	"""
	# Evidence counters
	strong_evidence = 0
	weak_evidence = 0
	
	# Check preterite singular (pt_1, pt_3) patterns
	for pt_tag in ['v%pt_1', 'v%pt_3']:
		if pt_tag in patterns:
			endings = patterns[pt_tag]
			total = sum(endings.values())
			if total > 0:
				# Strong verbs don't end in -e in preterite singular
				e_endings = endings.get('-e', 0) + endings.get('-en', 0)
				weak_endings = (endings.get('-ede', 0) + endings.get('-de', 0) + 
							   endings.get('-te', 0) + endings.get('-ed', 0) + 
							   endings.get('-d', 0) + endings.get('-t', 0))
				
				if e_endings>0:
					weak_evidence += 2  # Strong evidence for weak
				elif weak_endings == 0:
					strong_evidence += 2  # Strong evidence for strong
				elif (weak_endings + e_endings) < total/4: 
					strong_evidence += 1
	
	if 'v%pt_2' in patterns:
		endings = patterns['v%pt_2']
		total = sum(endings.values())
		if total>0:
			e_endings = endings.get('-e', 0) + endings.get('-en', 0)
			if e_endings==total:
				strong_evidence+=2*total
			elif e_endings == 0:
				weak_evidence+=2*total
			elif e_endings>total/2:
				strong_evidence+=total
			else:
				weak_evidence+=total
		
	# Check past participle pattern
	if 'v%ppl' in patterns:
		endings = patterns['v%ppl']
		total = sum(endings.values())
		if total > 0:
			en_e_endings = endings.get('-en', 0) + endings.get('-e', 0)
			weak_endings = (endings.get('-ede', 0) + endings.get('-de', 0) + 
						   endings.get('-te', 0) + endings.get('-ed', 0) + 
						   endings.get('-d', 0) + endings.get('-t', 0))
			
			# Strong verbs typically have -en/-e participles
			if en_e_endings == total:
				strong_evidence += total*2
			elif weak_endings == total:
				weak_evidence += total*2
	
	# Make classification decision
	if strong_evidence > 2*weak_evidence: 
		return 'strong'
	elif weak_evidence > 2*strong_evidence:
		return 'weak'
	elif weak_evidence == 0 and strong_evidence == 0:
		return 'unknown'
	else:
		return 'irregular'

def analyze_verbs(df, results, text_type, doc, current_filename):
	"""Analyze verb patterns in CSV data according to rules."""
	text_col = f'{text_type}_TEXT'
	tagging_col = f'{text_type}_TAGGING'
	filename_col = f'{text_type}_FILENAME'
	original_text_col = f'OG_{text_type}_TEXT'
	
	for idx, row in df.iterrows():
		text = row[text_col]
		tagging = row[tagging_col]
		line_number = row['LINE_NUMBER']
		filename = row[filename_col]
		original_text = row.get(original_text_col, text)
		
		words, headwords, tags = parse_tagged_text(text, tagging, text_type)
		
		for j in range(len(tags)):
			if j >= len(words) or j >= len(headwords):
				continue
			
			word = words[j].lower()
			headword = headwords[j]
			tag = clean_tag(tags[j])
			next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'
			
			if not tag.startswith('v'):
				continue
			
			ending = classify_ending(word)
			
			# Determine verb classification
			verb_class = 'unknown'
			if headword in verbs_dict:
				if is_strong(headword):
					verb_class = 'strong'
				elif is_weak(headword):
					verb_class = 'weak'
				else:
					verb_class = 'irregular'
			else:
				# Track unknown verbs
				results['unknown_verbs'][headword][tag][ending] += 1
			
			# Count ending distribution - overall
			results['ending_counts'][tag][ending] += 1
			results['verb_tag_counts'][headword][tag][ending] += 1
			
			# Count ending distribution by verb class
			results['ending_counts_by_class'][(tag, verb_class)][ending] += 1
			results['file_ending_counts_by_class'][current_filename][(tag, verb_class)][ending] += 1
			
			# Count ending distribution - per file
			results['file_ending_counts'][current_filename][tag][ending] += 1
			results['file_verb_tag_counts'][current_filename][headword][tag][ending] += 1
			
			# Rule checks
			violated = False
			reason = ""
			
			if tag == 'v%inf' and ending not in ['-en', '-e','vowel']:
				violated = True
				reason += "RULE 1: The infinitive ends in -en or -e unless the stem ends in a vowel."
			
			if tag == 'v%pt_pl' and ending not in ['-en', '-e']:
				violated = True
				reason += "RULE 4: Preterite plural verbs end in -e/-en"
			
			if tag == 'v%pr_1' and ending not in ['-en','-e']:
				violated = True
				reason += "RULE 5: First person singular present tense verbs end in -e/-en"
			
			if tag == 'v%pr_pl' and ending not in ['-en', '-e']:
				violated = True
				reason += "RULE 6: Present plural verbs end in -e/-en"
			
			if verb_class == "weak" and tag in ['v%pt_1', 'v%pt_3'] and ending not in ['-ede','-de','-te','-ed','-d','-t','-est','-et']:
				violated = True
				reason += "Rule 7: Singular weak preterite verbs in the first and third person end in -ede, -de, -te, -ed, -d, or -t"
			
			if verb_class == "strong" and tag in ['v%pt_1', 'v%pt_3'] and ending in ['-e']:
				violated = True
				reason += "Rule 8: Singular strong preterite verbs in the first and third person do not end in -e"
			
			if verb_class == "strong" and tag == 'v%ppl' and ending not in ['-en', '-e']:
				violated = True
				reason += "Strong participle must end in -en or -e "
			
			if violated:
				record = {
					'headword': headword,
					'word': word,
					'tag': tag,
					'line_number': line_number,
					'filename': filename,
					'context': original_text,
					'text_type': text_type,
					'source_file': current_filename
				}
				results['exceptions'].append(record)
				results['file_exceptions'][current_filename].append(record)
				add_exception_to_doc(doc, reason, word, original_text, line_number, filename)
	
	return results

def process_csv_directory(csv_dir, text_type, doc):
	results = {
		'exceptions': [],
		'ending_counts': defaultdict(Counter),
		'verb_tag_counts': defaultdict(lambda: defaultdict(Counter)),
		'ending_counts_by_class': defaultdict(Counter),
		'file_exceptions': defaultdict(list),
		'file_ending_counts': defaultdict(lambda: defaultdict(Counter)),
		'file_verb_tag_counts': defaultdict(lambda: defaultdict(lambda: defaultdict(Counter))),
		'file_ending_counts_by_class': defaultdict(lambda: defaultdict(Counter)),
		'unknown_verbs': defaultdict(lambda: defaultdict(Counter))  # NEW
	}
	
	file_count = 0
	for root, dirs, files in os.walk(csv_dir):
		for file in files:
			if not file.endswith('_gui_complete.csv'):
				continue
			
			file_count += 1
			try:
				df = pd.read_csv(os.path.join(root, file), encoding='utf-8')
				required_columns = [f'{text_type}_TAGGING', f'{text_type}_TEXT', 
								  'LINE_NUMBER', f'{text_type}_FILENAME']
				
				if not all(col in df.columns for col in required_columns):
					print(f"Skipping {file} for {text_type}: missing columns")
					continue
				
				results = analyze_verbs(df, results, text_type, doc, file)
				
			except Exception as e:
				print(f"Error processing {file} for {text_type}: {e}")
				continue
	
	print(f"Processed {file_count} files for {text_type}")
	return results

def write_unknown_verbs_csv(oxford_results, riverside_results, output_file):
	"""
	Write a CSV of unknown verbs with morphological classification.
	Combines data from both Oxford and Riverside results.
	"""
	# Combine unknown verbs from both sources
	all_unknown = defaultdict(lambda: defaultdict(Counter))
	
	for headword, patterns in oxford_results['unknown_verbs'].items():
		for tag, endings in patterns.items():
			all_unknown[headword][tag].update(endings)
	
	for headword, patterns in riverside_results['unknown_verbs'].items():
		for tag, endings in patterns.items():
			all_unknown[headword][tag].update(endings)
	
	# Classify each unknown verb
	classified_verbs = []
	today = date.today().strftime("%B %d, %Y")
	
	for headword, patterns in sorted(all_unknown.items()):
		classification = classify_verb_by_morphology(patterns)
		
		# Gather evidence details
		evidence = []
		for tag in sorted(patterns.keys()):
			endings = patterns[tag]
			total = sum(endings.values())
			ending_list = ', '.join([f"{ending}:{count}" for ending, count in endings.most_common()])
			evidence.append(f"{tag}({ending_list})")
		
		classified_verbs.append({
			'headword': headword,
			'classification': classification,
			'notes': f'classified by morphological evidence on {today}',
			'evidence': '; '.join(evidence),
			'total_occurrences': sum(sum(endings.values()) for endings in patterns.values())
		})
	
	# Write to CSV
	with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['headword', 'classification', 'notes', 'evidence', 'total_occurrences']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		writer.writerows(classified_verbs)
	
	print(f"\nUnknown verbs classified: {len(classified_verbs)}")
	print(f"  - Strong: {sum(1 for v in classified_verbs if v['classification'] == 'strong')}")
	print(f"  - Weak: {sum(1 for v in classified_verbs if v['classification'] == 'weak')}")
	print(f"  - Irregular: {sum(1 for v in classified_verbs if v['classification'] == 'irregular')}")
	print(f"  - Unknown: {sum(1 for v in classified_verbs if v['classification'] == 'unknown')}")
	
	return classified_verbs

def create_updated_verb_forms_csv(original_csv, unknown_verbs, output_file):
	"""
	Create an updated verb_forms_simple.csv with unknown verbs added.
	"""
	# Read original CSV
	existing_verbs = []
	with open(original_csv, 'r', encoding='utf-8') as csvfile:
		reader = csv.DictReader(csvfile)
		fieldnames = reader.fieldnames
		for row in reader:
			existing_verbs.append(row)
	
	# Ensure we have the right fieldnames
	if 'notes' not in fieldnames:
		fieldnames = list(fieldnames) + ['notes']
	
	# Add unknown verbs
	for unknown in unknown_verbs:
		existing_verbs.append({
			'headword': unknown['headword'],
			'classification': unknown['classification'],
			'notes': unknown['notes']
		})
	
	# Write updated CSV
	with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		writer.writerows(existing_verbs)
	
	print(f"\nUpdated verb forms CSV created: {output_file}")
	print(f"Total verbs: {len(existing_verbs)} (added {len(unknown_verbs)} new)")

def write_results(results, output_dir, text_type):
	os.makedirs(output_dir, exist_ok=True)
	
	# Exceptions - aggregated
	with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type', 'source_file']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		for rec in results['exceptions']:
			writer.writerow(rec)
	
	# Exceptions - by file
	with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['source_file', 'exception_count', 'headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		for source_file, exceptions in results['file_exceptions'].items():
			for rec in exceptions:
				writer.writerow({
					'source_file': source_file,
					'exception_count': len(exceptions),
					**rec
				})
	
	# Distribution by tag - aggregated
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Tag', 'Ending', 'Count', 'Percent'])
		for tag, counts in results['ending_counts'].items():
			total = sum(counts.values())
			for ending, count in counts.items():
				percent = (count / total) * 100 if total > 0 else 0
				writer.writerow([tag, ending, count, f"{percent:.2f}%"])
	
	# Distribution by tag - by file
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Source_File', 'Tag', 'Ending', 'Count', 'Percent'])
		for source_file, tags in results['file_ending_counts'].items():
			for tag, counts in tags.items():
				total = sum(counts.values())
				for ending, count in counts.items():
					percent = (count / total) * 100 if total > 0 else 0
					writer.writerow([source_file, tag, ending, count, f"{percent:.2f}%"])
	
	# Distribution by tag and verb class - aggregated
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution_by_class.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Tag', 'Verb_Class', 'Ending', 'Count', 'Percent'])
		for (tag, verb_class), counts in sorted(results['ending_counts_by_class'].items()):
			total = sum(counts.values())
			for ending, count in counts.items():
				percent = (count / total) * 100 if total > 0 else 0
				writer.writerow([tag, verb_class, ending, count, f"{percent:.2f}%"])
	
	# Distribution by tag and verb class - by file
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution_by_class_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Source_File', 'Tag', 'Verb_Class', 'Ending', 'Count', 'Percent'])
		for source_file, tag_class_dict in results['file_ending_counts_by_class'].items():
			for (tag, verb_class), counts in sorted(tag_class_dict.items()):
				total = sum(counts.values())
				for ending, count in counts.items():
					percent = (count / total) * 100 if total > 0 else 0
					writer.writerow([source_file, tag, verb_class, ending, count, f"{percent:.2f}%"])
	
	# Distribution by verb+tag - aggregated
	with open(os.path.join(output_dir, f'{text_type.lower()}_verb_tag_distribution.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Headword', 'Tag', 'Ending', 'Count', 'Percent'])
		for verb, tags in results['verb_tag_counts'].items():
			for tag, counts in tags.items():
				total = sum(counts.values())
				for ending, count in counts.items():
					percent = (count / total) * 100 if total > 0 else 0
					writer.writerow([verb, tag, ending, count, f"{percent:.2f}%"])
	
	# Distribution by verb+tag - by file
	with open(os.path.join(output_dir, f'{text_type.lower()}_verb_tag_distribution_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Source_File', 'Headword', 'Tag', 'Ending', 'Count', 'Percent'])
		for source_file, verbs in results['file_verb_tag_counts'].items():
			for verb, tags in verbs.items():
				for tag, counts in tags.items():
					total = sum(counts.values())
					for ending, count in counts.items():
						percent = (count / total) * 100 if total > 0 else 0
						writer.writerow([source_file, verb, tag, ending, count, f"{percent:.2f}%"])

def write_combined_results(oxford_results, riverside_results, output_dir):
	"""Create combined analysis showing overlaps and differences"""
	os.makedirs(output_dir, exist_ok=True)
	
	# Create lookup keys for exceptions (headword + word + tag)
	oxford_exceptions = {}
	for exc in oxford_results['exceptions']:
		key = (exc['headword'], exc['word'], exc['tag'])
		oxford_exceptions[key] = exc
	
	riverside_exceptions = {}
	for exc in riverside_results['exceptions']:
		key = (exc['headword'], exc['word'], exc['tag'])
		riverside_exceptions[key] = exc
	
	all_keys = set(oxford_exceptions.keys()) | set(riverside_exceptions.keys())
	
	# Combined exceptions with TYPE column
	with open(os.path.join(output_dir, 'combined_exceptions.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['TYPE', 'headword', 'word', 'tag', 'oxford_count', 'riverside_count', 
					 'oxford_files', 'riverside_files', 'oxford_contexts', 'riverside_contexts']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		
		for key in sorted(all_keys):
			in_oxford = key in oxford_exceptions
			in_riverside = key in riverside_exceptions
			
			if in_oxford and in_riverside:
				exc_type = 'BOTH'
			elif in_oxford:
				exc_type = 'OXFORD'
			else:
				exc_type = 'RIVERSIDE'
			
			# Collect all instances
			oxford_contexts = []
			oxford_files = []
			riverside_contexts = []
			riverside_files = []
			
			if in_oxford:
				for exc in oxford_results['exceptions']:
					if (exc['headword'], exc['word'], exc['tag']) == key:
						oxford_contexts.append(f"{exc['line_number']}: {exc['context']}")
						oxford_files.append(exc['filename'])
			
			if in_riverside:
				for exc in riverside_results['exceptions']:
					if (exc['headword'], exc['word'], exc['tag']) == key:
						riverside_contexts.append(f"{exc['line_number']}: {exc['context']}")
						riverside_files.append(exc['filename'])
			
			writer.writerow({
				'TYPE': exc_type,
				'headword': key[0],
				'word': key[1],
				'tag': key[2],
				'oxford_count': len(oxford_contexts),
				'riverside_count': len(riverside_contexts),
				'oxford_files': '; '.join(set(oxford_files)),
				'riverside_files': '; '.join(set(riverside_files)),
				'oxford_contexts': ' | '.join(oxford_contexts[:3]),
				'riverside_contexts': ' | '.join(riverside_contexts[:3])
			})
	
	# Summary statistics
	both_count = sum(1 for key in all_keys if key in oxford_exceptions and key in riverside_exceptions)
	oxford_only = sum(1 for key in all_keys if key in oxford_exceptions and key not in riverside_exceptions)
	riverside_only = sum(1 for key in all_keys if key in riverside_exceptions and key not in oxford_exceptions)
	
	with open(os.path.join(output_dir, 'combined_summary.txt'), 'w', encoding='utf-8') as f:
		f.write("COMBINED EXCEPTION ANALYSIS SUMMARY\n")
		f.write("=" * 50 + "\n\n")
		f.write(f"Total unique exceptions: {len(all_keys)}\n")
		f.write(f"Exceptions in BOTH: {both_count} ({both_count/len(all_keys)*100:.1f}%)\n")
		f.write(f"Exceptions OXFORD only: {oxford_only} ({oxford_only/len(all_keys)*100:.1f}%)\n")
		f.write(f"Exceptions RIVERSIDE only: {riverside_only} ({riverside_only/len(all_keys)*100:.1f}%)\n\n")
		f.write(f"Total Oxford exception instances: {len(oxford_results['exceptions'])}\n")
		f.write(f"Total Riverside exception instances: {len(riverside_results['exceptions'])}\n")
	
	# Add to combined doc
	combined_doc.add_heading('Summary Statistics', 1)
	combined_doc.add_paragraph(f"Total unique exceptions: {len(all_keys)}")
	combined_doc.add_paragraph(f"Exceptions in BOTH: {both_count}")
	combined_doc.add_paragraph(f"Exceptions OXFORD only: {oxford_only}")
	combined_doc.add_paragraph(f"Exceptions RIVERSIDE only: {riverside_only}")
	
	print(f"\nCombined analysis: {both_count} overlapping, {oxford_only} Oxford-only, {riverside_only} Riverside-only")

def write_summary_docx(results, output_dir, text_type):
	summary_doc = docx.Document()
	summary_doc.add_heading(f'{text_type} Verb Declension Analysis Summary', 0)
	
	summary_doc.add_heading('Exceptions', 1)
	summary_doc.add_paragraph(f"Total exceptions found: {len(results['exceptions'])}")
	
	# Per-file breakdown
	summary_doc.add_heading('Exceptions by File', 2)
	for source_file, exceptions in sorted(results['file_exceptions'].items()):
		summary_doc.add_paragraph(f"{source_file}: {len(exceptions)} exceptions")
	
	summary_doc.add_heading('Ending Distribution by Tag (Aggregated)', 1)
	for tag, counts in results['ending_counts'].items():
		para = summary_doc.add_paragraph()
		para.add_run(f"{tag}: ").bold = True
		total = sum(counts.values())
		for ending, count in counts.items():
			percent = (count / total) * 100 if total > 0 else 0
			para.add_run(f"{ending}={count} ({percent:.1f}%) ")
	
	summary_doc.save(os.path.join(output_dir, f'{text_type.lower()}_analysis_summary.docx'))

# Main execution
if __name__ == "__main__":
	print("Processing Oxford texts...")
	oxford_results = process_csv_directory(base_csv_dir, 'OXFORD', oxford_doc)
	
	print("Processing Riverside texts...")
	riverside_results = process_csv_directory(base_csv_dir, 'RIVERSIDE', riverside_doc)
	
	oxford_output_dir = 'oxford_verb_analysis_output'
	riverside_output_dir = 'riverside_verb_analysis_output'
	combined_output_dir = 'combined_verb_analysis_output'
	
	write_results(oxford_results, oxford_output_dir, 'OXFORD')
	write_summary_docx(oxford_results, oxford_output_dir, 'OXFORD')
	oxford_doc.save(os.path.join(oxford_output_dir, 'oxford_declension_exceptions.docx'))
	
	write_results(riverside_results, riverside_output_dir, 'RIVERSIDE')
	write_summary_docx(riverside_results, riverside_output_dir, 'RIVERSIDE')
	riverside_doc.save(os.path.join(riverside_output_dir, 'riverside_declension_exceptions.docx'))
	
	print("\nCreating combined analysis...")
	write_combined_results(oxford_results, riverside_results, combined_output_dir)
	combined_doc.save(os.path.join(combined_output_dir, 'combined_declension_exceptions.docx'))
	
	# NEW: Process unknown verbs
	print("\n" + "="*60)
	print("CLASSIFYING UNKNOWN VERBS")
	print("="*60)
	
	unknown_verbs_file = os.path.join(combined_output_dir, 'unknown_verbs_classified.csv')
	classified_verbs = write_unknown_verbs_csv(oxford_results, riverside_results, unknown_verbs_file)
	
	# Create updated verb_forms_simple.csv
	updated_verb_forms = os.path.join(combined_output_dir, 'verb_forms_simple_updated.csv')
	create_updated_verb_forms_csv(form_csv, classified_verbs, updated_verb_forms)
	
	print("\nAnalysis complete!")
	print(f"Oxford results: {oxford_output_dir}/")
	print(f"Riverside results: {riverside_output_dir}/")
	print(f"Combined results: {combined_output_dir}/")
	print(f"\nNew files created:")
	print(f"  - Unknown verbs: {unknown_verbs_file}")
	print(f"  - Updated verb forms: {updated_verb_forms}")
