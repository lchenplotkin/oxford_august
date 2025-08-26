# Re-import after code state reset
import docx 
import os
import re
import string
import csv
import pandas as pd
from collections import defaultdict, Counter
import matplotlib.pyplot as plt

# Initialize Word document
doc = docx.Document()
doc.add_heading('Exceptions to Rule 2 in the Oxford Chaucer', 0)

# Configuration
base_csv_dir = 'data/csvs'


'''
tag_tuple_list = [
    ('demonstrative', 'adj', 1),
    ('def_art', 'adj', 1),
    ('n%gen', 'adj', 1),
    ('pron%gen', 'adj', 1),
    ('pron%femgen', 'adj', 1),
    ('adj', 'n#propn', 0),
    ('interj', 'adj', 1),
]
'''

tag_tuple_list = [
    ('n%pl','adj',1),
    ('adj', 'n%pl', 0)
]


ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
                    "hadden", "his", "her", "him", "hers", "hide", "hir",
                    "hire", "hires", "hirs", "han"]

# Utilities
def vowel_cluster_count(w):
    return len(re.findall(r'[^aeiouy]*[aeiouy]+(?:[^aeiouy]+(?=[^aeiouy]*[aeiouy])|[^aeiouy]*)?', w)) or 1

def parse_tagged_text(oxford_text, oxford_tagging):
    """Extract words from oxford_text and tags from oxford_tagging"""
    if pd.isna(oxford_text) or pd.isna(oxford_tagging) or oxford_text == '' or oxford_tagging == '':
        return [], [], []

    # Clean and split the text
    words = oxford_text.lower().replace(',', '').replace('.', '').replace('!', '').replace('?', '').replace('°', '').replace('¶', '').strip().split()

    # Parse tags from oxford_tagging
    tag_tokens = oxford_tagging.strip().split()
    tags = []
    headwords = []

    for token in tag_tokens:
        if '@' in token:
            if token not in ["--@dash", ".@ellipsis"]:
                parts = token.split('@')
                headword = parts[0].lower()
                tag_part = parts[1] if len(parts) > 1 else ''

                ## Extract tag (after # if present, otherwise use the whole tag_part)
                #if '#' in tag_part:
                #    tag = tag_part.split('#')[1]
                #else:
                tag = tag_part
                re.sub(r'\d+', '', tag)

                headwords.append(headword)
                tags.append(tag)

    # Handle special demonstrative cases
    for i, (word, tag) in enumerate(zip(words[:len(tags)], tags)):
        if word in ['this', 'that', 'thilke'] and 'gram_adj' in tag:
            tags[i] = 'demonstrative'

    # Ensure all lists are the same length (trim to shortest)
    if len(words)!= len(tags):
        print(oxford_text, oxford_tagging,words,tags)
        print(len(words),len(tags))
    min_len = min(len(words), len(tags), len(headwords))
    return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(tag1, tag2, word, oxford_text, line_number, filename):
    """Add an exception case to the Word document"""
    doc_para = doc.add_paragraph()
    doc_para.add_run(f"{tag1} -> {tag2}: {word}").bold = True
    #doc_para.add_run("Word: ").bold = True
    ##doc_para.add_run(word).italic = True
    #doc_para.add_run(f"{filename}")

    doc_para.add_run(f"\n{line_number}\n")
    for i, ox_word in enumerate(oxford_text.split(' ')):
        if ox_word.lower() == word.lower():
            doc_para.add_run(f"{ox_word}").italic = True
        else:
            doc_para.add_run(f"{ox_word}").italic = False

        if i<len(oxford_text.split(' '))-1:
            doc_para.add_run(" ")
    #doc_para.add_run("\n\n")

def search_tag_sequence_csv(df, tag1, tag2, pos, adjectives, unspelled_log, unspelled_elided_log):
    """Search for tag sequences in CSV data"""
    spelled, spelled_elided, unspelled, unspelled_elided, final = 0, 0, 0, 0, 0
    unspelled_lines = []
    unspelled_elided_lines = []

    lines = 0
    for idx, row in df.iterrows():
        if row["MATCH"] != "DIFF":
            lines+=1
            oxford_text = row['OXFORD_TEXT']
            oxford_tagging = row['OXFORD_TAGGING']
            line_number = row['LINE_NUMBER']
            filename = row['OXFORD_FILENAME']

            words, headwords, tags = parse_tagged_text(oxford_text, oxford_tagging)

            for j in range(len(tags) - 1):
                if tags[j] == tag1 and tags[j + 1] == tag2:
                    if j + pos >= len(words):
                        continue

                    word = words[j + pos]
                    headword = headwords[j + pos]

                    # Check if it's a monosyllabic adjective (excluding exceptions)
                    if (vowel_cluster_count(headword) == 1 and word not in ['chief','coward','royal','cruel','shrewed']):# and
                        #word not in ['chief','coward','lewed','payen','real','royal','roial',
                         #           'shrewed','troian','troyan','cruel','crueel','crewel','cruele']):

                        adjectives[headword][5].append(word)
                        adjectives[headword][5] = list(set(adjectives[headword][5]))

                        if j + pos == len(tags) - 1:
                            final += 1
                            adjectives[headword][4] += 1
                        else:
                            nextword = words[j + pos + 1]

                            if word.endswith('e') and (nextword[0] in 'aeiou' or nextword in ELISION_FOLLOWERS):
                                spelled_elided += 1
                                adjectives[headword][1] += 1
                            elif word.endswith('e'):
                                spelled += 1
                                adjectives[headword][0] += 1
                            elif nextword[0] in 'aeiou' or nextword in ELISION_FOLLOWERS:
                                unspelled_elided += 1
                                adjectives[headword][3] += 1
                                exception_text = f"{tag1}->{tag2}\n{word}\nLine {line_number}: {oxford_text}\n\n"
                                unspelled_elided_lines.append(exception_text)
                                # Add to Word document
                                #add_exception_to_doc(tag1, tag2, word, oxford_text, line_number, filename)
                            else:
                                unspelled += 1
                                adjectives[headword][2] += 1
                                exception_text = f"{tag1}->{tag2}\n{word}\nLine {line_number}: {oxford_text}\n\n"
                                unspelled_lines.append(exception_text)
                                # Add to Word document
                                add_exception_to_doc(tag1, tag2, word, oxford_text, line_number, filename)

    unspelled_log.extend(unspelled_lines)
    unspelled_elided_log.extend(unspelled_elided_lines)

    return lines, spelled, spelled_elided, unspelled, unspelled_elided, final

def process_csv_directory(csv_dir):
    """Process all _gui.csv files in the directory"""
    data = {}
    unspelled_log = []
    unspelled_elided_log = []
    adjectives = defaultdict(lambda: [0, 0, 0, 0, 0, []])  # spelled, spelled_elided, unspelled, unspelled_elided, final, words

    for root, dirs, files in os.walk(csv_dir):
        for file in files:
            if not file.endswith('_gui.csv'):
                continue

            csv_path = os.path.join(root, file)
            rel_path = os.path.relpath(root, csv_dir)
            file_id = os.path.join(rel_path, file) if rel_path != '.' else file

            try:
                # Read CSV file
                df = pd.read_csv(csv_path, encoding='utf-8')

                # Skip if required columns are missing
                required_columns = ['OXFORD_TAGGING', 'OXFORD_TEXT', 'LINE_NUMBER', 'OXFORD_FILENAME']
                if not all(col in df.columns for col in required_columns):
                    print(f"Skipping {file}: missing required columns")
                    continue

                tag_data = {}
                total_lines, total_spelled, total_spelled_elided, total_unspelled, total_unspelled_elided, total_final = 0, 0, 0, 0, 0, 0

                for tag1, tag2, pos in tag_tuple_list:
                    lines, s, s_e, u, u_e, f = search_tag_sequence_csv(df, tag1, tag2, pos, adjectives, unspelled_log, unspelled_elided_log)
                    tag_data[f'{tag1}->{tag2} spelled'] = s
                    tag_data[f'{tag1}->{tag2} spelled, elided'] = s_e
                    tag_data[f'{tag1}->{tag2} unspelled'] = u
                    tag_data[f'{tag1}->{tag2} unspelled, elided'] = u_e
                    tag_data[f'{tag1}->{tag2} %'] = round((s / (s + u) * 100), 2) if (s + u) > 0 else 0.0
                    total_spelled += s
                    total_spelled_elided += s_e
                    total_unspelled += u
                    total_unspelled_elided += u_e
                    total_lines += lines

                match_ratio = total_spelled / (total_spelled + total_unspelled) if (total_spelled + total_unspelled) > 0 else 0
                data[file_id] = {
                    'total lines': total_lines,
                    'overall spelled unelided': total_spelled,
                    'overall unspelled unelided': total_unspelled,
                    'overall spelled elided': total_spelled_elided,
                    'overall unspelled elided': total_unspelled_elided,
                    'overall final': total_final,
                    'overall unelided spelled/unspelled ratio': round(match_ratio * 100, 2),
                    **tag_data
                }

            except Exception as e:
                print(f"Error processing {file}: {e}")
                continue

    return data, adjectives, unspelled_log, unspelled_elided_log

def write_csv(data, filename):
    """Write analysis results to CSV"""
    if not data:
        print("No data to write")
        return

    keys = sorted(next(iter(data.values())).keys())
    with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=['file'] + keys)
        writer.writeheader()
        for file, stats in data.items():
            row = {'file': file}
            row.update(stats)
            writer.writerow(row)

def write_adjectives(adjectives, filename):
    """Write adjective analysis to CSV"""
    with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['headword','weak_forms', 'spelled', 'spelled_elided', 'unspelled', 'unspelled_elided', 'final','total','unelided_pct','overall_pct'])
        for adj in sorted(adjectives.keys()):
            s, s_e, u, u_e, f, forms = adjectives[adj]
            total = s + s_e + u + u_e + f
            match_pct_unelided = round((s / (s + u) * 100), 2) if (s + u) > 0 else 'NA'
            match_pct_overall = round(((s+s_e) / (s + s_e  + u + u_e) * 100), 2) if (s + s_e + u + u_e) > 0 else 'NA'
            writer.writerow([adj, forms, s, s_e, u, u_e, f, total, match_pct_unelided, match_pct_overall])

def write_exceptions(log, filename):
    """Write exception log to text file"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.writelines(log)

# Create output directory if it doesn't exist
os.makedirs('rules_output', exist_ok=True)

# Main execution
print("Processing CSV files...")
data, adjectives, unspelled_log, unspelled_elided_log = process_csv_directory(base_csv_dir)

print(f"Found {len(data)} files to process")
print(f"Found {len(adjectives)} unique adjectives")
print(f"Found {len(unspelled_log)} unspelled exceptions")
print(f"Found {len(unspelled_elided_log)} unspelled elided exceptions")

# Write outputs
write_csv(data, 'rules_output/tag_analysis.csv')
write_adjectives(adjectives, 'rules_output/adjectives.csv')
write_exceptions(unspelled_log, 'rules_output/unspelled_oxford.txt')
#write_exceptions(unspelled_elided_log, 'rules_output/unspelled_elided_oxford.txt')

# Save Word document with exceptions
doc.save('rules_output/exceptions_isolated.docx')

print("Analysis complete! Check the rules_output directory for results.")
