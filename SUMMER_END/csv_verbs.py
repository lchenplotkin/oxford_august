#infinitive = v%inf
#preterite = v%pt
#past perticiple = v%ppl
#preterite plural = v%ptpl
#present participle = ger

'''
if infinitive, then ends in -e or -en or a vowel
if preterite plural, then ends in -e or -en
'''

# Re-import after code state reset
import os
import re
import string
import csv
import pandas as pd
from collections import defaultdict, Counter
import matplotlib.pyplot as plt
import docx

# Initialize Word document
doc = docx.Document()
doc.add_heading('Verb Rules Analysis in the Oxford Chaucer', 0)

# Configuration
base_csv_dir = 'data/csvs'

# Define verb tag patterns to search for
verb_tags = ['v%pr_1','v%pr_2','v%pr_3','v%pt_1','v%pt_2','v%pt_3','v%pr_pl','v%inf', 'v%pt', 'v%ppl', 'v%pt_pl', 'ger']

verb_dict = {}

# Utilities
def vowel_cluster_count(w):
    return len(re.findall(r'[^aeiouy]*[aeiouy]+(?:[^aeiouy]+(?=[^aeiouy]*[aeiouy])|[^aeiouy]*)?', w)) or 1

def parse_tagged_text(oxford_text, oxford_tagging):
    """Extract words from oxford_text and tags from oxford_tagging"""
    if pd.isna(oxford_text) or pd.isna(oxford_tagging) or oxford_text == '' or oxford_tagging == '':
        return [], [], []

    # Clean and split the text
    words = oxford_text.lower().replace(',', '').replace('.', '').replace('!', '').replace('?', '').replace('°', '').replace('¶', '').strip().split()

    # Parse tags from oxford_tagging
    tag_tokens = oxford_tagging.strip().split()
    tags = []
    headwords = []

    for token in tag_tokens:
        if '@' in token:
            if token not in ["--@dash", ".@ellipsis"]:
                parts = token.split('@')
                headword = parts[0].lower()
                tag_part = parts[1] if len(parts) > 1 else ''
                #tag_part = re.sub('r\d','',tag_part)
                #tag_part = tag_part.replace('_','')

                tag = tag_part

                headwords.append(headword)
                tags.append(tag)

    # Ensure all lists are the same length (trim to shortest)
    if len(words) != len(tags):
        print(oxford_text, oxford_tagging, words, tags)
        print(len(words), len(tags))
    min_len = min(len(words), len(tags), len(headwords))
    return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(tag, word, oxford_text, line_number, filename, rule_violated):
    """Add an exception case to the Word document"""
    doc_para = doc.add_paragraph()
    doc_para.add_run(f"{tag}: {word} - {rule_violated}").bold = True
    doc_para.add_run(f"\n{line_number} ({filename})\n")

    for i, ox_word in enumerate(oxford_text.split(' ')):
        if ox_word.lower().replace(',', '').replace('.', '').replace('!', '').replace('?', '').replace('°', '').replace('¶', '') == word.lower():
            doc_para.add_run(f"{ox_word}").italic = True
        else:
            doc_para.add_run(f"{ox_word}")

        if i < len(oxford_text.split(' ')) - 1:
            doc_para.add_run(" ")

def check_verb_rules(headword, word, tag):
    """Check if a word follows the expected verb rules"""
    global verb_dict
    word = word.lower().strip()

    if tag in ['v%ppl']:
        if headword not in verb_dict.keys():
            verb_dict[headword] = [{},{},{},0,0,0]

        if not (word.endswith('en') or word.endswith('e')):
            if word.endswith('t') or word.endswith('d'):
                verb_dict[headword][4]+=1
                if word not in verb_dict[headword][1].keys():
                    verb_dict[headword][1][word] = 0
                verb_dict[headword][1][word]+=1
            else:
                verb_dict[headword][5]+=1
                if word not in verb_dict[headword][2].keys():
                    verb_dict[headword][2][word] = 0
                verb_dict[headword][2][word]+=1
        else:
            verb_dict[headword][3]+=1
            if word not in verb_dict[headword][0].keys():
                verb_dict[headword][0][word] = 0
            verb_dict[headword][0][word]+=1

    '''
    if tag in ['v%pt_1','v%pt_3']:
        if headword not in verb_dict.keys():
            verb_dict[headword] = [{},{},{},0,0,0]

        if not word.endswith('n'):
            if word.endswith('t') or word.endswith('d'):
                verb_dict[headword][4]+=1
                if word not in verb_dict[headword][1].keys():
                    verb_dict[headword][1][word] = 0
                verb_dict[headword][1][word]+=1
            else:
                verb_dict[headword][5]+=1
                if word not in verb_dict[headword][2].keys():
                    verb_dict[headword][2][word] = 0
                verb_dict[headword][2][word]+=1
        else:
            verb_dict[headword][3]+=1
            if word not in verb_dict[headword][0].keys():
                verb_dict[headword][0][word] = 0
            verb_dict[headword][0][word]+=1
    '''


    if tag == 'v%inf':  # infinitive should end in -e, -en, or vowel
        if word.endswith('e') or word.endswith('en') or word[-1] in 'aeiouy':
            return True, "follows rule"
        else:
            return False, "infinitive should end in -e, -en, or vowel"

    elif tag == 'v%pt_pl':  # preterite plural should end in -e or -en
        if word.endswith('e') or word.endswith('en'):
            return True, "follows rule"
        else:
            return False, "preterite plural should end in -e or -en"
    elif tag == 'ger':
        if word.endswith('e'):
            return True, "follows rule"
        else:
            return False, "present participles (gerunds) should end in -e"
    elif tag == 'v%pr_pl':
        if word.endswith('e') or word.endswith('en'):
            return True, "follows rule"
        else:
            return False, "present plural should end in -e or -en"

    # For other verb forms, we don't have specific rules defined
    return True, "no rule specified"

def search_verb_patterns_csv(df, verbs, exception_log):
    """Search for verb patterns in CSV data"""
    rule_followers = defaultdict(int)
    rule_violators = defaultdict(int)
    exception_lines = []

    lines = 0
    for idx, row in df.iterrows():
        if row["MATCH"] != "DIFF":
            lines += 1
            oxford_text = row['OXFORD_TEXT']
            oxford_tagging = row['OXFORD_TAGGING']
            line_number = row['LINE_NUMBER']
            filename = row['OXFORD_FILENAME']

            words, headwords, tags = parse_tagged_text(oxford_text, oxford_tagging)

            for j, (word, headword, tag) in enumerate(zip(words, headwords, tags)):
                if tag in verb_tags:
                    follows_rule, rule_description = check_verb_rules(headword, word, tag)

                    verbs[headword]['forms'].append(word)
                    verbs[headword]['forms'] = list(set(verbs[headword]['forms']))

                    if follows_rule:
                        rule_followers[tag] += 1
                        verbs[headword]['following'] += 1
                    else:
                        rule_violators[tag] += 1
                        verbs[headword]['violating'] += 1

                        exception_text = f"{tag}: {word} - {rule_description}\nLine {line_number}: {oxford_text}\n\n"
                        exception_lines.append(exception_text)

                        # Add to Word document
                        add_exception_to_doc(tag, word, oxford_text, line_number, filename, rule_description)

    exception_log.extend(exception_lines)
    return lines, rule_followers, rule_violators

def process_csv_directory(csv_dir):
    """Process all _gui.csv files in the directory"""
    data = {}
    exception_log = []
    verbs = defaultdict(lambda: {'following': 0, 'violating': 0, 'forms': []})

    for root, dirs, files in os.walk(csv_dir):
        for file in files:
            if not file.endswith('_gui.csv'):
                continue

            csv_path = os.path.join(root, file)
            rel_path = os.path.relpath(root, csv_dir)
            file_id = os.path.join(rel_path, file) if rel_path != '.' else file

            #try:
            # Read CSV file
            df = pd.read_csv(csv_path, encoding='utf-8')

            # Skip if required columns are missing
            required_columns = ['OXFORD_TAGGING', 'OXFORD_TEXT', 'LINE_NUMBER', 'OXFORD_FILENAME']
            if not all(col in df.columns for col in required_columns):
                print(f"Skipping {file}: missing required columns")
                continue

            lines, rule_followers, rule_violators = search_verb_patterns_csv(df, verbs, exception_log)

            file_data = {
                'total_lines': lines,
                'total_following': sum(rule_followers.values()),
                'total_violating': sum(rule_violators.values()),
            }

            # Add per-tag statistics
            for tag in verb_tags:
                file_data[f'{tag}_following'] = rule_followers[tag]
                file_data[f'{tag}_violating'] = rule_violators[tag]
                total = rule_followers[tag] + rule_violators[tag]
                if total > 0:
                    file_data[f'{tag}_compliance_pct'] = round((rule_followers[tag] / total) * 100, 2)
                else:
                    file_data[f'{tag}_compliance_pct'] = 0.0

            # Overall compliance percentage
            total_verbs = file_data['total_following'] + file_data['total_violating']
            if total_verbs > 0:
                file_data['overall_compliance_pct'] = round((file_data['total_following'] / total_verbs) * 100, 2)
            else:
                file_data['overall_compliance_pct'] = 0.0

            data[file_id] = file_data

           # except Exception as e:
            #    print(f"Error processing {file}: {e}")
             #   continue

    return data, verbs, exception_log

def write_csv(data, filename):
    """Write analysis results to CSV"""
    if not data:
        print("No data to write")
        return

    keys = sorted(next(iter(data.values())).keys())
    with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=['file'] + keys)
        writer.writeheader()
        for file, stats in data.items():
            row = {'file': file}
            row.update(stats)
            writer.writerow(row)

def write_verbs(verbs, filename):
    """Write verb analysis to CSV"""
    with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['headword', 'forms', 'following_rules', 'violating_rules', 'total', 'compliance_pct'])
        for verb in sorted(verbs.keys()):
            following = verbs[verb]['following']
            violating = verbs[verb]['violating']
            forms = verbs[verb]['forms']
            total = following + violating
            compliance_pct = round((following / total * 100), 2) if total > 0 else 0.0
            writer.writerow([verb, forms, following, violating, total, compliance_pct])

def write_exceptions(log, filename):
    """Write exception log to text file"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.writelines(log)

# Create output directory if it doesn't exist
os.makedirs('verbs_output', exist_ok=True)

# Main execution
print("Processing CSV files for verb rules...")
data, verbs, exception_log = process_csv_directory(base_csv_dir)

print(f"Found {len(data)} files to process")
print(f"Found {len(verbs)} unique verbs")
print(f"Found {len(exception_log)} rule violations")

# Write outputs
write_csv(data, 'verbs_output/verb_analysis.csv')
write_verbs(verbs, 'verbs_output/verbs.csv')
write_exceptions(exception_log, 'verbs_output/exceptions.txt')

# Save Word document with exceptions
doc.save('verbs_output/exceptions.docx')

print("Verb analysis complete! Check the verbs_output directory for results.")
weak_verbs = {}
strong_verbs = {}
unclear_verbs = {}
conflict_verbs = {}
small_conflict_verbs = {}
unclear = 0
strong = 0
weak = 0
for verb in verb_dict.keys():
    if 0 not in [verb_dict[verb][3], verb_dict[verb][4]]:
        print(verb_dict[verb])

    verb_entry = verb_dict[verb]
    if verb_entry[5]>0 and verb_entry[3] == 0 and verb_entry[4] == 0:
        unclear_verbs[verb] = verb_entry
    elif verb_entry[3] == 0 and verb_entry[4] > 0:
        weak_verbs[verb] = verb_entry
    elif verb_entry[4] == 0 and verb_entry[3] > 0:
        strong_verbs[verb] = verb_entry
    elif verb_entry[5] >0 and (verb_entry[3] == 0 or verb_entry[4] == 0):
        small_conflict_verbs[verb] = verb_entry
    else:
        conflict_verbs[verb] = verb_entry

    strong += verb_entry[3]
    weak += verb_entry[4]
    unclear += verb_entry[5]

#print(verb_dict)
#print(strong_verbs)
print(len(strong_verbs))
print(len(weak_verbs))
print(len(unclear_verbs))
print(len(conflict_verbs))
print(len(small_conflict_verbs))
#print(conflict_verbs)
print(unclear_verbs)
print(strong,weak,unclear)
#print(verb_dict)
