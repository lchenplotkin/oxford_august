#!/usr/bin/env python3
"""
Export Middle English scansion to a .docx file with stress markings.
Usage: python export_scansion_docx.py input.txt output.docx
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scansion_final import scan_line, format_scansion, minimal_clean

def build_stress_line(text, marked):
    """
    Build stress notation line aligned with original text.
    For each character in original text, place '/' for stressed vowel, 'u' for unstressed, space otherwise.
    """
    stress_marks = []
    marked_idx = 0
    
    # Track which vowels we've seen in marked text
    vowel_positions = []
    
    # First pass: find all vowel positions and their stress in marked text
    i = 0
    char_pos = 0  # position in original text (ignoring spaces in marked)
    
    while i < len(marked):
        c = marked[i]
        
        if c == ' ':
            i += 1
            continue
        
        if c == '[':
            # Silent vowel - find the character and closing bracket
            i += 1
            if i < len(marked):
                vowel_char = marked[i]
                vowel_positions.append((char_pos, 'x', vowel_char.lower()))
                char_pos += 1
                # Skip to closing bracket
                while i < len(marked) and marked[i] != ']':
                    i += 1
            i += 1
            continue
        
        is_vowel = c.lower() in 'aeiouy'
        
        if is_vowel:
            if i>1:
                if marked[i-1]==c:
                    vowel_positions.append((char_pos, 'x', c.lower()))
            if c == c.upper():
                vowel_positions.append((char_pos, 'S', c.lower()))
            else:
                vowel_positions.append((char_pos, 'u', c.lower()))
            char_pos += 1
        else:
            # Consonant
            char_pos += 1
        
        i += 1
    
    # Second pass: build stress line aligned with original text
    vowel_idx = 0
    text_pos = 0
    
    for char in text:
        if char == ' ':
            stress_marks.append(' ')
            continue
        
        # Check if this position has a vowel
        found_mark = ' '
        
        if vowel_idx < len(vowel_positions):
            pos, stress, vowel_char = vowel_positions[vowel_idx]
            
            # Check if this is the vowel position
            if char.lower() in 'aeiouy':
                if stress == 'S':
                    found_mark = '/'
                elif stress == 'u':
                    found_mark = 'u'
                else:  # silent
                    found_mark = ' '
                vowel_idx += 1
        
        stress_marks.append(found_mark)
    
    return ''.join(stress_marks)

def create_scansion_docx(input_file, output_file):
    """Read lines from input file, scan them, and create a docx."""
    
    # Read input lines
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = [line.strip() for line in f if line.strip()]
    
    # Create document
    doc = Document()
    
    # Set default font to Courier New for monospace
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('Middle English Scansion', level=1)
    
    # Scan each line and add to document
    num = 0
    for line in lines:
        result = scan_line(line)
        
        if result:
            word_analyses, penalty = result
            formatted = format_scansion(word_analyses)
            
            # Build stress line
            stress_line = build_stress_line(line, formatted['marked_words'])
            
            # Add stress notation line
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(0)
            p1.paragraph_format.line_spacing = 1.0
            run1 = p1.add_run(stress_line)
            run1.font.name = 'Courier New'
            run1.font.size = Pt(11)
            
            # Add original text line
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(6)
            p2.paragraph_format.line_spacing = 1.0
            run2 = p2.add_run(line)
            run2.font.name = 'Courier New'
            run2.font.size = Pt(11)
            
            # Add metadata line
            p3 = doc.add_paragraph()
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(18)
            metadata_text = f"({formatted['syllable_count']} syllables, penalty: {penalty})"
            run3 = p3.add_run(metadata_text)
            run3.font.name = 'Arial'
            run3.font.size = Pt(9)
            run3.font.color.rgb = RGBColor(102, 102, 102)
        else:
            # Line failed to scan
            p1 = doc.add_paragraph()
            run1 = p1.add_run(f"FAILED TO SCAN: {line}")
            run1.font.name = 'Courier New'
            run1.font.size = Pt(11)
            run1.font.color.rgb = RGBColor(255, 0, 0)
            p1.paragraph_format.space_after = Pt(18)
    
        num+=1
        print(num, "lines done")
    # Save document
    doc.save(output_file)
    return True

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python export_scansion_docx.py input.txt output.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if create_scansion_docx(input_file, output_file):
        print(f"Scansion exported to {output_file}")
    else:
        sys.exit(1)
