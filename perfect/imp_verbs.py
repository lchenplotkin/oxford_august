import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict, Counter

# Initialize Word documents for exceptions
oxford_doc = docx.Document()
oxford_doc.add_heading('Oxford Verb Declension Exceptions', 0)

riverside_doc = docx.Document()
riverside_doc.add_heading('Riverside Verb Declension Exceptions', 0)

# Configuration
base_csv_dir = 'data/csvs'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
                    "hadden", "his", "her", "him", "hers", "hide", "hir",
                    "hire", "hires", "hirs", "han"]

form_csv = 'verb_forms_gold.csv'
verbs_dict = {}
with open(form_csv, 'r', encoding='utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        if 'headword' in row and 'classification' in row:
            verbs_dict[row['headword'].lower()] = row['classification'].lower()

def is_strong(verb):
    return verbs_dict.get(verb, '') == 'strong'

def is_weak(verb):
    return verbs_dict.get(verb, '') == 'weak'

def is_elided(word, next_word):
    if next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS:
        return True
    return False

def clean_tag(tag):
    return re.sub(r'\d+(?=%)', '', tag)

def parse_tagged_text(text, tagging, text_type):
    if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
        return [], [], []

    words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()
    tag_tokens = tagging.strip().split()
    tags = []
    headwords = []

    for token in tag_tokens:
        if '@' in token:
            parts = token.split('@')
            headword = parts[0].lower()
            tag = parts[1] if len(parts) > 1 else ''
            #tag = ''.join([i for i in tag if not i.isdigit()])
            headwords.append(headword)
            tags.append(tag)

    min_len = min(len(words), len(tags), len(headwords))
    return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(doc, exception_type, word, text, line_number, filename):
    doc_para = doc.add_paragraph()
    doc_para.add_run(f"{exception_type}: {word}").bold = True
    doc_para.add_run(f"\n{line_number} ({filename})\n")
    for i, ox_word in enumerate(text.split(' ')):
        if ox_word.lower() == word.lower():
            doc_para.add_run(f"{ox_word}").italic = True
        else:
            doc_para.add_run(f"{ox_word}")
        if i < len(text.split(' ')) - 1:
            doc_para.add_run(" ")
    doc_para.add_run("\n\n")

def classify_ending(word):
    if word.endswith('eth'):
        return '-eth'
    if word.endswith('en'):
        return '-en'
    if word.endswith('e'):
        return '-e'
    if word.endswith('ed'):
        return '-ed'
    if word.endswith('d'):
        return '-d'
    if word.endswith('t'):
        return '-t'
    if word.endswith(('a','e','i','o','u')):
        return 'vowel'
    return 'other'


def analyze_verbs(df, results, text_type, doc):
    """Analyze verb patterns in CSV data according to rules."""
    text_col = f'{text_type}_TEXT'
    tagging_col = f'{text_type}_TAGGING'
    filename_col = f'{text_type}_FILENAME'
    original_text_col = f'OG_{text_type}_TEXT'

    # Store all lines for later comparison
    if 'all_lines' not in results:
        results['all_lines'] = []

    for idx, row in df.iterrows():
        if row.get("MATCH") == "DIFF":
            continue

        text = row[text_col]
        tagging = row[tagging_col]
        line_number = row['LINE_NUMBER']
        filename = row[filename_col]
        original_text = row.get(original_text_col, text)

        # Store all lines for comparison
        results['all_lines'].append({
            'line_number': line_number,
            'text': original_text,
            'text_type': text_type
        })

        words, headwords, tags = parse_tagged_text(text, tagging, text_type)

        for j in range(len(tags)):
            if j >= len(words) or j >= len(headwords):
                continue

            word = words[j].lower()
            headword = headwords[j]
            tag = clean_tag(tags[j])
            next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'

            # Skip if elided
            if is_elided(word, next_word):
                continue


            ending = classify_ending(word)

            # Count ending distribution
            results['ending_counts'][tag][ending] += 1
            results['verb_tag_counts'][headword][tag][word] += 1

            # Rule checks
            # Rule checks
            violated = False
            reason = ""

            # Track rules individually
            def record_rule(rule_name, condition):
                if condition:
                    results['rule_stats'][rule_name]["success"] += 1
                else:
                    results['rule_stats'][rule_name]["fail"] += 1

            # Rules
            if tag == 'v%inf':
                record_rule("Infinitive ends in -en/-e", ending in ['-en', '-e'])
                if ending not in ['-en', '-e']:
                    violated = True
                    reason += "Infinitive must end in -en or -e "

            if tag == 'v%pt_pl':
                record_rule("Past plural ends in -en/-e", ending in ['-en', '-e'])
                if ending not in ['-en', '-e']:
                    violated = True
                    reason += "Past plural must end in -en or -e "

            if tag == 'v%pr_3':
                record_rule("Present 3rd sg ends in -eth", ending == '-eth')
                if ending != '-eth':
                    violated = True
                    reason += "Present 3rd sg must end in -eth "

            if tag == 'v%pr_pl':
                record_rule("Present plural ends in -en/-e", ending in ['-en', '-e'])
                if ending not in ['-en', '-e']:
                    violated = True
                    reason += "Present plural must end in -en or -e "

            if headword in verbs_dict:
                if is_strong(headword) and tag in ['v%pt_1', 'v%pt_3']:
                    record_rule("Strong pt sg not -en/-e", ending not in ['-en', '-e'])
                    if ending in ['-en', '-e']:
                        violated = True
                        reason += "Strong pt sg must not end in -en or -e "

                if is_weak(headword) and tag in ['v%pt_1', 'v%pt_3']:
                    record_rule("Weak pt sg ends in -ed/-d/-t", ending in ['-ed','-d','-t'])
                    if ending not in ['-ed','-d','-t']:
                        violated = True
                        reason += "Weak pt sg must end in -ed, -d, or -t "

                if is_strong(headword) and tag == 'v%ppl':
                    record_rule("Strong participle ends in -en/-e", ending in ['-en','-e'])
                    if ending not in ['-en', '-e']:
                        violated = True
                        reason += "Strong participle must end in -en or -e "

            if violated:
                record = {
                    'headword': headword,
                    'word': word,
                    'tag': tag,
                    'line_number': line_number,
                    'filename': filename,
                    'context': original_text,
                    'text_type': text_type,
                    'reason': reason.strip()
                }
                results['exceptions'].append(record)
                add_exception_to_doc(doc, reason, word, original_text, line_number, filename)

    return results

def process_csv_directory(csv_dir, text_type, doc):
    results = {
        'exceptions': [],
        'ending_counts': defaultdict(Counter),
        'verb_tag_counts': defaultdict(lambda: defaultdict(Counter)),
        'all_lines': [],
        'rule_stats': defaultdict(lambda: {"success": 0, "fail": 0})
    }

    file_count = 0
    for root, dirs, files in os.walk(csv_dir):
        for file in files:
            if not file.endswith('_gui.csv'):
                continue
            file_count += 1
            try:
                df = pd.read_csv(os.path.join(root, file), encoding='utf-8')
                required_columns = [f'{text_type}_TAGGING', f'{text_type}_TEXT', 'LINE_NUMBER', f'{text_type}_FILENAME']
                if not all(col in df.columns for col in required_columns):
                    print(f"Skipping {file} for {text_type}: missing columns")
                    continue
                results = analyze_verbs(df, results, text_type, doc)
            except Exception as e:
                print(f"Error processing {file} for {text_type}: {e}")
                continue
    print(f"Processed {file_count} files for {text_type}")
    return results

def create_record_key(record):
    """Create a unique key for a record based on line number and headword"""
    return f"{record['line_number']}_{record['headword']}"

def find_corresponding_line(line_number, all_lines, target_text_type):
    """Find the corresponding line in the other text version based on line number and filename"""
    for line in all_lines:
        if (line['line_number'] == line_number and
            line['text_type'] == target_text_type):
            return line
    return None

def compare_results(oxford_results, riverside_results):
    """Compare Oxford and Riverside results to find common and unique instances"""
    comparison = {
        'oxford_only': [],
        'riverside_only': [],
        'both_texts': []
    }

    # Create sets of record keys for comparison
    oxford_keys = set()
    riverside_keys = set()
    oxford_records = {}
    riverside_records = {}

    for record in oxford_results['exceptions']:
        key = create_record_key(record)
        oxford_keys.add(key)
        oxford_records[key] = record

    for record in riverside_results['exceptions']:
        key = create_record_key(record)
        riverside_keys.add(key)
        riverside_records[key] = record

    # Find common and unique records
    common_keys = oxford_keys.intersection(riverside_keys)
    oxford_only_keys = oxford_keys - riverside_keys
    riverside_only_keys = riverside_keys - oxford_keys

    comparison['oxford_only'] = [oxford_records[key] for key in oxford_only_keys]
    comparison['riverside_only'] = [riverside_records[key] for key in riverside_only_keys]
    comparison['both_texts'] = [(oxford_records[key], riverside_records[key]) for key in common_keys]

    return comparison

def write_results(results, output_dir, text_type):
    os.makedirs(output_dir, exist_ok=True)

    # Exceptions
    with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type', 'reason']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for rec in results['exceptions']:
            writer.writerow(rec)

    # Distribution by tag
    with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Tag', 'Ending', 'Count', 'Percent'])
        for tag, counts in results['ending_counts'].items():
            total = sum(counts.values())
            for ending, count in counts.items():
                percent = (count / total) * 100 if total > 0 else 0
                writer.writerow([tag, ending, count, f"{percent:.2f}%"])

    # Distribution by verb+tag
    with open(os.path.join(output_dir, f'{text_type.lower()}_verb_tag_distribution.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Headword', 'Tag', 'Ending', 'Count', 'Percent'])
        for verb, tags in results['verb_tag_counts'].items():
            for tag, counts in tags.items():
                total = sum(counts.values())
                for ending, count in counts.items():
                    percent = (count / total) * 100 if total > 0 else 0
                    writer.writerow([verb, tag, ending, count, f"{percent:.2f}%"])

    # Rule success/failure rates
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule_success_rates.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Rule', 'Successes', 'Failures', 'Total', 'SuccessRate'])
        for rule, counts in results['rule_stats'].items():
            total = counts['success'] + counts['fail']
            success_rate = (counts['success'] / total * 100) if total > 0 else 0
            writer.writerow([rule, counts['success'], counts['fail'], total, f"{success_rate:.2f}%"])



# --- ANALYSIS + RESULT WRITING functions here (same as before, unchanged) ---
# (keeping for brevity – your existing `analyze_verbs`, `process_csv_directory`,
#  `write_results`, `write_summary_docx`, etc. stay as-is)

# --- KEY CHANGE: Better X_only / Y_only formatting in Word ---
def write_comparison_results(comparison, output_dir, oxford_results, riverside_results):
    os.makedirs(output_dir, exist_ok=True)

    oxford_only_doc = docx.Document()
    oxford_only_doc.add_heading('Oxford-Only Verb Declension Exceptions', 0)

    riverside_only_doc = docx.Document()
    riverside_only_doc.add_heading('Riverside-Only Verb Declension Exceptions', 0)

    both_texts_doc = docx.Document()
    both_texts_doc.add_heading('Verb Declension Exceptions in Both Texts', 0)

    def write_comparison_records(records, filename, record_type, doc=None):
        with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type', 'reason']
            if record_type == 'both':
                fieldnames.extend(['oxford_context', 'riverside_context', 'oxford_reason', 'riverside_reason'])
            elif record_type == 'oxford_only':
                fieldnames.extend(['riverside_context', 'riverside_reason'])
            elif record_type == 'riverside_only':
                fieldnames.extend(['oxford_context', 'oxford_reason'])

            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()

            if record_type == 'both':
                for oxford_record, riverside_record in records:
                    combined_record = oxford_record.copy()
                    combined_record['oxford_context'] = oxford_record['context']
                    combined_record['riverside_context'] = riverside_record['context']
                    combined_record['oxford_reason'] = oxford_record['reason']
                    combined_record['riverside_reason'] = riverside_record['reason']
                    combined_record['text_type'] = 'BOTH'
                    writer.writerow(combined_record)

                    if doc:
                        doc_para = doc.add_paragraph()
                        doc_para.add_run(f"Exception: {oxford_record['word']}").bold = True
                        doc_para.add_run(f"\nHeadword: {oxford_record['headword']}")
                        doc_para.add_run(f"\nTag: {oxford_record['tag']}")
                        doc_para.add_run(f"\nLine: {oxford_record['line_number']}")
                        doc_para.add_run(f"\nOxford Filename: {oxford_record['filename']}")
                        doc_para.add_run(f"\nRiverside Filename: {riverside_record['filename']}")
                        doc_para.add_run(f"\nOxford Reason: {oxford_record['reason']}")
                        doc_para.add_run(f"\nRiverside Reason: {riverside_record['reason']}")
                        doc_para.add_run(f"\nOxford Context: {oxford_record['context']}")
                        doc_para.add_run(f"\nRiverside Context: {riverside_record['context']}")
                        doc_para.add_run("\n" + "="*50 + "\n")
            else:
                for record in records:
                    other_line = find_corresponding_line(
                        record['line_number'],
                        riverside_results['all_lines'] if record_type == 'oxford_only' else oxford_results['all_lines'],
                        "RIVERSIDE" if record_type == 'oxford_only' else "OXFORD"
                    )

                    csv_record = record.copy()
                    if other_line:
                        if record_type == 'oxford_only':
                            csv_record['riverside_context'] = other_line['text']
                            csv_record['riverside_reason'] = "No exception (normal form)"
                        else:
                            csv_record['oxford_context'] = other_line['text']
                            csv_record['oxford_reason'] = "No exception (normal form)"
                    writer.writerow(csv_record)

                    if doc:
                        # Pretty formatting: same as individual exception files
                        doc_para = doc.add_paragraph()
                        doc_para.add_run(f"{record['reason']}: {record['word']}").bold = True
                        doc_para.add_run(f"\n{record['line_number']} ({record['filename']})\n")
                        for i, token in enumerate(record['context'].split(' ')):
                            if token.lower() == record['word'].lower():
                                doc_para.add_run(token).italic = True
                            else:
                                doc_para.add_run(token)
                            if i < len(record['context'].split(' ')) - 1:
                                doc_para.add_run(" ")

                        if other_line:
                            other_text_type_display = "Riverside" if record_type == 'oxford_only' else "Oxford"
                            doc_para.add_run(f"\n\n{other_text_type_display} line: {other_line['text']}")
                            doc_para.add_run(f"\n{other_text_type_display} Reason: No exception (normal form)")
                        else:
                            other_text_type_display = "Riverside" if record_type == 'oxford_only' else "Oxford"
                            doc_para.add_run(f"\n\n{other_text_type_display} Context: Line not found")
                            doc_para.add_run(f"\n{other_text_type_display} Reason: Line not found")

                        doc_para.add_run("\n\n")

    write_comparison_records(comparison['oxford_only'], 'oxford_only_exceptions.csv', 'oxford_only', oxford_only_doc)
    write_comparison_records(comparison['riverside_only'], 'riverside_only_exceptions.csv', 'riverside_only', riverside_only_doc)
    write_comparison_records(comparison['both_texts'], 'both_texts_exceptions.csv', 'both', both_texts_doc)

    oxford_only_doc.save(os.path.join(output_dir, 'oxford_only_exceptions.docx'))
    riverside_only_doc.save(os.path.join(output_dir, 'riverside_only_exceptions.docx'))
    both_texts_doc.save(os.path.join(output_dir, 'both_texts_exceptions.docx'))

# --- rest of the program (comparison summary writing, main execution) is unchanged ---
def write_summary_docx(results, output_dir, text_type):
    summary_doc = docx.Document()
    summary_doc.add_heading(f'{text_type} Verb Declension Analysis Summary', 0)

    summary_doc.add_heading('Exceptions', 1)
    summary_doc.add_paragraph(f"Total exceptions found: {len(results['exceptions'])}")

    summary_doc.add_heading('Ending Distribution by Tag', 1)
    for tag, counts in results['ending_counts'].items():
        para = summary_doc.add_paragraph()
        para.add_run(f"{tag}: ").bold = True
        total = sum(counts.values())
        for ending, count in counts.items():
            percent = (count / total) * 100 if total > 0 else 0
            para.add_run(f"{ending}={count} ({percent:.1f}%)  ")

    summary_doc.add_heading('Rule Success Rates', 1)
    for rule, counts in results['rule_stats'].items():
        total = counts['success'] + counts['fail']
        success_rate = (counts['success'] / total * 100) if total > 0 else 0
        para = summary_doc.add_paragraph()
        para.add_run(f"{rule}: ").bold = True
        para.add_run(f"{counts['success']} / {total} successes ({success_rate:.1f}%)")

    summary_doc.save(os.path.join(output_dir, f'{text_type.lower()}_analysis_summary.docx'))

def write_comparison_summary_docx(comparison, output_dir):
    """Write comparison summary to Word document"""
    comp_doc = docx.Document()
    comp_doc.add_heading('Oxford vs Riverside Verb Comparison Summary', 0)

    comp_doc.add_heading('Exception Comparison Statistics', 1)

    oxford_count = len(comparison['oxford_only'])
    riverside_count = len(comparison['riverside_only'])
    both_count = len(comparison['both_texts'])
    total_oxford = oxford_count + both_count
    total_riverside = riverside_count + both_count

    para = comp_doc.add_paragraph()
    para.add_run("Exception Counts:").bold = True
    para.add_run(f"\n  Oxford only: {oxford_count}")
    para.add_run(f"\n  Riverside only: {riverside_count}")
    para.add_run(f"\n  Both texts: {both_count}")
    para.add_run(f"\n  Total Oxford exceptions: {total_oxford}")
    para.add_run(f"\n  Total Riverside exceptions: {total_riverside}\n")

    comp_doc.add_heading('Rule Success Rate Comparison', 1)

    # Gather all rules from both texts
    all_rules = set(oxford_results['rule_stats'].keys()).union(set(riverside_results['rule_stats'].keys()))

    table = comp_doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rule'
    hdr_cells[1].text = 'Oxford Successes / Total (%)'
    hdr_cells[2].text = 'Riverside Successes / Total (%)'
    hdr_cells[3].text = 'Oxford Rate (%)'
    hdr_cells[4].text = 'Riverside Rate (%)'

    for rule in sorted(all_rules):
        ox_counts = oxford_results['rule_stats'].get(rule, {"success":0,"fail":0})
        rv_counts = riverside_results['rule_stats'].get(rule, {"success":0,"fail":0})

        ox_total = ox_counts['success'] + ox_counts['fail']
        rv_total = rv_counts['success'] + rv_counts['fail']

        ox_rate = (ox_counts['success'] / ox_total * 100) if ox_total > 0 else 0
        rv_rate = (rv_counts['success'] / rv_total * 100) if rv_total > 0 else 0

        row_cells = table.add_row().cells
        row_cells[0].text = rule
        row_cells[1].text = f"{ox_counts['success']} / {ox_total}"
        row_cells[2].text = f"{rv_counts['success']} / {rv_total}"
        row_cells[3].text = f"{ox_rate:.1f}%"
        row_cells[4].text = f"{rv_rate:.1f}%"


    # Add detailed breakdown
    comp_doc.add_heading('Detailed Breakdown', 1)

    # Count exceptions by type for each category
    def count_by_reason(records):
        reason_counts = Counter()
        for record in records:
            reason_counts[record['reason']] += 1
        return reason_counts

    oxford_reasons = count_by_reason(comparison['oxford_only'])
    riverside_reasons = count_by_reason(comparison['riverside_only'])
    both_reasons = Counter()
    for oxford_record, riverside_record in comparison['both_texts']:
        both_reasons[oxford_record['reason']] += 1

    comp_doc.add_heading('Oxford-Only Exceptions by Type', 2)
    for reason, count in oxford_reasons.most_common():
        comp_doc.add_paragraph(f"{reason}: {count}")

    comp_doc.add_heading('Riverside-Only Exceptions by Type', 2)
    for reason, count in riverside_reasons.most_common():
        comp_doc.add_paragraph(f"{reason}: {count}")

    comp_doc.add_heading('Both-Texts Exceptions by Type', 2)
    for reason, count in both_reasons.most_common():
        comp_doc.add_paragraph(f"{reason}: {count}")

    comp_doc.save(os.path.join(output_dir, 'comparison_summary.docx'))

# Main execution
if __name__ == "__main__":
    print("Processing Oxford texts...")
    oxford_results = process_csv_directory(base_csv_dir, 'OXFORD', oxford_doc)

    print("Processing Riverside texts...")
    riverside_results = process_csv_directory(base_csv_dir, 'RIVERSIDE', riverside_doc)

    oxford_output_dir = 'oxford_verb_analysis_output'
    riverside_output_dir = 'riverside_verb_analysis_output'
    comparison_output_dir = 'comparison_verb_analysis_output'

    # Write individual results
    write_results(oxford_results, oxford_output_dir, 'OXFORD')
    write_summary_docx(oxford_results, oxford_output_dir, 'OXFORD')
    oxford_doc.save(os.path.join(oxford_output_dir, 'oxford_declension_exceptions.docx'))

    write_results(riverside_results, riverside_output_dir, 'RIVERSIDE')
    write_summary_docx(riverside_results, riverside_output_dir, 'RIVERSIDE')
    riverside_doc.save(os.path.join(riverside_output_dir, 'riverside_declension_exceptions.docx'))

    # Compare results and write comparison
    print("Comparing Oxford and Riverside results...")
    comparison = compare_results(oxford_results, riverside_results)
    write_comparison_results(comparison, comparison_output_dir, oxford_results, riverside_results)
    write_comparison_summary_docx(comparison, comparison_output_dir)

    # Write completion logs
    for output_dir, text_type, results in [
        (oxford_output_dir, 'OXFORD', oxford_results),
        (riverside_output_dir, 'RIVERSIDE', riverside_results)
    ]:
        with open(os.path.join(output_dir, f'{text_type.lower()}_analysis_log.txt'), 'w', encoding='utf-8') as log_file:
            log_file.write(f"Chaucer {text_type} Verb Declension Analysis - Complete\n")
            log_file.write("="*50 + "\n\n")

            log_file.write(f"Total verb exceptions found: {len(results['exceptions'])}\n\n")

            log_file.write("Files generated in this directory:\n")
            files = [
                f"{text_type.lower()}_exceptions.csv - All verb exception instances",
                f"{text_type.lower()}_ending_distribution.csv - Ending distribution by tag",
                f"{text_type.lower()}_verb_tag_distribution.csv - Verb-tag ending distribution",
                f"{text_type.lower()}_analysis_summary.docx - Summary statistics document",
                f"{text_type.lower()}_declension_exceptions.docx - Detailed exception examples",
                f"{text_type.lower()}_analysis_log.txt - This log file"
            ]

            for file_desc in files:
                log_file.write(f"- {file_desc}\n")

    # Write comparison log
    with open(os.path.join(comparison_output_dir, 'comparison_analysis_log.txt'), 'w', encoding='utf-8') as log_file:
        log_file.write("Oxford vs Riverside Verb Comparison Analysis - Complete\n")
        log_file.write("="*50 + "\n\n")

        log_file.write("Comparison Summary:\n")
        log_file.write(f"Oxford Only Exceptions: {len(comparison['oxford_only'])}\n")
        log_file.write(f"Riverside Only Exceptions: {len(comparison['riverside_only'])}\n")
        log_file.write(f"Exceptions in Both Texts: {len(comparison['both_texts'])}\n")

        log_file.write("\nFiles generated in this directory:\n")
        comparison_files = [
            "oxford_only_exceptions.csv - Exceptions found only in Oxford",
            "riverside_only_exceptions.csv - Exceptions found only in Riverside",
            "both_texts_exceptions.csv - Exceptions found in both texts",
            "oxford_only_exceptions.docx - Oxford-only exceptions in Word format",
            "riverside_only_exceptions.docx - Riverside-only exceptions in Word format",
            "both_texts_exceptions.docx - Both-texts exceptions in Word format",
            "comparison_summary.csv - Overall comparison statistics",
            "comparison_summary.docx - Detailed comparison summary document",
            "comparison_analysis_log.txt - This log file"
        ]

        for file_desc in comparison_files:
            log_file.write(f"- {file_desc}\n")

    print("\nAnalysis complete!")
    print(f"Oxford results: {oxford_output_dir}/")
    print(f"Riverside results: {riverside_output_dir}/")
    print(f"Comparison results: {comparison_output_dir}/")
