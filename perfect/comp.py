# Merged Chaucer Adjective Declension Analysis - Oxford and Riverside
import docx 
import os
import re
import string
import csv
import pandas as pd
from collections import defaultdict, Counter
import matplotlib.pyplot as plt

# Initialize Word documents for exceptions
oxford_doc = docx.Document()
oxford_doc.add_heading('Oxford Adjective Declension Exceptions', 0)

riverside_doc = docx.Document()
riverside_doc.add_heading('Riverside Adjective Declension Exceptions', 0)

# Configuration
base_csv_dir = 'data/csvs'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
                    "hadden", "his", "her", "him", "hers", "hide", "hir",
                    "hire", "hires", "hirs", "han"]

monosyllabic_set = set()
mono_csv = 'monosyllabic_adjectives.csv'
with open(mono_csv,'r',encoding = 'utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        # Assume the CSV has a column named 'headword' or similar
        if 'headword' in row:
            monosyllabic_set.add(row['headword'].lower())
        elif 'Headword' in row:
            monosyllabic_set.add(row['Headword'].lower())

# Utilities
def vowel_cluster_count(w):
    """Count vowel clusters to determine syllables"""
    return len(re.findall(r'[^aeiouy]*[aeiouy]+(?:[^aeiouy]+(?=[^aeiouy]*[aeiouy])|[^aeiouy]*)?', w)) or 1

def is_elided(word, next_word):
    """Check if word is elided before next word"""
    if next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS:
        return True
    return False

def is_weak_form(prev_tag, prev_word, tag, next_tag):
    """Check if adjective is in weak declension context"""
    weak_triggers = ['demonstrative', 'def_art', 'n%gen', 'pron%gen', 'interj', 'pron%fem_gen']
    for cause in weak_triggers:
        if prev_tag == cause:
            return True
    if next_tag == 'n#propn':
        return True
    return False

def is_plural_form(prev_tag, tag, next_tag):
    """Check if adjective modifies plural noun"""
    if next_tag.startswith('n%pl') or prev_tag.startswith('n%pl'):
        return True
    return False

def parse_tagged_text(text, tagging, text_type):
    """Extract words from text and tags from tagging"""
    if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
        return [], [], []

    # Clean and split the text
    words = text.lower().replace(',', '').replace('.', '').replace('!', '').replace('?', '').replace('°', '').replace('¶', '').strip().split()

    # Parse tags from tagging
    tag_tokens = tagging.strip().split()
    tags = []
    headwords = []

    for token in tag_tokens:
        if '@' in token:
            if token not in ["--@dash", ".@ellipsis"]:
                parts = token.split('@')
                headword = parts[0].lower()
                tag_part = parts[1] if len(parts) > 1 else ''

                tag = tag_part
                tag = ''.join([i for i in tag if not i.isdigit()])
                headwords.append(headword)
                tags.append(tag)

    # Handle special demonstrative cases
    for i, (word, tag) in enumerate(zip(words[:len(tags)], tags)):
        if word in ['this', 'that', 'thilke'] and 'gram_adj' in tag:
            tags[i] = 'demonstrative'

    # Ensure all lists are the same length (trim to shortest)
    if len(words) != len(tags):
        print(f"Length mismatch ({text_type}): {text} | {tagging}")
        print(f"Words: {len(words)}, Tags: {len(tags)}")
    min_len = min(len(words), len(tags), len(headwords))
    return words[:min_len], headwords[:min_len], tags[:min_len]

def is_monosyllabic_root(headword, word_forms):
    if headword in monosyllabic_set:
        return True
    else:
        return False

def add_exception_to_doc(doc, exception_type, word, text, line_number, filename):
    """Add an exception case to the Word document"""
    doc_para = doc.add_paragraph()
    doc_para.add_run(f"{exception_type}: {word}").bold = True
    doc_para.add_run(f"\n{line_number} ({filename})\n")
    
    for i, ox_word in enumerate(text.split(' ')):
        if ox_word.lower() == word.lower():
            doc_para.add_run(f"{ox_word}").italic = True
        else:
            doc_para.add_run(f"{ox_word}")
        if i < len(text.split(' ')) - 1:
            doc_para.add_run(" ")
    doc_para.add_run("\n\n")

def analyze_adjectives(df, results, text_type, doc):
    """Analyze adjective patterns in CSV data"""
    
    # Column names based on text type
    text_col = f'{text_type}_TEXT'
    tagging_col = f'{text_type}_TAGGING'
    filename_col = f'{text_type}_FILENAME'
    original_text_col = f'OG_{text_type}_TEXT'
    
    for idx, row in df.iterrows():
        if row["MATCH"] != "DIFF":
            text = row[text_col]
            tagging = row[tagging_col]
            line_number = row['LINE_NUMBER']
            filename = row[filename_col]
            original_text = row.get(original_text_col, text)

            words, headwords, tags = parse_tagged_text(text, tagging, text_type)

            prev_tag = 'NA'
            prev_word = 'NA'
            
            for j in range(len(tags)):
                if j >= len(words) or j >= len(headwords):
                    continue
                    
                headword = headwords[j]
                tag = tags[j]
                word = words[j].lower()
                next_tag = tags[j + 1] if j + 1 < len(tags) else 'END'
                next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'

                # Only process adjectives
                if tag != 'adj': 
                    prev_word = word
                    prev_tag = tag
                    continue

                # Initialize adjective tracking
                if headword not in results['all_adjectives']:
                    results['all_adjectives'][headword] = {
                        'stem_forms': set(), 
                        'all_forms': set(),
                        'stem_with_e': set(),
                        'stem_without_e': set()
                    }

                results['all_adjectives'][headword]['all_forms'].add(word)
                if tag == 'adj':  # Stem form
                    results['all_adjectives'][headword]['stem_forms'].add(word)
                    if word.endswith('e'):
                        results['all_adjectives'][headword]['stem_with_e'].add(word)
                    else:
                        results['all_adjectives'][headword]['stem_without_e'].add(word)

                # Check if monosyllabic root
                stem_forms = results['all_adjectives'][headword]['stem_forms']
                if not is_monosyllabic_root(headword, stem_forms):
                    prev_word = word
                    prev_tag = tag
                    continue

                # Determine grammatical context
                is_weak = is_weak_form(prev_tag, prev_word, tag, next_tag)
                is_plural = is_plural_form(prev_tag, tag, next_tag)
                is_final = (next_word == 'END')
                is_word_elided = is_elided(word, next_word) if next_word != 'END' else False

                # Create record for this instance
                record = {
                    'headword': headword,
                    'word': word,
                    'line_number': line_number,
                    'filename': filename,
                    'context': original_text,
                    'is_elided': is_word_elided,
                    'is_final': is_final,
                    'text_type': text_type
                }

                # Categorize the adjective usage
                if is_weak:
                    if not word.endswith('e'):
                        # Weak declension without -e (exception)
                        results['weak_no_e_all'].append(record)
                        if not is_word_elided and not is_final:
                            results['weak_no_e_strict'].append(record)
                            add_exception_to_doc(doc, "Weak without -e", word, original_text, line_number, filename)
                elif is_plural:
                    if not word.endswith('e'):
                        # Plural form without -e (exception)
                        results['plural_no_e_all'].append(record)
                        if not is_word_elided and not is_final:
                            results['plural_no_e_strict'].append(record)
                            add_exception_to_doc(doc, "Plural without -e", word, original_text, line_number, filename)
                else:
                    # Strong/stem form - but exclude adjectives that always end in -e in strong form
                    if word.endswith('e') and tag == 'adj':
                        # Only count as exception if this adjective also has stem forms without -e
                        # We'll filter this in post-processing after collecting all data
                        results['strong_with_e_all'].append(record)
                        if not is_word_elided and not is_final:
                            results['strong_with_e_strict'].append(record)

                prev_word = word
                prev_tag = tag

    return results

def process_csv_directory(csv_dir, text_type, doc):
    """Process all _gui.csv files in the directory for a specific text type"""
    results = {
        'all_adjectives': {},
        'weak_no_e_all': [],          
        'weak_no_e_strict': [],       
        'plural_no_e_all': [],        
        'plural_no_e_strict': [],     
        'strong_with_e_all': [],      
        'strong_with_e_strict': []    
    }

    file_count = 0
    for root, dirs, files in os.walk(csv_dir):
        for file in files:
            if not file.endswith('_gui.csv'):
                continue

            csv_path = os.path.join(root, file)
            file_count += 1

            try:
                # Read CSV file
                df = pd.read_csv(csv_path, encoding='utf-8')

                # Column names based on text type
                text_col = f'{text_type}_TEXT'
                tagging_col = f'{text_type}_TAGGING'
                filename_col = f'{text_type}_FILENAME'
                
                # Skip if required columns are missing
                required_columns = [tagging_col, text_col, 'LINE_NUMBER', filename_col]
                if not all(col in df.columns for col in required_columns):
                    print(f"Skipping {file} for {text_type}: missing required columns")
                    continue

                results = analyze_adjectives(df, results, text_type, doc)

            except Exception as e:
                print(f"Error processing {file} for {text_type}: {e}")
                continue

    print(f"Processed {file_count} files for {text_type}")
    return results

def filter_strong_form_exceptions(results):
    """Filter out adjectives that always end in -e in strong form from strong form exceptions"""
    # Identify adjectives that always end in -e in strong form
    always_e_adjectives = set()
    for headword, data in results['all_adjectives'].items():
        if is_monosyllabic_root(headword, data['stem_forms']):
            # If this adjective has stem forms and ALL stem forms end in -e, it's an "always -e" adjective
            if data['stem_forms'] and len(data['stem_without_e']) == 0:
                always_e_adjectives.add(headword)
    
    # Filter the strong form exceptions to exclude always-e adjectives
    results['strong_with_e_all_filtered'] = [
        record for record in results['strong_with_e_all'] 
        if record['headword'] not in always_e_adjectives
    ]
    results['strong_with_e_strict_filtered'] = [
        record for record in results['strong_with_e_strict'] 
        if record['headword'] not in always_e_adjectives
    ]
    
    # Add the always-e adjectives to results for reporting
    results['always_e_adjectives'] = always_e_adjectives
    
    return results

def create_record_key(record):
    """Create a unique key for a record based on line number and filename"""
    return f"{record['line_number']}_{record['headword']}"
    #return f"{record['filename']}_{record['line_number']}_{record['headword']}_{record['word']}"

def compare_results(oxford_results, riverside_results):
    """Compare Oxford and Riverside results to find common and unique instances"""
    comparison = {
        'oxford_only': {},
        'riverside_only': {},
        'both_texts': {}
    }
    
    categories = ['weak_no_e_all', 'weak_no_e_strict', 'plural_no_e_all', 
                  'plural_no_e_strict', 'strong_with_e_all', 'strong_with_e_strict',
                  'strong_with_e_all_filtered', 'strong_with_e_strict_filtered']
    
    for category in categories:
        # Create sets of record keys for comparison
        oxford_keys = set()
        riverside_keys = set()
        oxford_records = {}
        riverside_records = {}
        
        for record in oxford_results.get(category, []):
            key = create_record_key(record)
            oxford_keys.add(key)
            oxford_records[key] = record
            
        for record in riverside_results.get(category, []):
            key = create_record_key(record)
            riverside_keys.add(key)
            riverside_records[key] = record
        
        # Find common and unique records
        common_keys = oxford_keys.intersection(riverside_keys)
        oxford_only_keys = oxford_keys - riverside_keys
        riverside_only_keys = riverside_keys - oxford_keys
        
        comparison['oxford_only'][category] = [oxford_records[key] for key in oxford_only_keys]
        comparison['riverside_only'][category] = [riverside_records[key] for key in riverside_only_keys]
        comparison['both_texts'][category] = [(oxford_records[key], riverside_records[key]) for key in common_keys]
    
    return comparison

def write_results_to_csv(results, output_dir, text_type):
    """Write analysis results to CSV files"""
    os.makedirs(output_dir, exist_ok=True)
    
    # Filter strong form exceptions first
    results = filter_strong_form_exceptions(results)
    
    # Helper function to write a list of records to CSV
    def write_records(records, filename):
        if not records:
            # Create empty file with headers
            with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final', 'text_type']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
            return
            
        with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final', 'text_type']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            for record in records:
                writer.writerow(record)

    # Write each category
    write_records(results['weak_no_e_all'], f'{text_type.lower()}_weak_no_e_all_instances.csv')
    write_records(results['weak_no_e_strict'], f'{text_type.lower()}_weak_no_e_strict_instances.csv')
    write_records(results['plural_no_e_all'], f'{text_type.lower()}_plural_no_e_all_instances.csv')
    write_records(results['plural_no_e_strict'], f'{text_type.lower()}_plural_no_e_strict_instances.csv')
    write_records(results['strong_with_e_all'], f'{text_type.lower()}_strong_with_e_all_instances.csv')
    write_records(results['strong_with_e_strict'], f'{text_type.lower()}_strong_with_e_strict_instances.csv')
    write_records(results['strong_with_e_all_filtered'], f'{text_type.lower()}_strong_with_e_all_filtered_instances.csv')
    write_records(results['strong_with_e_strict_filtered'], f'{text_type.lower()}_strong_with_e_strict_filtered_instances.csv')

    # Write summary statistics
    with open(os.path.join(output_dir, f'{text_type.lower()}_summary_statistics.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Category', 'All Instances', 'Strict Instances (not elided, not final)'])
        writer.writerow(['Weak declension without -e', len(results['weak_no_e_all']), len(results['weak_no_e_strict'])])
        writer.writerow(['Plural form without -e', len(results['plural_no_e_all']), len(results['plural_no_e_strict'])])
        writer.writerow(['Strong form with -e (all)', len(results['strong_with_e_all']), len(results['strong_with_e_strict'])])
        writer.writerow(['Strong form with -e (filtered)', len(results['strong_with_e_all_filtered']), len(results['strong_with_e_strict_filtered'])])
        writer.writerow(['Always -e adjectives found', len(results['always_e_adjectives']), 'N/A'])

    # Write monosyllabic adjectives found
    with open(os.path.join(output_dir, f'{text_type.lower()}_monosyllabic_adjectives.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Headword', 'Stem Forms', 'All Forms', 'Always Ends in E'])
        for headword, data in results['all_adjectives'].items():
            if is_monosyllabic_root(headword, data['stem_forms']):
                always_e = 'Yes' if headword in results['always_e_adjectives'] else 'No'
                writer.writerow([headword, '; '.join(sorted(data['stem_forms'])), '; '.join(sorted(data['all_forms'])), always_e])

    # Write list of always -e adjectives
    with open(os.path.join(output_dir, f'{text_type.lower()}_always_e_adjectives.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Headword', 'Stem Forms with -e', 'Reason'])
        for headword in sorted(results['always_e_adjectives']):
            data = results['all_adjectives'][headword]
            stem_forms = '; '.join(sorted(data['stem_with_e']))
            writer.writerow([headword, stem_forms, 'All stem forms end in -e'])

def write_comparison_results(comparison, output_dir):
    """Write comparison results to CSV files"""
    os.makedirs(output_dir, exist_ok=True)
    
    categories = ['weak_no_e_all', 'weak_no_e_strict', 'plural_no_e_all', 
                  'plural_no_e_strict', 'strong_with_e_all', 'strong_with_e_strict',
                  'strong_with_e_all_filtered', 'strong_with_e_strict_filtered']
    
    def write_comparison_records(records, filename, record_type):
        with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final', 'text_type']
            if record_type == 'both':
                fieldnames.extend(['oxford_context', 'riverside_context'])
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            
            if record_type == 'both':
                for oxford_record, riverside_record in records:
                    combined_record = oxford_record.copy()
                    combined_record['oxford_context'] = oxford_record['context']
                    combined_record['riverside_context'] = riverside_record['context']
                    combined_record['text_type'] = 'BOTH'
                    writer.writerow(combined_record)
            else:
                for record in records:
                    writer.writerow(record)
    
    # Write comparison files for each category
    for category in categories:
        write_comparison_records(comparison['oxford_only'][category], f'oxford_only_{category}.csv', 'oxford_only')
        write_comparison_records(comparison['riverside_only'][category], f'riverside_only_{category}.csv', 'riverside_only')
        write_comparison_records(comparison['both_texts'][category], f'both_texts_{category}.csv', 'both')

    # Write summary comparison statistics
    with open(os.path.join(output_dir, 'comparison_summary.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Category', 'Oxford Only', 'Riverside Only', 'Both Texts'])
        for category in categories:
            oxford_count = len(comparison['oxford_only'][category])
            riverside_count = len(comparison['riverside_only'][category])
            both_count = len(comparison['both_texts'][category])
            writer.writerow([category, oxford_count, riverside_count, both_count])

def write_summary_to_docx(results, output_dir, text_type):
    """Write summary statistics to Word document"""
    # Filter strong form exceptions first
    results = filter_strong_form_exceptions(results)
    
    summary_doc = docx.Document()
    summary_doc.add_heading(f'{text_type} Adjective Declension Analysis Summary', 0)
    
    monosyllabic_count = sum(1 for headword, data in results['all_adjectives'].items() 
                           if is_monosyllabic_root(headword, data['stem_forms']))
    
    # Overview section
    summary_doc.add_heading('Overview', 1)
    para = summary_doc.add_paragraph()
    para.add_run(f"Total adjectives found: {len(results['all_adjectives'])}\n")
    para.add_run(f"Monosyllabic adjectives found: {monosyllabic_count}\n")
    para.add_run(f"Always -e adjectives found: {len(results['always_e_adjectives'])}")
    
    # Statistics section
    summary_doc.add_heading('Exception Statistics', 1)
    
    categories = [
        ('Weak declension without -e', 'weak_no_e_all', 'weak_no_e_strict'),
        ('Plural form without -e', 'plural_no_e_all', 'plural_no_e_strict'),
        ('Strong form with -e (all)', 'strong_with_e_all', 'strong_with_e_strict'),
        ('Strong form with -e (filtered)', 'strong_with_e_all_filtered', 'strong_with_e_strict_filtered')
    ]
    
    for desc, all_key, strict_key in categories:
        all_count = len(results[all_key])
        strict_count = len(results[strict_key])
        
        para = summary_doc.add_paragraph()
        para.add_run(desc + ":").bold = True
        para.add_run(f"\n  All instances: {all_count}")
        para.add_run(f"\n  Strict instances (not elided, not final): {strict_count}\n")
    
    # Always -e adjectives section
    if results['always_e_adjectives']:
        summary_doc.add_heading('Adjectives That Always End in -e in Strong Form', 1)
        para = summary_doc.add_paragraph()
        para.add_run("The following adjectives were excluded from strong form exception analysis because they always end in -e in the strong form:\n\n")
        
        for headword in sorted(results['always_e_adjectives']):
            data = results['all_adjectives'][headword]
            stem_forms = ', '.join(sorted(data['stem_with_e']))
            para.add_run(f"• {headword}").bold = True
            para.add_run(f" (forms: {stem_forms})\n")
    
    summary_doc.save(os.path.join(output_dir, f'{text_type.lower()}_analysis_summary.docx'))

def write_comparison_summary_docx(comparison, output_dir):
    """Write comparison summary to Word document"""
    comp_doc = docx.Document()
    comp_doc.add_heading('Oxford vs Riverside Comparison Summary', 0)
    
    categories = ['weak_no_e_all', 'weak_no_e_strict', 'plural_no_e_all', 
                  'plural_no_e_strict', 'strong_with_e_all', 'strong_with_e_strict',
                  'strong_with_e_all_filtered', 'strong_with_e_strict_filtered']
    
    comp_doc.add_heading('Comparison Statistics', 1)
    
    for category in categories:
        oxford_count = len(comparison['oxford_only'][category])
        riverside_count = len(comparison['riverside_only'][category])
        both_count = len(comparison['both_texts'][category])
        total_oxford = oxford_count + both_count
        total_riverside = riverside_count + both_count
        
        para = comp_doc.add_paragraph()
        para.add_run(f"{category.replace('_', ' ').title()}:").bold = True
        para.add_run(f"\n  Oxford only: {oxford_count}")
        para.add_run(f"\n  Riverside only: {riverside_count}")
        para.add_run(f"\n  Both texts: {both_count}")
        para.add_run(f"\n  Total Oxford: {total_oxford}")
        para.add_run(f"\n  Total Riverside: {total_riverside}\n")
    
    comp_doc.save(os.path.join(output_dir, 'comparison_summary.docx'))

# Main execution
if __name__ == "__main__":
    # Process Oxford files
    print("Processing Oxford texts...")
    oxford_results = process_csv_directory(base_csv_dir, 'OXFORD', oxford_doc)
    
    # Process Riverside files
    print("Processing Riverside texts...")
    riverside_results = process_csv_directory(base_csv_dir, 'RIVERSIDE', riverside_doc)
    
    # Create output directories
    oxford_output_dir = 'oxford_analysis_output'
    riverside_output_dir = 'riverside_analysis_output'
    comparison_output_dir = 'comparison_analysis_output'
    
    # Write Oxford results
    write_results_to_csv(oxford_results, oxford_output_dir, 'OXFORD')
    write_summary_to_docx(oxford_results, oxford_output_dir, 'OXFORD')
    oxford_doc.save(os.path.join(oxford_output_dir, 'oxford_declension_exceptions.docx'))
    
    # Write Riverside results
    write_results_to_csv(riverside_results, riverside_output_dir, 'RIVERSIDE')
    write_summary_to_docx(riverside_results, riverside_output_dir, 'RIVERSIDE')
    riverside_doc.save(os.path.join(riverside_output_dir, 'riverside_declension_exceptions.docx'))
    
    # Compare results
    print("Comparing Oxford and Riverside results...")
    comparison = compare_results(oxford_results, riverside_results)
    write_comparison_results(comparison, comparison_output_dir)
    write_comparison_summary_docx(comparison, comparison_output_dir)
    
    # Write completion logs
    for output_dir, text_type, results in [
        (oxford_output_dir, 'OXFORD', oxford_results), 
        (riverside_output_dir, 'RIVERSIDE', riverside_results)
    ]:
        with open(os.path.join(output_dir, f'{text_type.lower()}_analysis_log.txt'), 'w', encoding='utf-8') as log_file:
            log_file.write(f"Chaucer {text_type} Adjective Declension Analysis - Complete\n")
            log_file.write("="*50 + "\n\n")
            
            monosyllabic_count = sum(1 for headword, data in results['all_adjectives'].items() 
                                   if is_monosyllabic_root(headword, data['stem_forms']))
            
            # Filter results for logging
            results = filter_strong_form_exceptions(results)
            
            log_file.write(f"Total adjectives found: {len(results['all_adjectives'])}\n")
            log_file.write(f"Monosyllabic adjectives found: {monosyllabic_count}\n")
            log_file.write(f"Always -e adjectives found: {len(results['always_e_adjectives'])}\n\n")
            
            log_file.write("Files generated in this directory:\n")
            files = [
                f"{text_type.lower()}_weak_no_e_all_instances.csv - All weak declension forms without -e",
                f"{text_type.lower()}_weak_no_e_strict_instances.csv - Strict weak declension exceptions", 
                f"{text_type.lower()}_plural_no_e_all_instances.csv - All plural forms without -e",
                f"{text_type.lower()}_plural_no_e_strict_instances.csv - Strict plural exceptions",
                f"{text_type.lower()}_strong_with_e_all_instances.csv - All strong forms with -e",
                f"{text_type.lower()}_strong_with_e_strict_instances.csv - Strict strong form exceptions",
                f"{text_type.lower()}_strong_with_e_all_filtered_instances.csv - Strong forms with -e (filtered)",
                f"{text_type.lower()}_strong_with_e_strict_filtered_instances.csv - Strict strong exceptions (filtered)",
                f"{text_type.lower()}_summary_statistics.csv - Overall counts",
                f"{text_type.lower()}_monosyllabic_adjectives.csv - Monosyllabic adjectives found",
                f"{text_type.lower()}_always_e_adjectives.csv - Adjectives that always end in -e",
                f"{text_type.lower()}_analysis_summary.docx - Summary statistics document",
                f"{text_type.lower()}_declension_exceptions.docx - Detailed exception examples",
                f"{text_type.lower()}_analysis_log.txt - This log file"
            ]
            
            for file_desc in files:
                log_file.write(f"- {file_desc}\n")
    
    # Write comparison log
    with open(os.path.join(comparison_output_dir, 'comparison_analysis_log.txt'), 'w', encoding='utf-8') as log_file:
        log_file.write("Oxford vs Riverside Comparison Analysis - Complete\n")
        log_file.write("="*50 + "\n\n")
        
        categories = ['weak_no_e_all', 'weak_no_e_strict', 'plural_no_e_all', 
                      'plural_no_e_strict', 'strong_with_e_all', 'strong_with_e_strict',
                      'strong_with_e_all_filtered', 'strong_with_e_strict_filtered']
        
        log_file.write("Comparison Summary:\n")
        for category in categories:
            oxford_count = len(comparison['oxford_only'][category])
            riverside_count = len(comparison['riverside_only'][category])
            both_count = len(comparison['both_texts'][category])
            log_file.write(f"{category}: Oxford={oxford_count}, Riverside={riverside_count}, Both={both_count}\n")
        
        log_file.write("\nFiles generated in this directory:\n")
        comparison_files = [
            "comparison_summary.csv - Overall comparison statistics",
            "comparison_summary.docx - Detailed comparison summary document"
        ]
        
        for category in categories:
            comparison_files.extend([
                f"oxford_only_{category}.csv - Exceptions found only in Oxford",
                f"riverside_only_{category}.csv - Exceptions found only in Riverside", 
                f"both_texts_{category}.csv - Exceptions found in both texts"
            ])
        
        comparison_files.append("comparison_analysis_log.txt - This log file")
        
        for file_desc in comparison_files:
            log_file.write(f"- {file_desc}\n")
    
    print("\nAnalysis complete!")
    print(f"Oxford results: {oxford_output_dir}/")
    print(f"Riverside results: {riverside_output_dir}/")
    print(f"Comparison results: {comparison_output_dir}/")
