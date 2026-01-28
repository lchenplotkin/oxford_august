import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict, Counter

# Initialize Word documents for exceptions
oxford_doc = docx.Document()
oxford_doc.add_heading('Oxford Adverb Analysis', 0)

riverside_doc = docx.Document()
riverside_doc.add_heading('Riverside Adverb Analysis', 0)

combined_doc = docx.Document()
combined_doc.add_heading('Combined Adverb Analysis', 0)

# Configuration
base_csv_dir = 'dataset'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde", "hadden",
                     "his", "her", "him", "hers", "hide", "hir", "hire", "hires", "hirs", "han"]

# Common adverbs that are EXCLUDED from the -e rule
COMMON_ADVERBS = ['ful', 'wel', 'ly']

form_csv = 'verb_forms_simple.csv'
verbs_dict = {}

try:
    with open(form_csv, 'r', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            if 'headword' in row and 'classification' in row:
                verbs_dict[row['headword'].lower()] = row['classification'].lower()
except FileNotFoundError:
    print(f"Warning: {form_csv} not found. Continuing without verb classification.")

def is_elided(word, next_word):
    """Check if word is elided before next word"""
    if next_word and (next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS):
        return True
    return False

def clean_tag(tag):
    """Remove digits before % in a tag (e.g., v2%pt_1 -> v%pt_1)."""
    return re.sub(r'\d+(?=%)', '', tag)

def parse_tagged_text(text, tagging, text_type):
    """Extract words from text and tags from tagging"""
    if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
        return [], [], []
    
    words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()
    tag_tokens = tagging.strip().split()
    
    tags = []
    headwords = []
    
    for token in tag_tokens:
        if '@' in token:
            parts = token.split('@')
            headword = parts[0].lower()
            tag = parts[1] if len(parts) > 1 else ''
            headwords.append(headword)
            tags.append(tag)
    
    min_len = min(len(words), len(tags), len(headwords))
    return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(doc, exception_type, word, text, line_number, filename):
    """Add an exception case to the Word document"""
    doc_para = doc.add_paragraph()
    doc_para.add_run(f"{exception_type}: {word}").bold = True
    doc_para.add_run(f"\n{line_number} ({filename})\n")
    
    for i, ox_word in enumerate(text.split(' ')):
        if ox_word.lower() == word.lower():
            doc_para.add_run(f"{ox_word}").italic = True
        else:
            doc_para.add_run(f"{ox_word}")
        if i < len(text.split(' ')) - 1:
            doc_para.add_run(" ")
    
    doc_para.add_run("\n\n")

def is_common_adverb(word):
    """Check if word ends with common adverb suffixes that are excluded from -e rule"""
    for suffix in COMMON_ADVERBS:
        if word.endswith(suffix):
            return True
    return False

def analyze_adverbs(df, results, text_type, doc, current_filename):
    """Analyze adverb patterns in CSV data according to rules."""
    text_col = f'{text_type}_TEXT'
    tagging_col = f'{text_type}_TAGGING'
    filename_col = f'{text_type}_FILENAME'
    original_text_col = f'OG_{text_type}_TEXT'
    scansion_col = f'{text_type}_SCANSION'
    
    for idx, row in df.iterrows():
        text = row[text_col]
        tagging = row[tagging_col]
        line_number = row['LINE_NUMBER']
        filename = row[filename_col]
        original_text = row.get(original_text_col, text)
        scansion = row.get(scansion_col, '')
        
        words, headwords, tags = parse_tagged_text(text, tagging, text_type)
        
        for j in range(len(tags)):
            if j >= len(words) or j >= len(headwords):
                continue
            
            word = words[j].lower()
            headword = headwords[j]
            tag = clean_tag(tags[j])
            next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'
            
            # Skip if elided
            if is_elided(word, next_word):
                continue
            
            # RULE 13: Adverbs (except common ones) should end in -e
            if tag.startswith('adv'):
                # Track all adverb endings
                results['all_adverb_counts'][word] += 1
                results['file_all_adverb_counts'][current_filename][word] += 1
                
                # Check if it's not a common adverb
                if not is_common_adverb(word):
                    if word.endswith('e'):
                        results['adverb_with_e'] += 1
                        results['file_adverb_with_e'][current_filename] += 1
                    else:
                        results['adverb_without_e'] += 1
                        results['file_adverb_without_e'][current_filename] += 1
                        
                        # This is an exception - adverb without -e
                        record = {
                            'headword': headword,
                            'word': word,
                            'tag': tag,
                            'line_number': line_number,
                            'filename': filename,
                            'context': original_text,
                            'text_type': text_type,
                            'source_file': current_filename,
                            'rule': 'RULE_13_ADV_NO_E'
                        }
                        results['exceptions'].append(record)
                        results['file_exceptions'][current_filename].append(record)
                        add_exception_to_doc(doc, "Adverb without -e", word, original_text, line_number, filename)
            
            # RULE 14: Track 'lich' vs 'liche' endings
            if 'lich' in word:
                if word.endswith('liche'):
                    results['liche_count'] += 1
                    results['file_liche_count'][current_filename] += 1
                elif word.endswith('lich'):
                    results['lich_count'] += 1
                    results['file_lich_count'][current_filename] += 1
                    
                    # Exception: lich without -e
                    record = {
                        'headword': headword,
                        'word': word,
                        'tag': tag,
                        'line_number': line_number,
                        'filename': filename,
                        'context': original_text,
                        'text_type': text_type,
                        'source_file': current_filename,
                        'rule': 'RULE_14_LICH_NO_E'
                    }
                    results['exceptions'].append(record)
                    results['file_exceptions'][current_filename].append(record)
                    add_exception_to_doc(doc, "lich without -e", word, original_text, line_number, filename)
            
            # RULE 15: 'where' should be monosyllabic (using scansion)
            if headword == 'where' or word == 'where':
                # Parse scansion to count syllables for this position
                if not pd.isna(scansion) and scansion:
                    # Extract syllables for this word position
                    # Scansion format might be like: "/ x / x" where / and x are syllables
                    scansion_parts = scansion.strip().split()
                    if j < len(scansion_parts):
                        syllable_marker = scansion_parts[j]
                        # Count syllables - each character that's not a space is a syllable marker
                        syllable_count = len(syllable_marker.replace(' ', ''))
                        
                        results['where_syllable_counts'][syllable_count] += 1
                        results['file_where_syllable_counts'][current_filename][syllable_count] += 1
                        
                        if syllable_count > 1:
                            # Exception: where is polysyllabic
                            record = {
                                'headword': headword,
                                'word': word,
                                'tag': tag,
                                'line_number': line_number,
                                'filename': filename,
                                'context': original_text,
                                'text_type': text_type,
                                'source_file': current_filename,
                                'rule': 'RULE_15_WHERE_POLYSYLLABIC',
                                'syllables': syllable_count,
                                'scansion': scansion
                            }
                            results['exceptions'].append(record)
                            results['file_exceptions'][current_filename].append(record)
                            add_exception_to_doc(doc, f"'where' with {syllable_count} syllables", word, original_text, line_number, filename)
            
            # RULE 16: Track 'than' vs 'thanne'
            if headword in ['than', 'thanne'] or word in ['than', 'thanne']:
                if word == 'than':
                    results['than_count'] += 1
                    results['file_than_count'][current_filename] += 1
                elif word == 'thanne':
                    results['thanne_count'] += 1
                    results['file_thanne_count'][current_filename] += 1
    
    return results

def process_csv_directory(csv_dir, text_type, doc):
    results = {
        'exceptions': [],
        'file_exceptions': defaultdict(list),
        
        # Rule 13: Adverbs with/without -e
        'adverb_with_e': 0,
        'adverb_without_e': 0,
        'all_adverb_counts': Counter(),
        'file_adverb_with_e': defaultdict(int),
        'file_adverb_without_e': defaultdict(int),
        'file_all_adverb_counts': defaultdict(Counter),
        
        # Rule 14: lich vs liche
        'lich_count': 0,
        'liche_count': 0,
        'file_lich_count': defaultdict(int),
        'file_liche_count': defaultdict(int),
        
        # Rule 15: where syllable counts
        'where_syllable_counts': Counter(),
        'file_where_syllable_counts': defaultdict(Counter),
        
        # Rule 16: than vs thanne
        'than_count': 0,
        'thanne_count': 0,
        'file_than_count': defaultdict(int),
        'file_thanne_count': defaultdict(int)
    }
    
    file_count = 0
    for root, dirs, files in os.walk(csv_dir):
        for file in files:
            if not file.endswith('_gui_complete.csv'):
                continue
            
            file_count += 1
            try:
                df = pd.read_csv(os.path.join(root, file), encoding='utf-8')
                required_columns = [f'{text_type}_TAGGING', f'{text_type}_TEXT', 'LINE_NUMBER', f'{text_type}_FILENAME']
                
                if not all(col in df.columns for col in required_columns):
                    print(f"Skipping {file} for {text_type}: missing columns")
                    continue
                
                results = analyze_adverbs(df, results, text_type, doc, file)
            except Exception as e:
                print(f"Error processing {file} for {text_type}: {e}")
                continue
    
    print(f"Processed {file_count} files for {text_type}")
    return results

def write_results(results, output_dir, text_type):
    os.makedirs(output_dir, exist_ok=True)
    
    # Exceptions - aggregated
    with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['rule', 'headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type', 'source_file', 'syllables', 'scansion']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames, extrasaction='ignore')
        writer.writeheader()
        for rec in results['exceptions']:
            writer.writerow(rec)
    
    # Exceptions - by file
    with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['source_file', 'exception_count', 'rule', 'headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames, extrasaction='ignore')
        writer.writeheader()
        for source_file, exceptions in results['file_exceptions'].items():
            for rec in exceptions:
                writer.writerow({
                    'source_file': source_file,
                    'exception_count': len(exceptions),
                    **rec
                })
    
    # Rule 13: Adverb -e distribution
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule13_adverb_e_summary.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Category', 'Count', 'Percent'])
        total = results['adverb_with_e'] + results['adverb_without_e']
        if total > 0:
            writer.writerow(['Adverbs with -e', results['adverb_with_e'], f"{results['adverb_with_e']/total*100:.2f}%"])
            writer.writerow(['Adverbs without -e', results['adverb_without_e'], f"{results['adverb_without_e']/total*100:.2f}%"])
            writer.writerow(['Total', total, '100%'])
    
    # Rule 13: By file
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule13_adverb_e_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Source_File', 'With_E', 'Without_E', 'Total', 'Percent_With_E'])
        all_files = set(results['file_adverb_with_e'].keys()) | set(results['file_adverb_without_e'].keys())
        for source_file in sorted(all_files):
            with_e = results['file_adverb_with_e'][source_file]
            without_e = results['file_adverb_without_e'][source_file]
            total = with_e + without_e
            percent = (with_e / total * 100) if total > 0 else 0
            writer.writerow([source_file, with_e, without_e, total, f"{percent:.2f}%"])
    
    # Rule 13: All adverb word counts
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule13_all_adverbs.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Adverb', 'Count', 'Has_E'])
        for word, count in results['all_adverb_counts'].most_common():
            has_e = 'Yes' if word.endswith('e') else 'No'
            writer.writerow([word, count, has_e])
    
    # Rule 14: lich/liche distribution
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule14_lich_summary.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Form', 'Count', 'Percent'])
        total = results['lich_count'] + results['liche_count']
        if total > 0:
            writer.writerow(['liche (with -e)', results['liche_count'], f"{results['liche_count']/total*100:.2f}%"])
            writer.writerow(['lich (without -e)', results['lich_count'], f"{results['lich_count']/total*100:.2f}%"])
            writer.writerow(['Total', total, '100%'])
    
    # Rule 14: By file
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule14_lich_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Source_File', 'liche', 'lich', 'Total', 'Percent_liche'])
        all_files = set(results['file_liche_count'].keys()) | set(results['file_lich_count'].keys())
        for source_file in sorted(all_files):
            liche = results['file_liche_count'][source_file]
            lich = results['file_lich_count'][source_file]
            total = liche + lich
            percent = (liche / total * 100) if total > 0 else 0
            writer.writerow([source_file, liche, lich, total, f"{percent:.2f}%"])
    
    # Rule 15: where syllable counts
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule15_where_syllables.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Syllable_Count', 'Occurrences', 'Percent'])
        total = sum(results['where_syllable_counts'].values())
        if total > 0:
            for syllables in sorted(results['where_syllable_counts'].keys()):
                count = results['where_syllable_counts'][syllables]
                writer.writerow([syllables, count, f"{count/total*100:.2f}%"])
            writer.writerow(['Total', total, '100%'])
    
    # Rule 15: By file
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule15_where_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Source_File', 'Syllable_Count', 'Occurrences'])
        for source_file in sorted(results['file_where_syllable_counts'].keys()):
            for syllables, count in sorted(results['file_where_syllable_counts'][source_file].items()):
                writer.writerow([source_file, syllables, count])
    
    # Rule 16: than/thanne distribution
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule16_than_summary.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Form', 'Count', 'Percent'])
        total = results['than_count'] + results['thanne_count']
        if total > 0:
            writer.writerow(['than (normal)', results['than_count'], f"{results['than_count']/total*100:.2f}%"])
            writer.writerow(['thanne (emphatic)', results['thanne_count'], f"{results['thanne_count']/total*100:.2f}%"])
            writer.writerow(['Total', total, '100%'])
    
    # Rule 16: By file
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule16_than_by_file.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Source_File', 'than', 'thanne', 'Total', 'Percent_than'])
        all_files = set(results['file_than_count'].keys()) | set(results['file_thanne_count'].keys())
        for source_file in sorted(all_files):
            than = results['file_than_count'][source_file]
            thanne = results['file_thanne_count'][source_file]
            total = than + thanne
            percent = (than / total * 100) if total > 0 else 0
            writer.writerow([source_file, than, thanne, total, f"{percent:.2f}%"])

def write_summary_docx(results, output_dir, text_type):
    summary_doc = docx.Document()
    summary_doc.add_heading(f'{text_type} Adverb Analysis Summary', 0)
    
    # Rule 13
    summary_doc.add_heading('Rule 13: Adverbs (except common) end in -e', 1)
    total = results['adverb_with_e'] + results['adverb_without_e']
    if total > 0:
        summary_doc.add_paragraph(f"Adverbs with -e: {results['adverb_with_e']} ({results['adverb_with_e']/total*100:.1f}%)")
        summary_doc.add_paragraph(f"Adverbs without -e: {results['adverb_without_e']} ({results['adverb_without_e']/total*100:.1f}%)")
    
    # Rule 14
    summary_doc.add_heading('Rule 14: lich should be liche', 1)
    total = results['lich_count'] + results['liche_count']
    if total > 0:
        summary_doc.add_paragraph(f"liche (with -e): {results['liche_count']} ({results['liche_count']/total*100:.1f}%)")
        summary_doc.add_paragraph(f"lich (without -e): {results['lich_count']} ({results['lich_count']/total*100:.1f}%)")
    
    # Rule 15
    summary_doc.add_heading('Rule 15: where is monosyllabic', 1)
    total_where = sum(results['where_syllable_counts'].values())
    if total_where > 0:
        for syllables in sorted(results['where_syllable_counts'].keys()):
            count = results['where_syllable_counts'][syllables]
            summary_doc.add_paragraph(f"{syllables} syllable(s): {count} ({count/total_where*100:.1f}%)")
    
    # Rule 16
    summary_doc.add_heading('Rule 16: than (normal) vs thanne (emphatic)', 1)
    total = results['than_count'] + results['thanne_count']
    if total > 0:
        summary_doc.add_paragraph(f"than: {results['than_count']} ({results['than_count']/total*100:.1f}%)")
        summary_doc.add_paragraph(f"thanne: {results['thanne_count']} ({results['thanne_count']/total*100:.1f}%)")
    
    # Exceptions summary
    summary_doc.add_heading('Total Exceptions', 1)
    summary_doc.add_paragraph(f"Total exceptions found: {len(results['exceptions'])}")
    
    summary_doc.save(os.path.join(output_dir, f'{text_type.lower()}_analysis_summary.docx'))

def write_combined_results(oxford_results, riverside_results, output_dir):
    """Create combined analysis showing overlaps and differences"""
    os.makedirs(output_dir, exist_ok=True)
    
    # Combined summary
    with open(os.path.join(output_dir, 'combined_summary.txt'), 'w', encoding='utf-8') as f:
        f.write("COMBINED ADVERB ANALYSIS SUMMARY\n")
        f.write("=" * 70 + "\n\n")
        
        f.write("RULE 13: Adverbs (except common) end in -e\n")
        f.write("-" * 70 + "\n")
        ox_total = oxford_results['adverb_with_e'] + oxford_results['adverb_without_e']
        rs_total = riverside_results['adverb_with_e'] + riverside_results['adverb_without_e']
        if ox_total > 0:
            f.write(f"Oxford: {oxford_results['adverb_with_e']}/{ox_total} with -e ({oxford_results['adverb_with_e']/ox_total*100:.1f}%)\n")
        if rs_total > 0:
            f.write(f"Riverside: {riverside_results['adverb_with_e']}/{rs_total} with -e ({riverside_results['adverb_with_e']/rs_total*100:.1f}%)\n")
        f.write("\n")
        
        f.write("RULE 14: lich should be liche\n")
        f.write("-" * 70 + "\n")
        ox_total = oxford_results['lich_count'] + oxford_results['liche_count']
        rs_total = riverside_results['lich_count'] + riverside_results['liche_count']
        if ox_total > 0:
            f.write(f"Oxford: {oxford_results['liche_count']}/{ox_total} liche ({oxford_results['liche_count']/ox_total*100:.1f}%)\n")
        if rs_total > 0:
            f.write(f"Riverside: {riverside_results['liche_count']}/{rs_total} liche ({riverside_results['liche_count']/rs_total*100:.1f}%)\n")
        f.write("\n")
        
        f.write("RULE 15: where is monosyllabic\n")
        f.write("-" * 70 + "\n")
        ox_mono = oxford_results['where_syllable_counts'].get(1, 0)
        ox_total = sum(oxford_results['where_syllable_counts'].values())
        rs_mono = riverside_results['where_syllable_counts'].get(1, 0)
        rs_total = sum(riverside_results['where_syllable_counts'].values())
        if ox_total > 0:
            f.write(f"Oxford: {ox_mono}/{ox_total} monosyllabic ({ox_mono/ox_total*100:.1f}%)\n")
        if rs_total > 0:
            f.write(f"Riverside: {rs_mono}/{rs_total} monosyllabic ({rs_mono/rs_total*100:.1f}%)\n")
        f.write("\n")
        
        f.write("RULE 16: than (normal) vs thanne (emphatic)\n")
        f.write("-" * 70 + "\n")
        ox_total = oxford_results['than_count'] + oxford_results['thanne_count']
        rs_total = riverside_results['than_count'] + riverside_results['thanne_count']
        if ox_total > 0:
            f.write(f"Oxford: {oxford_results['than_count']}/{ox_total} than ({oxford_results['than_count']/ox_total*100:.1f}%)\n")
        if rs_total > 0:
            f.write(f"Riverside: {riverside_results['than_count']}/{rs_total} than ({riverside_results['than_count']/rs_total*100:.1f}%)\n")
        f.write("\n")
        
        f.write(f"Total Oxford exceptions: {len(oxford_results['exceptions'])}\n")
        f.write(f"Total Riverside exceptions: {len(riverside_results['exceptions'])}\n")
    
    # Combined doc
    combined_doc.add_heading('Summary Statistics', 1)
    combined_doc.add_paragraph(f"Total Oxford exceptions: {len(oxford_results['exceptions'])}")
    combined_doc.add_paragraph(f"Total Riverside exceptions: {len(riverside_results['exceptions'])}")

# Main execution
if __name__ == "__main__":
    print("Processing Oxford texts...")
    oxford_results = process_csv_directory(base_csv_dir, 'OXFORD', oxford_doc)
    
    print("Processing Riverside texts...")
    riverside_results = process_csv_directory(base_csv_dir, 'RIVERSIDE', riverside_doc)
    
    oxford_output_dir = 'oxford_adverb_analysis_output'
    riverside_output_dir = 'riverside_adverb_analysis_output'
    combined_output_dir = 'combined_adverb_analysis_output'
    
    write_results(oxford_results, oxford_output_dir, 'OXFORD')
    write_summary_docx(oxford_results, oxford_output_dir, 'OXFORD')
    oxford_doc.save(os.path.join(oxford_output_dir, 'oxford_adverb_exceptions.docx'))
    
    write_results(riverside_results, riverside_output_dir, 'RIVERSIDE')
    write_summary_docx(riverside_results, riverside_output_dir, 'RIVERSIDE')
    riverside_doc.save(os.path.join(riverside_output_dir, 'riverside_adverb_exceptions.docx'))
    
    print("\nCreating combined analysis...")
    write_combined_results(oxford_results, riverside_results, combined_output_dir)
    combined_doc.save(os.path.join(combined_output_dir, 'combined_adverb_exceptions.docx'))
    
    print("\nAnalysis complete!")
    print(f"Oxford results: {oxford_output_dir}/")
    print(f"Riverside results: {riverside_output_dir}/")
    print(f"Combined results: {combined_output_dir}/")
