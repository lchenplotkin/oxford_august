import docx
import os
import re
import csv
import pandas as pd

# Initialize Word document for exceptions
doc = docx.Document()
doc.add_heading('Filtered Adjective Declension Exceptions in the Oxford Chaucer', 0)

# Configuration
base_csv_dir = '../dataset'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
                     "hadden", "his", "her", "him", "hers", "hide", "hir",
                     "hire", "hires", "hirs", "han"]

# -----------------------------
# Load Monosyllabic Adjectives
# -----------------------------
monosyllabic_set = set()
with open('monosyllabic_adjectives.csv','r',encoding='utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        key = 'headword' if 'headword' in row else 'Headword'
        monosyllabic_set.add(row[key].lower())

# -----------------------------
# Load Always-E Adjectives FIRST
# -----------------------------
def load_always_e_adjectives(filepath='always_e_adjectives.csv'):
    always_e = set()
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                always_e.add(row['Headword'].lower())
    return always_e

always_e_adjectives_global = load_always_e_adjectives()

# -----------------------------
# Utilities
# -----------------------------
def is_elided(word, next_word, next_tag):
    return next_word[0] in 'aeiou' or \
           (next_word[0]=='h' and next_tag.startswith('pron')) or \
           next_word in ELISION_FOLLOWERS

def is_weak_form(prev_tag, prev_word, tag, next_tag):
    weak_triggers = ['demonstrative', 'def_art', 'n%gen', 'pron%gen', 'interj', 'pron%fem_gen']
    return prev_tag in weak_triggers or next_tag == 'n#propn'

def is_plural_form(prev_tag, tag, next_tag):
    return next_tag.startswith('n%pl') or prev_tag.startswith('n%pl')

def is_monosyllabic_root(headword):
    return headword in monosyllabic_set

def add_exception_to_doc(exception_type, word, oxford_text, line_number, filename):
    para = doc.add_paragraph()
    para.add_run(f"{exception_type}: {word}").bold = True
    para.add_run(f"\nLine {line_number} ({filename})\n")

    for token in oxford_text.split():
        run = para.add_run(token + " ")
        if token.lower().strip(".,;:!?") == word.lower():
            run.italic = True
    para.add_run("\n\n")

# -----------------------------
# Tag Parsing
# -----------------------------
def parse_tagged_text(oxford_text, oxford_tagging):
    if pd.isna(oxford_text) or pd.isna(oxford_tagging):
        return [], [], []

    words = re.sub(r'[^\w\s]', '', oxford_text.lower()).split()
    tag_tokens = oxford_tagging.strip().split()

    tags, headwords = [], []
    for token in tag_tokens:
        if '@' in token and token not in ["--@dash", ".@ellipsis"]:
            head, tag = token.split('@', 1)
            tag = ''.join(i for i in tag if not i.isdigit())
            headwords.append(head.lower())
            tags.append(tag)

    min_len = min(len(words), len(tags), len(headwords))
    return words[:min_len], headwords[:min_len], tags[:min_len]

# -----------------------------
# Core Analysis
# -----------------------------
def analyze_adjectives(df, results, current_file, always_e_adjectives):

    if current_file not in results['file_stats']:
        results['file_stats'][current_file] = {
            'weak_total':0,'weak_correct':0,'weak_exceptions':0,
            'plural_total':0,'plural_correct':0,'plural_exceptions':0,
            'strong_total':0,'strong_correct':0,'strong_exceptions':0,
            'monosyllabic_adjectives_found':0
        }

    for _, row in df.iterrows():
        words, headwords, tags = parse_tagged_text(row['OXFORD_TEXT'], row['OXFORD_TAGGING'])
        prev_tag = 'NA'

        for j, tag in enumerate(tags):
            if not tag.startswith('adj'):
                prev_tag = tag
                continue

            headword = headwords[j]
            word = words[j]
            next_tag = tags[j+1] if j+1 < len(tags) else 'END'
            next_word = words[j+1] if j+1 < len(words) else 'END'

            if not is_monosyllabic_root(headword):
                prev_tag = tag
                continue

            results['file_stats'][current_file]['monosyllabic_adjectives_found'] += 1

            is_weak = is_weak_form(prev_tag, '', tag, next_tag)
            is_plural = is_plural_form(prev_tag, tag, next_tag)
            is_word_elided = next_word != 'END' and is_elided(word, next_word, next_tag)
            is_final = next_word == 'END'

            record = {
                'headword': headword,
                'word': word,
                'line_number': row['LINE_NUMBER'],
                'filename': row['OXFORD_FILENAME'],
                'context': row['OXFORD_TEXT'],
                'is_elided': is_word_elided,
                'is_final': is_final,
                'source_csv': current_file
            }

            # WEAK
            if is_weak:
                results['file_stats'][current_file]['weak_total'] += 1
                if word.endswith('e'):
                    results['file_stats'][current_file]['weak_correct'] += 1
                else:
                    results['file_stats'][current_file]['weak_exceptions'] += 1
                    results['weak_no_e_all'].append(record)
                    if not is_word_elided and not is_final:
                        results['weak_no_e_strict'].append(record)
                        add_exception_to_doc("Weak without -e", word, row['OXFORD_TEXT'], row['LINE_NUMBER'], row['OXFORD_FILENAME'])

            # PLURAL
            elif is_plural:
                results['file_stats'][current_file]['plural_total'] += 1
                if word.endswith('e'):
                    results['file_stats'][current_file]['plural_correct'] += 1
                else:
                    results['file_stats'][current_file]['plural_exceptions'] += 1
                    results['plural_no_e_all'].append(record)
                    if not is_word_elided and not is_final:
                        results['plural_no_e_strict'].append(record)
                        add_exception_to_doc("Plural without -e", word, row['OXFORD_TEXT'], row['LINE_NUMBER'], row['OXFORD_FILENAME'])

            # STRONG
            else:
                results['file_stats'][current_file]['strong_total'] += 1
                if word.endswith('e'):
                    results['file_stats'][current_file]['strong_exceptions'] += 1
                    results['strong_with_e_all'].append(record)
                    if not is_word_elided and not is_final:
                        results['strong_with_e_strict'].append(record)

                        # 🔥 FILTERED DOC LOGGING
                        if headword not in always_e_adjectives:
                            add_exception_to_doc("Strong with -e (true exception)", word, row['OXFORD_TEXT'], row['LINE_NUMBER'], row['OXFORD_FILENAME'])
                else:
                    results['file_stats'][current_file]['strong_correct'] += 1

            prev_tag = tag

    return results

# -----------------------------
# Driver
# -----------------------------
def process_csv_directory(csv_dir):
    results = {
        'all_adjectives': {},
        'weak_no_e_all': [], 'weak_no_e_strict': [],
        'plural_no_e_all': [], 'plural_no_e_strict': [],
        'strong_with_e_all': [], 'strong_with_e_strict': [],
        'file_stats': {}
    }

    for root, _, files in os.walk(csv_dir):
        for file in files:
            if file.endswith('_gui_complete.csv'):
                df = pd.read_csv(os.path.join(root, file))
                results = analyze_adjectives(df, results, file, always_e_adjectives_global)

    return results

# -----------------------------
# MAIN
# -----------------------------
if __name__ == "__main__":
    results = process_csv_directory(base_csv_dir)

    os.makedirs("adjective_analysis_output", exist_ok=True)
    doc.save("adjective_analysis_output/declension_exceptions.docx")

    print("\nAnalysis complete with filtered strong-form exceptions in the Word document.")
