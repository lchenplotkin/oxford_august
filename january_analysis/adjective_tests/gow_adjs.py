# Gower Adjective Declension Analysis with Per-File Accuracy Breakdown

import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict

doc = docx.Document()
doc.add_heading('Adjective Declension Exceptions in Gower', 0)

base_csv_dir = '../gow_csvs'

ELISION_FOLLOWERS = ["have","haven","haveth","havest","had","hadde",
                     "hadden","his","her","him","hers","hide","hir",
                     "hire","hires","hirs","han"]

# ---------------- MONOSYLLABIC LIST ---------------- #
monosyllabic_set = set()
with open('monosyllabic_adjectives.csv','r',encoding='utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        key = 'headword' if 'headword' in row else 'Headword'
        monosyllabic_set.add(row[key].lower())

# ---------------- UTILITIES ---------------- #

def is_elided(word, next_word):
    return next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS

def is_weak_form(prev_tag, tag, next_tag):
    weak_triggers = ['demonstrative','def_art','n%gen','pron%gen','interj','pron%fem_gen']
    return prev_tag in weak_triggers or next_tag == 'n#propn'

def is_plural_form(prev_tag, next_tag):
    return next_tag.startswith('n%pl') or prev_tag.startswith('n%pl')

def parse_tagged_text(text, tagging):
    if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
        return [], [], []

    words = re.sub(r'[.,!?°¶]', '', text.lower()).split()
    tags, headwords = [], []

    for tok in tagging.split():
        if '@' in tok and tok not in ["--@dash",".@ellipsis"]:
            h, t = tok.split('@')
            tags.append(re.sub(r'\d+', '', t))
            headwords.append(h.lower())

    min_len = min(len(words), len(tags))
    return words[:min_len], headwords[:min_len], tags[:min_len]

def is_monosyllabic_root(headword):
    return headword in monosyllabic_set

def add_exception_to_doc(reason, word, text, line, filename):
    p = doc.add_paragraph()
    p.add_run(f"{reason}: {word}").bold = True
    p.add_run(f"\nLine {line} ({filename})\n")
    for tok in text.split():
        run = p.add_run(tok + " ")
        if tok.lower() == word.lower():
            run.italic = True
    p.add_run("\n")

# ---------------- ANALYSIS ---------------- #

def analyze_adjectives(df, results, current_file):
    if current_file not in results['file_stats']:
        results['file_stats'][current_file] = dict(
            weak_total=0, weak_correct=0, weak_exceptions=0,
            plural_total=0, plural_correct=0, plural_exceptions=0,
            strong_total=0, strong_correct=0, strong_exceptions=0,
            monosyllabic_adjectives_found=0
        )

    for _, row in df.iterrows():
        words, headwords, tags = parse_tagged_text(row['TEXT'], row['TAGGING'])

        prev_tag, prev_word = 'NA', 'NA'

        for i, tag in enumerate(tags):
            if not tag.startswith('adj'):
                prev_tag = tag
                prev_word = words[i]
                continue

            word = words[i]
            headword = headwords[i]
            next_tag = tags[i+1] if i+1 < len(tags) else 'END'
            next_word = words[i+1] if i+1 < len(words) else 'END'

            if not is_monosyllabic_root(headword):
                prev_tag = tag
                prev_word = word
                continue

            if prev_word in ['this','that','thilke']:
                prev_tag = 'demonstrative'

            # Add to all_adjectives dictionary
            if headword not in results['all_adjectives']:
                results['all_adjectives'][headword] = {
                    'stem_forms': set(),  # Stem forms (without -e)
                    'stem_with_e': set(),  # Stem forms with -e
                    'all_forms': set(),
                    'stem_without_e': set()  # For tracking stem forms without -e
                }
            
            results['all_adjectives'][headword]['all_forms'].add(word)
            
            # Classify as stem form with or without -e
            # (You may need to adjust this logic based on your actual classification)
            if word.endswith('e'):
                results['all_adjectives'][headword]['stem_with_e'].add(word)
            else:
                results['all_adjectives'][headword]['stem_without_e'].add(word)
            
            # Determine if it's a stem form (you need to define this logic)
            # For now, I'm assuming all forms are stem forms
            results['all_adjectives'][headword]['stem_forms'].add(word)

            stats = results['file_stats'][current_file]
            stats['monosyllabic_adjectives_found'] += 1

            is_weak = is_weak_form(prev_tag, tag, next_tag)
            is_plural = is_plural_form(prev_tag, next_tag)
            is_final = next_word == 'END'
            elided = is_elided(word, next_word) if next_word != 'END' else False

            record = dict(headword=headword, word=word,
                          line_number=row['LINE_NUMBER'],
                          filename=row['FILENAME'],
                          context=row['OG_TEXT'],
                          is_elided=elided, is_final=is_final,
                          source_csv=current_file)

            if is_weak:
                stats['weak_total'] += 1
                if word.endswith('e'):
                    stats['weak_correct'] += 1
                else:
                    stats['weak_exceptions'] += 1
                    results['weak_no_e_all'].append(record)
                    if not elided and not is_final:
                        results['weak_no_e_strict'].append(record)
                        add_exception_to_doc("Weak without -e", word, row['OG_TEXT'], row['LINE_NUMBER'], row['FILENAME'])

            elif is_plural:
                stats['plural_total'] += 1
                if word.endswith('e'):
                    stats['plural_correct'] += 1
                else:
                    stats['plural_exceptions'] += 1
                    results['plural_no_e_all'].append(record)
                    if not elided and not is_final:
                        results['plural_no_e_strict'].append(record)
                        add_exception_to_doc("Plural without -e", word, row['OG_TEXT'], row['LINE_NUMBER'], row['FILENAME'])

            else:
                stats['strong_total'] += 1
                if word.endswith('e'):
                    stats['strong_exceptions'] += 1
                    results['strong_with_e_all'].append(record)
                    if not elided and not is_final:
                        results['strong_with_e_strict'].append(record)
                else:
                    stats['strong_correct'] += 1

            prev_tag = tag
            prev_word = word

    return results

# ---------------- DIRECTORY PROCESSING ---------------- #

def process_csv_directory(csv_dir):
    results = dict(
        weak_no_e_all=[], weak_no_e_strict=[],
        plural_no_e_all=[], plural_no_e_strict=[],
        strong_with_e_all=[], strong_with_e_strict=[],
        all_adjectives={},  # Added this line
        file_stats={}
    )

    for root, _, files in os.walk(csv_dir):
        for file in files:
            if not file.endswith('.csv'):
                continue
            df = pd.read_csv(os.path.join(root, file), encoding='utf-8')
            if not {'TEXT','TAGGING','LINE_NUMBER','FILENAME','OG_TEXT'}.issubset(df.columns):
                continue
            results = analyze_adjectives(df, results, file)

    return results
# ---------------- OUTPUT (UNCHANGED LOGIC) ---------------- #

def filter_strong_form_exceptions(results):
	"""Filter out adjectives that always end in -e in strong form from strong form exceptions"""
	# Identify adjectives that always end in -e in strong form
	always_e_adjectives = set()
	for headword, data in results['all_adjectives'].items():
		if is_monosyllabic_root(headword):
			# If this adjective has stem forms and ALL stem forms end in -e, it's an "always -e" adjective
			if data['stem_forms'] and len(data['stem_without_e']) == 0:
				always_e_adjectives.add(headword)
	
	# Filter the strong form exceptions to exclude always-e adjectives
	results['strong_with_e_all_filtered'] = [
		record for record in results['strong_with_e_all'] 
		if record['headword'] not in always_e_adjectives
	]
	results['strong_with_e_strict_filtered'] = [
		record for record in results['strong_with_e_strict'] 
		if record['headword'] not in always_e_adjectives
	]
	
	# Update file statistics to account for filtered strong form exceptions
	# Recalculate strong_exceptions for each file based on filtered data
	for file in results['file_stats']:
		# Count how many strong exceptions from this file were filtered out
		original_exceptions = [r for r in results['strong_with_e_all'] if r['source_csv'] == file]
		filtered_exceptions = [r for r in results['strong_with_e_all_filtered'] if r['source_csv'] == file]
		
		filtered_out_count = len(original_exceptions) - len(filtered_exceptions)
		
		# Adjust the strong stats
		results['file_stats'][file]['strong_exceptions_filtered'] = len(filtered_exceptions)
		results['file_stats'][file]['strong_correct_filtered'] = (
			results['file_stats'][file]['strong_correct'] + filtered_out_count
		)
		results['file_stats'][file]['strong_total_filtered'] = (
			results['file_stats'][file]['strong_correct_filtered'] + 
			results['file_stats'][file]['strong_exceptions_filtered']
		)
	
	# Add the always-e adjectives to results for reporting
	results['always_e_adjectives'] = always_e_adjectives
	
	return results

def calculate_accuracy_rates(results):
	"""Calculate accuracy rates for each file"""
	accuracy_data = []
	
	for file, stats in results['file_stats'].items():
		file_data = {'filename': file}
		
		# Weak declension accuracy (should end in -e)
		if stats['weak_total'] > 0:
			weak_accuracy = (stats['weak_correct'] / stats['weak_total']) * 100
		else:
			weak_accuracy = None
		file_data['weak_accuracy'] = weak_accuracy
		file_data['weak_correct'] = stats['weak_correct']
		file_data['weak_total'] = stats['weak_total']
		
		# Plural form accuracy (should end in -e)
		if stats['plural_total'] > 0:
			plural_accuracy = (stats['plural_correct'] / stats['plural_total']) * 100
		else:
			plural_accuracy = None
		file_data['plural_accuracy'] = plural_accuracy
		file_data['plural_correct'] = stats['plural_correct']
		file_data['plural_total'] = stats['plural_total']
		
		# Strong form accuracy (should NOT end in -e)
		if stats['strong_total'] > 0:
			strong_accuracy = (stats['strong_correct'] / stats['strong_total']) * 100
		else:
			strong_accuracy = None
		file_data['strong_accuracy'] = strong_accuracy
		file_data['strong_correct'] = stats['strong_correct']
		file_data['strong_total'] = stats['strong_total']
		
		# Strong form accuracy (filtered - excluding always-e adjectives)
		if 'strong_total_filtered' in stats and stats['strong_total_filtered'] > 0:
			strong_accuracy_filtered = (stats['strong_correct_filtered'] / stats['strong_total_filtered']) * 100
		else:
			strong_accuracy_filtered = None
		file_data['strong_accuracy_filtered'] = strong_accuracy_filtered
		file_data['strong_correct_filtered'] = stats.get('strong_correct_filtered', 0)
		file_data['strong_total_filtered'] = stats.get('strong_total_filtered', 0)
		
		# Overall accuracy
		total_instances = stats['weak_total'] + stats['plural_total'] + stats['strong_total']
		total_correct = stats['weak_correct'] + stats['plural_correct'] + stats['strong_correct']
		
		if total_instances > 0:
			overall_accuracy = (total_correct / total_instances) * 100
		else:
			overall_accuracy = None
		file_data['overall_accuracy'] = overall_accuracy
		file_data['total_correct'] = total_correct
		file_data['total_instances'] = total_instances
		
		# Overall accuracy (filtered)
		total_instances_filtered = (stats['weak_total'] + stats['plural_total'] + 
								   stats.get('strong_total_filtered', 0))
		total_correct_filtered = (stats['weak_correct'] + stats['plural_correct'] + 
								 stats.get('strong_correct_filtered', 0))
		
		if total_instances_filtered > 0:
			overall_accuracy_filtered = (total_correct_filtered / total_instances_filtered) * 100
		else:
			overall_accuracy_filtered = None
		file_data['overall_accuracy_filtered'] = overall_accuracy_filtered
		file_data['total_correct_filtered'] = total_correct_filtered
		file_data['total_instances_filtered'] = total_instances_filtered
		
		file_data['monosyllabic_adjectives_found'] = stats['monosyllabic_adjectives_found']
		
		accuracy_data.append(file_data)
	
	return accuracy_data

def write_results_to_csv(results, output_dir):
	"""Write analysis results to CSV files"""
	os.makedirs(output_dir, exist_ok=True)
	
	# Filter strong form exceptions first
	results = filter_strong_form_exceptions(results)
	
	# Helper function to write a list of records to CSV
	def write_records(records, filename):
		if not records:
			# Create empty file with headers
			with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
				fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final', 'source_csv']
				writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
				writer.writeheader()
			return
			
		with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
			fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final', 'source_csv']
			writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
			writer.writeheader()
			for record in records:
				writer.writerow(record)

	# Write each category
	write_records(results['weak_no_e_all'], 'weak_no_e_all_instances.csv')
	write_records(results['weak_no_e_strict'], 'weak_no_e_strict_instances.csv')
	write_records(results['plural_no_e_all'], 'plural_no_e_all_instances.csv')
	write_records(results['plural_no_e_strict'], 'plural_no_e_strict_instances.csv')
	write_records(results['strong_with_e_all'], 'strong_with_e_all_instances.csv')
	write_records(results['strong_with_e_strict'], 'strong_with_e_strict_instances.csv')
	write_records(results['strong_with_e_all_filtered'], 'strong_with_e_all_filtered_instances.csv')
	write_records(results['strong_with_e_strict_filtered'], 'strong_with_e_strict_filtered_instances.csv')

	# Write per-file accuracy breakdown
	accuracy_data = calculate_accuracy_rates(results)
	with open(os.path.join(output_dir, 'per_file_accuracy_breakdown.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = [
			'filename',
			'overall_accuracy', 'total_correct', 'total_instances',
			'overall_accuracy_filtered', 'total_correct_filtered', 'total_instances_filtered',
			'weak_accuracy', 'weak_correct', 'weak_total',
			'plural_accuracy', 'plural_correct', 'plural_total',
			'strong_accuracy', 'strong_correct', 'strong_total',
			'strong_accuracy_filtered', 'strong_correct_filtered', 'strong_total_filtered',
			'monosyllabic_adjectives_found'
		]
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		
		# Sort by filename
		accuracy_data.sort(key=lambda x: x['filename'])
		
		for row in accuracy_data:
			# Format percentages to 2 decimal places
			formatted_row = {}
			for key, value in row.items():
				if key.endswith('_accuracy') and value is not None:
					formatted_row[key] = f"{value:.2f}"
				else:
					formatted_row[key] = value
			writer.writerow(formatted_row)
	
	# Calculate and write overall accuracy statistics
	overall_stats = {
		'weak_total': sum(s['weak_total'] for s in results['file_stats'].values()),
		'weak_correct': sum(s['weak_correct'] for s in results['file_stats'].values()),
		'plural_total': sum(s['plural_total'] for s in results['file_stats'].values()),
		'plural_correct': sum(s['plural_correct'] for s in results['file_stats'].values()),
		'strong_total': sum(s['strong_total'] for s in results['file_stats'].values()),
		'strong_correct': sum(s['strong_correct'] for s in results['file_stats'].values()),
		'strong_total_filtered': sum(s.get('strong_total_filtered', 0) for s in results['file_stats'].values()),
		'strong_correct_filtered': sum(s.get('strong_correct_filtered', 0) for s in results['file_stats'].values()),
	}

	# Write summary statistics
	with open(os.path.join(output_dir, 'summary_statistics.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Category', 'All Instances', 'Strict Instances (not elided, not final)', 'Accuracy Rate'])
		
		weak_acc = (overall_stats['weak_correct'] / overall_stats['weak_total'] * 100) if overall_stats['weak_total'] > 0 else 0
		writer.writerow(['Weak declension (should have -e)', 
						overall_stats['weak_total'], 
						len(results['weak_no_e_strict']),
						f"{weak_acc:.2f}%"])
		
		plural_acc = (overall_stats['plural_correct'] / overall_stats['plural_total'] * 100) if overall_stats['plural_total'] > 0 else 0
		writer.writerow(['Plural form (should have -e)', 
						overall_stats['plural_total'], 
						len(results['plural_no_e_strict']),
						f"{plural_acc:.2f}%"])
		
		strong_acc = (overall_stats['strong_correct'] / overall_stats['strong_total'] * 100) if overall_stats['strong_total'] > 0 else 0
		writer.writerow(['Strong form (should not have -e, all)', 
						overall_stats['strong_total'], 
						len(results['strong_with_e_strict']),
						f"{strong_acc:.2f}%"])
		
		strong_acc_filtered = (overall_stats['strong_correct_filtered'] / overall_stats['strong_total_filtered'] * 100) if overall_stats['strong_total_filtered'] > 0 else 0
		writer.writerow(['Strong form (should not have -e, filtered)', 
						overall_stats['strong_total_filtered'], 
						len(results['strong_with_e_strict_filtered']),
						f"{strong_acc_filtered:.2f}%"])
		
		writer.writerow(['Always -e adjectives found', len(results['always_e_adjectives']), 'N/A', 'N/A'])

	# Write monosyllabic adjectives found
	with open(os.path.join(output_dir, 'monosyllabic_adjectives.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Headword', 'Stem Forms', 'All Forms', 'Always Ends in E'])
		for headword, data in results['all_adjectives'].items():
			if is_monosyllabic_root(headword):
				always_e = 'Yes' if headword in results['always_e_adjectives'] else 'No'
				writer.writerow([headword, '; '.join(sorted(data['stem_forms'])), '; '.join(sorted(data['all_forms'])), always_e])

	# Write list of always -e adjectives
	with open(os.path.join(output_dir, 'always_e_adjectives.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Headword', 'Stem Forms with -e', 'Reason'])
		for headword in sorted(results['always_e_adjectives']):
			data = results['all_adjectives'][headword]
			stem_forms = '; '.join(sorted(data['stem_with_e']))
			writer.writerow([headword, stem_forms, 'All stem forms end in -e'])

def write_summary_to_docx(results, output_dir):
	"""Write summary statistics to Word document"""
	# Filter strong form exceptions first
	results = filter_strong_form_exceptions(results)
	
	summary_doc = docx.Document()
	summary_doc.add_heading('Adjective Declension Analysis Summary', 0)
	
	monosyllabic_count = sum(1 for headword, data in results['all_adjectives'].items() 
						   if is_monosyllabic_root(headword))
	
	# Overview section
	summary_doc.add_heading('Overview', 1)
	para = summary_doc.add_paragraph()
	para.add_run(f"Total adjectives found: {len(results['all_adjectives'])}\n")
	para.add_run(f"Monosyllabic adjectives found: {monosyllabic_count}\n")
	para.add_run(f"Always -e adjectives found: {len(results['always_e_adjectives'])}\n")
	para.add_run(f"Files processed: {len(results['file_stats'])}")
	
	# Calculate overall accuracy
	overall_stats = {
		'weak_total': sum(s['weak_total'] for s in results['file_stats'].values()),
		'weak_correct': sum(s['weak_correct'] for s in results['file_stats'].values()),
		'plural_total': sum(s['plural_total'] for s in results['file_stats'].values()),
		'plural_correct': sum(s['plural_correct'] for s in results['file_stats'].values()),
		'strong_total': sum(s['strong_total'] for s in results['file_stats'].values()),
		'strong_correct': sum(s['strong_correct'] for s in results['file_stats'].values()),
		'strong_total_filtered': sum(s.get('strong_total_filtered', 0) for s in results['file_stats'].values()),
		'strong_correct_filtered': sum(s.get('strong_correct_filtered', 0) for s in results['file_stats'].values()),
	}
	
	# Accuracy rates section
	summary_doc.add_heading('Overall Accuracy Rates', 1)
	para = summary_doc.add_paragraph()
	
	if overall_stats['weak_total'] > 0:
		weak_acc = (overall_stats['weak_correct'] / overall_stats['weak_total']) * 100
		para.add_run(f"Weak declension accuracy: {weak_acc:.2f}% ({overall_stats['weak_correct']}/{overall_stats['weak_total']})\n")
	
	if overall_stats['plural_total'] > 0:
		plural_acc = (overall_stats['plural_correct'] / overall_stats['plural_total']) * 100
		para.add_run(f"Plural form accuracy: {plural_acc:.2f}% ({overall_stats['plural_correct']}/{overall_stats['plural_total']})\n")
	
	if overall_stats['strong_total'] > 0:
		strong_acc = (overall_stats['strong_correct'] / overall_stats['strong_total']) * 100
		para.add_run(f"Strong form accuracy (all): {strong_acc:.2f}% ({overall_stats['strong_correct']}/{overall_stats['strong_total']})\n")
	
	if overall_stats['strong_total_filtered'] > 0:
		strong_acc_filtered = (overall_stats['strong_correct_filtered'] / overall_stats['strong_total_filtered']) * 100
		para.add_run(f"Strong form accuracy (filtered): {strong_acc_filtered:.2f}% ({overall_stats['strong_correct_filtered']}/{overall_stats['strong_total_filtered']})")
	
	# Statistics section
	summary_doc.add_heading('Exception Statistics', 1)
	
	categories = [
		('Weak declension without -e', 'weak_no_e_all', 'weak_no_e_strict'),
		('Plural form without -e', 'plural_no_e_all', 'plural_no_e_strict'),
		('Strong form with -e (all)', 'strong_with_e_all', 'strong_with_e_strict'),
		('Strong form with -e (filtered)', 'strong_with_e_all_filtered', 'strong_with_e_strict_filtered')
	]
	
	for desc, all_key, strict_key in categories:
		all_count = len(results[all_key])
		strict_count = len(results[strict_key])
		
		para = summary_doc.add_paragraph()
		para.add_run(desc + ":").bold = True
		para.add_run(f"\n  All instances: {all_count}")
		para.add_run(f"\n  Strict instances (not elided, not final): {strict_count}\n")
	
	# Always -e adjectives section
	if results['always_e_adjectives']:
		summary_doc.add_heading('Adjectives That Always End in -e in Strong Form', 1)
		para = summary_doc.add_paragraph()
		para.add_run("The following adjectives were excluded from strong form exception analysis because they always end in -e in the strong form:\n\n")
		
		for headword in sorted(results['always_e_adjectives']):
			data = results['all_adjectives'][headword]
			stem_forms = ', '.join(sorted(data['stem_with_e']))
			para.add_run(f"• {headword}").bold = True
			para.add_run(f" (forms: {stem_forms})\n")
	
	summary_doc.save(os.path.join(output_dir, 'analysis_summary.docx'))

# Reuse your SAME write_results_to_csv(), calculate_accuracy_rates(),
# filter_strong_form_exceptions(), and write_summary_to_docx()
# They work without modification because the record structure is identical.

# ---------------- MAIN ---------------- #

if __name__ == "__main__":
    results = process_csv_directory(base_csv_dir)

    output_dir = 'gower_adjective_analysis_output'
    write_results_to_csv(results, output_dir)
    write_summary_to_docx(results, output_dir)

    doc.save(os.path.join(output_dir, 'declension_exceptions.docx'))

    print(f"\nAnalysis complete! Results saved to '{output_dir}' directory.")

