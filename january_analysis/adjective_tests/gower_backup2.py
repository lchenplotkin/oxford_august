import docx 
import os
import re
import string
import csv
import pandas as pd
from collections import defaultdict, Counter
import matplotlib.pyplot as plt

# Initialize Word document for exceptions
doc = docx.Document()
doc.add_heading('Adjective Declension Exceptions in Gower', 0)

# Configuration
base_csv_dir = '../gow_csvs'

orderfile = open('gower_adjective_analysis_output/adj_order.csv', 'w', newline='', encoding='utf-8')
fieldnames = ['adj1', 'adj2', 'text','line_number', 'filename' ]
orderwriter = csv.writer(orderfile)
orderwriter.writerow(fieldnames)

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
					"hadden", "his", "her", "him", "hers", "hide", "hir",
					"hire", "hires", "hirs", "han"]

monosyllabic_set = set()
mono_csv = 'monosyllabic_adjectives.csv'
with open(mono_csv, 'r', encoding='utf-8') as csvfile:
	reader = csv.DictReader(csvfile)
	for row in reader:
		# Assume the CSV has a column named 'headword' or similar
		if 'headword' in row:
			monosyllabic_set.add(row['headword'].lower())
		elif 'Headword' in row:
			monosyllabic_set.add(row['Headword'].lower())

# Utilities
def vowel_cluster_count(w):
	"""Count vowel clusters to determine syllables"""
	return len(re.findall(r'[^aeiouy]*[aeiouy]+(?:[^aeiouy]+(?=[^aeiouy]*[aeiouy])|[^aeiouy]*)?', w)) or 1

def is_elided(word, next_word, next_tag):
	"""Check if word is elided before next word"""
	if next_word[0] in 'aeiou' or (next_word[0] == 'h' and next_tag.startswith('pron')) or next_word in ELISION_FOLLOWERS:
		return True
	return False

def is_weak_form(prev_tag, prev_word, tag, next_tag):
	"""Check if adjective is in weak declension context"""
	weak_triggers = ['demonstrative', 'def_art', 'n%gen', 'pron%gen', 'interj', 'pron%pl_gen','pron%fem_gen']
	for cause in weak_triggers:
		if prev_tag == cause:
			return True
	if next_tag == 'n#propn':
		return True
	return False

def is_plural_form(prev_tag, tag, next_tag):
	"""Check if adjective modifies plural noun"""
	if next_tag.startswith('n%pl') or prev_tag.startswith('n%pl'):
		return True
	return False

def parse_tagged_text(text, tagging):
	"""Extract words from text and tags from tagging"""
	if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
		return [], [], []

	# Clean and split the text
	words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()

	# Parse tags from tagging
	tag_tokens = tagging.strip().split()
	tags = []
	headwords = []

	for token in tag_tokens:
		if '@' in token:
			if token not in ["--@dash", ".@ellipsis"]:
				parts = token.split('@')
				headword = parts[0].lower()
				tag_part = parts[1] if len(parts) > 1 else ''

				# Remove digits from tag
				tag = ''.join([i for i in tag_part if not i.isdigit()])
				if token == 'your@pron' or token == 'min@pron':
					tag = 'pron%gen'
				headwords.append(headword)
				tags.append(tag)

	# Handle special demonstrative cases
	for i, (word, tag) in enumerate(zip(words[:len(tags)], tags)):
		if word in ['this', 'that', 'thilke'] and 'gram_adj' in tag:
			tags[i] = 'demonstrative'

	# Ensure all lists are the same length (trim to shortest)
	if len(words) != len(tags):
		print(f"Length mismatch: {text} | {tagging}")
		print(f"Words: {len(words)}, Tags: {len(tags)}")
	min_len = min(len(words), len(tags), len(headwords))
	return words[:min_len], headwords[:min_len], tags[:min_len]

def is_monosyllabic_root(headword, word_forms):
	if headword in monosyllabic_set:
		return True
	else:
		return False

def add_exception_to_doc(exception_type, word, text, line_number, filename):
	"""Add an exception case to the Word document"""
	doc_para = doc.add_paragraph()
	doc_para.add_run(f"{exception_type}: {word}").bold = True
	doc_para.add_run(f"\nLine {line_number} ({filename})\n")
	
	for i, word_token in enumerate(text.split(' ')):
		if word_token.lower() == word.lower():
			doc_para.add_run(f"{word_token}").italic = True
		else:
			doc_para.add_run(f"{word_token}")
		if i < len(text.split(' ')) - 1:
			doc_para.add_run(" ")
	doc_para.add_run("\n\n")

def analyze_adjectives(df, results, current_file):
	"""Analyze adjective patterns in CSV data"""
	
	# Initialize file-specific tracking if needed
	if current_file not in results['file_stats']:
		results['file_stats'][current_file] = {
			'weak_total': 0,
			'weak_correct': 0,
			'weak_exceptions': 0,
			'plural_total': 0,
			'plural_correct': 0,
			'plural_exceptions': 0,
			'strong_total': 0,
			'strong_correct': 0,
			'strong_exceptions': 0,
			'monosyllabic_adjectives_found': 0
		}
	
	# Initialize headword_stats if needed
	if 'headword_stats' not in results:
		results['headword_stats'] = {}
	
	# Initialize headword_stats structure if not present
	if 'headword_stats' not in results:
		results['headword_stats'] = {}
	
	for idx, row in df.iterrows():
		text = row['TEXT']
		tagging = row['TAGGING']
		line_number = row['LINE_NUMBER']
		filename = row['FILENAME']#.split('gow')[0]
		#= row['FILENAME'].split('gow')[0]
		original_text = row.get('OG_TEXT', text)

		words, headwords, tags = parse_tagged_text(text, tagging)

		prev_tag = 'NA'
		prev_word = 'NA'
		
		for j in range(len(tags)):
			if j >= len(words) or j >= len(headwords):
				continue
				
			headword = headwords[j]
			tag = tags[j]
			word = words[j].lower()
			next_tag = tags[j + 1] if j + 1 < len(tags) else 'END'
			next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'

			# Only process adjectives
			if not tag.startswith('adj'):
				prev_word = word
				prev_tag = tag
				continue

			if tag == 'adj' and next_tag == 'adj':
				#print(original_text)
				for later_headword, later_tag in zip(headwords[j+1:],tags[j+1:]):
					if later_tag != 'adj':
						break
					else:
						orderwriter.writerow([headword, later_headword,original_text, line_number,filename])


			# Initialize adjective tracking
			if headword not in results['all_adjectives']:
				results['all_adjectives'][headword] = {
					'stem_forms': set(), 
					'all_forms': set(),
					'stem_with_e': set(),
					'stem_without_e': set()
				}

			results['all_adjectives'][headword]['all_forms'].add(word)
			if tag == 'adj':  # Stem form
				results['all_adjectives'][headword]['stem_forms'].add(word)
				if word.endswith('e'):
					results['all_adjectives'][headword]['stem_with_e'].add(word)
				else:
					results['all_adjectives'][headword]['stem_without_e'].add(word)

			# Check if monosyllabic root
			stem_forms = results['all_adjectives'][headword]['stem_forms']
			if not is_monosyllabic_root(headword, stem_forms):
				prev_word = word
				prev_tag = tag
				continue

			# Track that we found a monosyllabic adjective in this file
			results['file_stats'][current_file]['monosyllabic_adjectives_found'] += 1

			# Initialize headword-specific statistics if needed
			if headword not in results['headword_stats']:
				results['headword_stats'][headword] = {
					'weak_total': 0,
					'weak_correct': 0,
					'weak_exceptions': 0,
					'plural_total': 0,
					'plural_correct': 0,
					'plural_exceptions': 0,
					'strong_total': 0,
					'strong_correct': 0,
					'strong_exceptions': 0
				}

			# Determine grammatical context
			is_weak = is_weak_form(prev_tag, prev_word, tag, next_tag)
			is_plural = is_plural_form(prev_tag, tag, next_tag)
			is_final = (next_word == 'END')
			is_word_elided = is_elided(word, next_word, next_tag) if next_word != 'END' else False
			after_noun = False
			if not is_word_elided and is_weak and headword in['right','lift','up','down','west','east','north','south']:
				print(original_text)

			# Create record for this instance
			record = {
				'headword': headword,
				'word': word,
				'line_number': line_number,
				'filename': filename,
				'context': original_text,
				'is_elided': is_word_elided,
				'is_final': is_final,
				'source_csv': current_file
			}

			# Categorize the adjective usage and update file statistics
			if is_weak:
				results['file_stats'][current_file]['weak_total'] += 1
				results['headword_stats'][headword]['weak_total'] += 1
				
				if word.endswith('e'):
					# Correct weak declension with -e
					results['file_stats'][current_file]['weak_correct'] += 1
					results['headword_stats'][headword]['weak_correct'] += 1
				elif not after_noun:
					# Weak declension without -e (exception)
					results['file_stats'][current_file]['weak_exceptions'] += 1
					results['headword_stats'][headword]['weak_exceptions'] += 1
					results['weak_no_e_all'].append(record)
					if not is_word_elided and not is_final:
						results['weak_no_e_strict'].append(record)
						add_exception_to_doc("Weak without -e", word, original_text, line_number, filename)
						
			elif is_plural:
				results['file_stats'][current_file]['plural_total'] += 1
				results['headword_stats'][headword]['plural_total'] += 1
				
				if word.endswith('e'):
					# Correct plural form with -e
					results['file_stats'][current_file]['plural_correct'] += 1
					results['headword_stats'][headword]['plural_correct'] += 1
				else:
					# Plural form without -e (exception)
					results['file_stats'][current_file]['plural_exceptions'] += 1
					results['headword_stats'][headword]['plural_exceptions'] += 1
					results['plural_no_e_all'].append(record)
					if not is_word_elided and not is_final:
						results['plural_no_e_strict'].append(record)
						add_exception_to_doc("Plural without -e", word, original_text, line_number, filename)
			else:
				# Strong/stem form
				results['file_stats'][current_file]['strong_total'] += 1
				results['headword_stats'][headword]['strong_total'] += 1
				
				if word.endswith('e') and tag == 'adj':
					# Strong form with -e (potential exception)
					results['file_stats'][current_file]['strong_exceptions'] += 1
					results['headword_stats'][headword]['strong_exceptions'] += 1
					results['strong_with_e_all'].append(record)
					if not is_word_elided and not is_final:
						results['strong_with_e_strict'].append(record)
				else:
					# Correct strong form without -e
					results['file_stats'][current_file]['strong_correct'] += 1
					results['headword_stats'][headword]['strong_correct'] += 1

			prev_word = word
			prev_tag = tag

	return results

def process_csv_directory(csv_dir):
	"""Process all CSV files in the directory"""
	results = {
		'all_adjectives': {},
		'weak_no_e_all': [],		  # Weak declension without -e (all instances)
		'weak_no_e_strict': [],	   # Weak declension without -e (strict: not elided, not final)
		'plural_no_e_all': [],		# Plural form without -e (all instances)  
		'plural_no_e_strict': [],	 # Plural form without -e (strict: not elided, not final)
		'strong_with_e_all': [],	  # Strong form with -e (all instances)
		'strong_with_e_strict': [],   # Strong form with -e (strict: not elided, not final)
		'file_stats': {},			  # Per-file statistics
		'headword_stats': {}		   # Per-headword statistics
	}

	file_count = 0
	for root, dirs, files in os.walk(csv_dir):
		for file in files:
			if not file.endswith('.csv'):
				continue

			csv_path = os.path.join(root, file)
			file_count += 1

			try:
				# Read CSV file
				df = pd.read_csv(csv_path, encoding='utf-8')

				# Skip if required columns are missing
				required_columns = ['TAGGING', 'TEXT', 'LINE_NUMBER', 'FILENAME']
				if not all(col in df.columns for col in required_columns):
					print(f"Skipping {file}: missing required columns")
					continue

				results = analyze_adjectives(df, results, file)

			except Exception as e:
				print(f"Error processing {file}: {e}")
				continue

	print(f"Processed {file_count} files")
	return results

def filter_strong_form_exceptions(results):
	"""Filter out adjectives that always end in -e in strong form from strong form exceptions"""
	# Identify adjectives that always end in -e in strong form
	always_e_adjectives = set()
	for headword, data in results['all_adjectives'].items():
		if is_monosyllabic_root(headword, data['stem_forms']):
			# If this adjective has stem forms and ALL stem forms end in -e, it's an "always -e" adjective
			if data['stem_forms'] and len(data['stem_without_e']) == 0:
				always_e_adjectives.add(headword)
	
	# Filter the strong form exceptions to exclude always-e adjectives
	results['strong_with_e_all_filtered'] = [
		record for record in results['strong_with_e_all'] 
		if record['headword'] not in always_e_adjectives
	]
	results['strong_with_e_strict_filtered'] = [
		record for record in results['strong_with_e_strict'] 
		if record['headword'] not in always_e_adjectives
	]
	
	# Update file statistics to account for filtered strong form exceptions
	# Recalculate strong_exceptions for each file based on filtered data
	for file in results['file_stats']:
		# Count how many strong exceptions from this file were filtered out
		original_exceptions = [r for r in results['strong_with_e_all'] if r['source_csv'] == file]
		filtered_exceptions = [r for r in results['strong_with_e_all_filtered'] if r['source_csv'] == file]
		
		filtered_out_count = len(original_exceptions) - len(filtered_exceptions)
		
		# Adjust the strong stats
		results['file_stats'][file]['strong_exceptions_filtered'] = len(filtered_exceptions)
		results['file_stats'][file]['strong_correct_filtered'] = (
			results['file_stats'][file]['strong_correct'] + filtered_out_count
		)
		results['file_stats'][file]['strong_total_filtered'] = (
			results['file_stats'][file]['strong_correct_filtered'] + 
			results['file_stats'][file]['strong_exceptions_filtered']
		)
	
	# Add the always-e adjectives to results for reporting
	results['always_e_adjectives'] = always_e_adjectives
	
	return results

def calculate_accuracy_rates(results):
	"""Calculate accuracy rates for each file"""
	accuracy_data = []
	
	for file, stats in results['file_stats'].items():
		file_data = {'filename': file}
		
		# Weak declension accuracy (should end in -e)
		if stats['weak_total'] > 0:
			weak_accuracy = (stats['weak_correct'] / stats['weak_total']) * 100
		else:
			weak_accuracy = None
		file_data['weak_accuracy'] = weak_accuracy
		file_data['weak_correct'] = stats['weak_correct']
		file_data['weak_total'] = stats['weak_total']
		
		# Plural form accuracy (should end in -e)
		if stats['plural_total'] > 0:
			plural_accuracy = (stats['plural_correct'] / stats['plural_total']) * 100
		else:
			plural_accuracy = None
		file_data['plural_accuracy'] = plural_accuracy
		file_data['plural_correct'] = stats['plural_correct']
		file_data['plural_total'] = stats['plural_total']
		
		# Strong form accuracy (should NOT end in -e)
		if stats['strong_total'] > 0:
			strong_accuracy = (stats['strong_correct'] / stats['strong_total']) * 100
		else:
			strong_accuracy = None
		file_data['strong_accuracy'] = strong_accuracy
		file_data['strong_correct'] = stats['strong_correct']
		file_data['strong_total'] = stats['strong_total']
		
		# Strong form accuracy (filtered - excluding always-e adjectives)
		if 'strong_total_filtered' in stats and stats['strong_total_filtered'] > 0:
			strong_accuracy_filtered = (stats['strong_correct_filtered'] / stats['strong_total_filtered']) * 100
		else:
			strong_accuracy_filtered = None
		file_data['strong_accuracy_filtered'] = strong_accuracy_filtered
		file_data['strong_correct_filtered'] = stats.get('strong_correct_filtered', 0)
		file_data['strong_total_filtered'] = stats.get('strong_total_filtered', 0)
		
		# Overall accuracy
		total_instances = stats['weak_total'] + stats['plural_total'] + stats['strong_total']
		total_correct = stats['weak_correct'] + stats['plural_correct'] + stats['strong_correct']
		
		if total_instances > 0:
			overall_accuracy = (total_correct / total_instances) * 100
		else:
			overall_accuracy = None
		file_data['overall_accuracy'] = overall_accuracy
		file_data['total_correct'] = total_correct
		file_data['total_instances'] = total_instances
		
		# Overall accuracy (filtered)
		total_instances_filtered = (stats['weak_total'] + stats['plural_total'] + 
								   stats.get('strong_total_filtered', 0))
		total_correct_filtered = (stats['weak_correct'] + stats['plural_correct'] + 
								 stats.get('strong_correct_filtered', 0))
		
		if total_instances_filtered > 0:
			overall_accuracy_filtered = (total_correct_filtered / total_instances_filtered) * 100
		else:
			overall_accuracy_filtered = None
		file_data['overall_accuracy_filtered'] = overall_accuracy_filtered
		file_data['total_correct_filtered'] = total_correct_filtered
		file_data['total_instances_filtered'] = total_instances_filtered
		
		file_data['monosyllabic_adjectives_found'] = stats['monosyllabic_adjectives_found']
		
		accuracy_data.append(file_data)
	
	return accuracy_data

def calculate_headword_accuracy_rates(results):
	"""Calculate accuracy rates for each headword"""
	headword_accuracy_data = []
	
	for headword, stats in results['headword_stats'].items():
		headword_data = {'headword': headword}
		
		# Weak declension accuracy (should end in -e)
		if stats['weak_total'] > 0:
			weak_accuracy = (stats['weak_correct'] / stats['weak_total']) * 100
		else:
			weak_accuracy = None
		headword_data['weak_accuracy'] = weak_accuracy
		headword_data['weak_correct'] = stats['weak_correct']
		headword_data['weak_total'] = stats['weak_total']
		
		# Plural form accuracy (should end in -e)
		if stats['plural_total'] > 0:
			plural_accuracy = (stats['plural_correct'] / stats['plural_total']) * 100
		else:
			plural_accuracy = None
		headword_data['plural_accuracy'] = plural_accuracy
		headword_data['plural_correct'] = stats['plural_correct']
		headword_data['plural_total'] = stats['plural_total']
		
		# Strong form accuracy (should NOT end in -e)
		if stats['strong_total'] > 0:
			strong_accuracy = (stats['strong_correct'] / stats['strong_total']) * 100
		else:
			strong_accuracy = None
		headword_data['strong_accuracy'] = strong_accuracy
		headword_data['strong_correct'] = stats['strong_correct']
		headword_data['strong_total'] = stats['strong_total']
		
		# Overall accuracy
		total_instances = stats['weak_total'] + stats['plural_total'] + stats['strong_total']
		total_correct = stats['weak_correct'] + stats['plural_correct'] + stats['strong_correct']
		
		if total_instances > 0:
			overall_accuracy = (total_correct / total_instances) * 100
		else:
			overall_accuracy = None
		headword_data['overall_accuracy'] = overall_accuracy
		headword_data['total_correct'] = total_correct
		headword_data['total_instances'] = total_instances
		
		# Get word forms
		if headword in results['all_adjectives']:
			headword_data['stem_forms'] = '; '.join(sorted(results['all_adjectives'][headword]['stem_forms']))
			headword_data['all_forms'] = '; '.join(sorted(results['all_adjectives'][headword]['all_forms']))
		else:
			headword_data['stem_forms'] = ''
			headword_data['all_forms'] = ''
		
		# Check if always ends in -e
		if 'always_e_adjectives' in results:
			headword_data['always_e'] = 'Yes' if headword in results['always_e_adjectives'] else 'No'
		else:
			headword_data['always_e'] = 'Unknown'
		
		headword_accuracy_data.append(headword_data)
	
	return headword_accuracy_data


def calculate_headword_accuracy_rates(results):
	"""Calculate accuracy rates for each headword"""
	headword_data = []
	
	for headword, stats in results['headword_stats'].items():
		# Only process monosyllabic adjectives
		if headword not in results['all_adjectives']:
			continue
		if not is_monosyllabic_root(headword, results['all_adjectives'][headword]['stem_forms']):
			continue
			
		hw_data = {'headword': headword}
		
		# Weak declension accuracy (should end in -e)
		if stats['weak_total'] > 0:
			weak_accuracy = (stats['weak_correct'] / stats['weak_total']) * 100
		else:
			weak_accuracy = None
		hw_data['weak_accuracy'] = weak_accuracy
		hw_data['weak_correct'] = stats['weak_correct']
		hw_data['weak_total'] = stats['weak_total']
		hw_data['weak_exceptions'] = stats['weak_exceptions']
		
		# Plural form accuracy (should end in -e)
		if stats['plural_total'] > 0:
			plural_accuracy = (stats['plural_correct'] / stats['plural_total']) * 100
		else:
			plural_accuracy = None
		hw_data['plural_accuracy'] = plural_accuracy
		hw_data['plural_correct'] = stats['plural_correct']
		hw_data['plural_total'] = stats['plural_total']
		hw_data['plural_exceptions'] = stats['plural_exceptions']
		
		# Strong form accuracy (should NOT end in -e)
		if stats['strong_total'] > 0:
			strong_accuracy = (stats['strong_correct'] / stats['strong_total']) * 100
		else:
			strong_accuracy = None
		hw_data['strong_accuracy'] = strong_accuracy
		hw_data['strong_correct'] = stats['strong_correct']
		hw_data['strong_total'] = stats['strong_total']
		hw_data['strong_exceptions'] = stats['strong_exceptions']
		
		# Overall accuracy
		total_instances = stats['weak_total'] + stats['plural_total'] + stats['strong_total']
		total_correct = stats['weak_correct'] + stats['plural_correct'] + stats['strong_correct']
		
		if total_instances > 0:
			overall_accuracy = (total_correct / total_instances) * 100
		else:
			overall_accuracy = None
		hw_data['overall_accuracy'] = overall_accuracy
		hw_data['total_correct'] = total_correct
		hw_data['total_instances'] = total_instances
		
		# Check if this is an always-e adjective
		hw_data['always_ends_in_e'] = 'Yes' if headword in results.get('always_e_adjectives', set()) else 'No'
		
		headword_data.append(hw_data)
	
	return headword_data

def calculate_headword_accuracy_rates(results):
	"""Calculate accuracy rates for each headword"""
	headword_accuracy_data = []
	
	# Only process monosyllabic adjectives
	for headword, data in results['all_adjectives'].items():
		if not is_monosyllabic_root(headword, data['stem_forms']):
			continue
		
		stats = results['headword_stats'][headword]
		headword_data = {'headword': headword}
		
		# Weak declension accuracy (should end in -e)
		if stats['weak_total'] > 0:
			weak_accuracy = (stats['weak_correct'] / stats['weak_total']) * 100
		else:
			weak_accuracy = None
		headword_data['weak_accuracy'] = weak_accuracy
		headword_data['weak_correct'] = stats['weak_correct']
		headword_data['weak_total'] = stats['weak_total']
		
		# Plural form accuracy (should end in -e)
		if stats['plural_total'] > 0:
			plural_accuracy = (stats['plural_correct'] / stats['plural_total']) * 100
		else:
			plural_accuracy = None
		headword_data['plural_accuracy'] = plural_accuracy
		headword_data['plural_correct'] = stats['plural_correct']
		headword_data['plural_total'] = stats['plural_total']
		
		# Strong form accuracy (should NOT end in -e)
		if stats['strong_total'] > 0:
			strong_accuracy = (stats['strong_correct'] / stats['strong_total']) * 100
		else:
			strong_accuracy = None
		headword_data['strong_accuracy'] = strong_accuracy
		headword_data['strong_correct'] = stats['strong_correct']
		headword_data['strong_total'] = stats['strong_total']
		
		# Overall accuracy
		total_instances = stats['weak_total'] + stats['plural_total'] + stats['strong_total']
		total_correct = stats['weak_correct'] + stats['plural_correct'] + stats['strong_correct']
		
		if total_instances > 0:
			overall_accuracy = (total_correct / total_instances) * 100
		else:
			overall_accuracy = None
		headword_data['overall_accuracy'] = overall_accuracy
		headword_data['total_correct'] = total_correct
		headword_data['total_instances'] = total_instances
		
		# Check if this is an "always -e" adjective
		headword_data['always_e'] = headword in results.get('always_e_adjectives', set())
		
		# Add all forms for reference
		headword_data['all_forms'] = '; '.join(sorted(data['all_forms']))
		headword_data['stem_forms'] = '; '.join(sorted(data['stem_forms']))
		
		headword_accuracy_data.append(headword_data)
	
	return headword_accuracy_data

def write_results_to_csv(results, output_dir):
	"""Write analysis results to CSV files"""
	os.makedirs(output_dir, exist_ok=True)
	
	# Filter strong form exceptions first
	results = filter_strong_form_exceptions(results)
	
	# Helper function to write a list of records to CSV
	def write_records(records, filename):
		if not records:
			# Create empty file with headers
			with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
				fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final', 'source_csv']
				writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
				writer.writeheader()
			return
			
		with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
			fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final', 'source_csv']
			writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
			writer.writeheader()
			for record in records:
				writer.writerow(record)

	# Write each category
	write_records(results['weak_no_e_all'], 'weak_no_e_all_instances.csv')
	write_records(results['weak_no_e_strict'], 'weak_no_e_strict_instances.csv')
	write_records(results['plural_no_e_all'], 'plural_no_e_all_instances.csv')
	write_records(results['plural_no_e_strict'], 'plural_no_e_strict_instances.csv')
	write_records(results['strong_with_e_all'], 'strong_with_e_all_instances.csv')
	write_records(results['strong_with_e_strict'], 'strong_with_e_strict_instances.csv')
	write_records(results['strong_with_e_all_filtered'], 'strong_with_e_all_filtered_instances.csv')
	write_records(results['strong_with_e_strict_filtered'], 'strong_with_e_strict_filtered_instances.csv')

	# Write per-file accuracy breakdown
	accuracy_data = calculate_accuracy_rates(results)
	with open(os.path.join(output_dir, 'per_file_accuracy_breakdown.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = [
			'filename',
			'overall_accuracy', 'total_correct', 'total_instances',
			'overall_accuracy_filtered', 'total_correct_filtered', 'total_instances_filtered',
			'weak_accuracy', 'weak_correct', 'weak_total',
			'plural_accuracy', 'plural_correct', 'plural_total',
			'strong_accuracy', 'strong_correct', 'strong_total',
			'strong_accuracy_filtered', 'strong_correct_filtered', 'strong_total_filtered',
			'monosyllabic_adjectives_found'
		]
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		
		# Sort by filename
		accuracy_data.sort(key=lambda x: x['filename'])
		
		for row in accuracy_data:
			# Format percentages to 2 decimal places
			formatted_row = {}
			for key, value in row.items():
				if key.endswith('_accuracy') and value is not None:
					formatted_row[key] = f"{value:.2f}"
				else:
					formatted_row[key] = value
			writer.writerow(formatted_row)
	
	# Write per-headword accuracy breakdown
	headword_accuracy_data = calculate_headword_accuracy_rates(results)
	with open(os.path.join(output_dir, 'per_headword_accuracy_breakdown.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = [
			'headword', 'always_e',
			'overall_accuracy', 'total_correct', 'total_instances',
			'weak_accuracy', 'weak_correct', 'weak_total',
			'plural_accuracy', 'plural_correct', 'plural_total',
			'strong_accuracy', 'strong_correct', 'strong_total',
			'stem_forms', 'all_forms'
		]
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		
		# Sort by headword
		headword_accuracy_data.sort(key=lambda x: x['headword'])
		
		for row in headword_accuracy_data:
			# Format percentages to 2 decimal places
			formatted_row = {}
			for key, value in row.items():
				if key.endswith('_accuracy') and value is not None:
					formatted_row[key] = f"{value:.2f}"
				else:
					formatted_row[key] = value
			writer.writerow(formatted_row)
	
	# Write per-headword accuracy breakdown
	headword_data = calculate_headword_accuracy_rates(results)
	with open(os.path.join(output_dir, 'per_headword_accuracy_breakdown.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = [
			'headword',
			'overall_accuracy', 'total_correct', 'total_instances',
			'weak_accuracy', 'weak_correct', 'weak_total', 'weak_exceptions',
			'plural_accuracy', 'plural_correct', 'plural_total', 'plural_exceptions',
			'strong_accuracy', 'strong_correct', 'strong_total', 'strong_exceptions',
			'always_ends_in_e', 'stem_forms','always_e','all_forms'
		]
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		
		# Sort by headword
		headword_data.sort(key=lambda x: x['headword'])
		
		for row in headword_data:
			# Format percentages to 2 decimal places
			formatted_row = {}
			for key, value in row.items():
				if key.endswith('_accuracy') and value is not None:
					formatted_row[key] = f"{value:.2f}"
				else:
					formatted_row[key] = value
			writer.writerow(formatted_row)
	
	# Write per-headword accuracy breakdown
	headword_accuracy_data = calculate_headword_accuracy_rates(results)
	with open(os.path.join(output_dir, 'per_headword_accuracy_breakdown.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = [
			'headword',
			'overall_accuracy', 'total_correct', 'total_instances',
			'weak_accuracy', 'weak_correct', 'weak_total',
			'plural_accuracy', 'plural_correct', 'plural_total',
			'strong_accuracy', 'strong_correct', 'strong_total',
			'always_e', 'stem_forms', 'all_forms'
		]
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		
		# Sort by headword
		headword_accuracy_data.sort(key=lambda x: x['headword'])
		
		for row in headword_accuracy_data:
			# Format percentages to 2 decimal places
			formatted_row = {}
			for key, value in row.items():
				if key.endswith('_accuracy') and value is not None:
					formatted_row[key] = f"{value:.2f}"
				else:
					formatted_row[key] = value
			writer.writerow(formatted_row)
	
	# Calculate and write overall accuracy statistics
	overall_stats = {
		'weak_total': sum(s['weak_total'] for s in results['file_stats'].values()),
		'weak_correct': sum(s['weak_correct'] for s in results['file_stats'].values()),
		'plural_total': sum(s['plural_total'] for s in results['file_stats'].values()),
		'plural_correct': sum(s['plural_correct'] for s in results['file_stats'].values()),
		'strong_total': sum(s['strong_total'] for s in results['file_stats'].values()),
		'strong_correct': sum(s['strong_correct'] for s in results['file_stats'].values()),
		'strong_total_filtered': sum(s.get('strong_total_filtered', 0) for s in results['file_stats'].values()),
		'strong_correct_filtered': sum(s.get('strong_correct_filtered', 0) for s in results['file_stats'].values()),
	}

	# Write summary statistics
	with open(os.path.join(output_dir, 'summary_statistics.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Category', 'All Instances', 'Strict Instances (not elided, not final)', 'Accuracy Rate'])
		
		weak_acc = (overall_stats['weak_correct'] / overall_stats['weak_total'] * 100) if overall_stats['weak_total'] > 0 else 0
		writer.writerow(['Weak declension (should have -e)', 
						overall_stats['weak_total'], 
						len(results['weak_no_e_strict']),
						f"{weak_acc:.2f}%"])
		
		plural_acc = (overall_stats['plural_correct'] / overall_stats['plural_total'] * 100) if overall_stats['plural_total'] > 0 else 0
		writer.writerow(['Plural form (should have -e)', 
						overall_stats['plural_total'], 
						len(results['plural_no_e_strict']),
						f"{plural_acc:.2f}%"])
		
		strong_acc = (overall_stats['strong_correct'] / overall_stats['strong_total'] * 100) if overall_stats['strong_total'] > 0 else 0
		writer.writerow(['Strong form (should not have -e, all)', 
						overall_stats['strong_total'], 
						len(results['strong_with_e_strict']),
						f"{strong_acc:.2f}%"])
		
		strong_acc_filtered = (overall_stats['strong_correct_filtered'] / overall_stats['strong_total_filtered'] * 100) if overall_stats['strong_total_filtered'] > 0 else 0
		writer.writerow(['Strong form (should not have -e, filtered)', 
						overall_stats['strong_total_filtered'], 
						len(results['strong_with_e_strict_filtered']),
						f"{strong_acc_filtered:.2f}%"])
		
		writer.writerow(['Always -e adjectives found', len(results['always_e_adjectives']), 'N/A', 'N/A'])

	# Write monosyllabic adjectives found
	with open(os.path.join(output_dir, 'monosyllabic_adjectives.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Headword', 'Stem Forms', 'All Forms', 'Always Ends in E'])
		for headword, data in results['all_adjectives'].items():
			if is_monosyllabic_root(headword, data['stem_forms']):
				always_e = 'Yes' if headword in results['always_e_adjectives'] else 'No'
				writer.writerow([headword, '; '.join(sorted(data['stem_forms'])), '; '.join(sorted(data['all_forms'])), always_e])

	# Write list of always -e adjectives
	with open(os.path.join(output_dir, 'always_e_adjectives.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Headword', 'Stem Forms with -e', 'Reason'])
		for headword in sorted(results['always_e_adjectives']):
			data = results['all_adjectives'][headword]
			stem_forms = '; '.join(sorted(data['stem_with_e']))
			writer.writerow([headword, stem_forms, 'All stem forms end in -e'])

def write_strong_filtered_to_docx(results, output_dir):
	"""Write strong_with_e_strict_filtered_instances to Word document"""
	strong_doc = docx.Document()
	strong_doc.add_heading('Strong Form with -e Exceptions (Filtered)', 0)
	
	# Add summary paragraph
	para = strong_doc.add_paragraph()
	para.add_run(f"Total instances: {len(results['strong_with_e_strict_filtered'])}\n").bold = True
	para.add_run("These are monosyllabic adjectives in strong form that end in -e (excluding elided and final positions, and excluding adjectives that always end in -e).\n\n")
	
	# Group by headword for better organization
	by_headword = defaultdict(list)
	for record in results['strong_with_e_strict_filtered']:
		by_headword[record['headword']].append(record)
	
	# Write each headword's exceptions
	for headword in sorted(by_headword.keys()):
		strong_doc.add_heading(f"Headword: {headword}", level=2)
		
		records = by_headword[headword]
		para = strong_doc.add_paragraph()
		para.add_run(f"Count: {len(records)}\n\n").italic = True
		
		for record in records:
			# Add exception entry
			exception_para = strong_doc.add_paragraph()
			exception_para.add_run(f"Word: {record['word']}").bold = True
			exception_para.add_run(f"\nLine {record['line_number']} ({record['filename']})\n")
			
			# Add context with the word italicized
			for i, word_token in enumerate(record['context'].split(' ')):
				if word_token.lower() == record['word'].lower():
					exception_para.add_run(f"{word_token}").italic = True
				else:
					exception_para.add_run(f"{word_token}")
				if i < len(record['context'].split(' ')) - 1:
					exception_para.add_run(" ")
			exception_para.add_run("\n\n")
	
	strong_doc.save(os.path.join(output_dir, 'strong_with_e_filtered_exceptions.docx'))
	print(f"Strong form filtered exceptions document saved to: {output_dir}/strong_with_e_filtered_exceptions.docx")

def write_summary_to_docx(results, output_dir):
	"""Write summary statistics to Word document"""
	# Filter strong form exceptions first
	results = filter_strong_form_exceptions(results)
	
	summary_doc = docx.Document()
	summary_doc.add_heading('Gower Adjective Declension Analysis Summary', 0)
	
	monosyllabic_count = sum(1 for headword, data in results['all_adjectives'].items() 
						   if is_monosyllabic_root(headword, data['stem_forms']))
	
	# Overview section
	summary_doc.add_heading('Overview', 1)
	para = summary_doc.add_paragraph()
	para.add_run(f"Total adjectives found: {len(results['all_adjectives'])}\n")
	para.add_run(f"Monosyllabic adjectives found: {monosyllabic_count}\n")
	para.add_run(f"Always -e adjectives found: {len(results['always_e_adjectives'])}\n")
	para.add_run(f"Files processed: {len(results['file_stats'])}")
	
	# Calculate overall accuracy
	overall_stats = {
		'weak_total': sum(s['weak_total'] for s in results['file_stats'].values()),
		'weak_correct': sum(s['weak_correct'] for s in results['file_stats'].values()),
		'plural_total': sum(s['plural_total'] for s in results['file_stats'].values()),
		'plural_correct': sum(s['plural_correct'] for s in results['file_stats'].values()),
		'strong_total': sum(s['strong_total'] for s in results['file_stats'].values()),
		'strong_correct': sum(s['strong_correct'] for s in results['file_stats'].values()),
		'strong_total_filtered': sum(s.get('strong_total_filtered', 0) for s in results['file_stats'].values()),
		'strong_correct_filtered': sum(s.get('strong_correct_filtered', 0) for s in results['file_stats'].values()),
	}
	
	# Accuracy rates section
	summary_doc.add_heading('Overall Accuracy Rates', 1)
	para = summary_doc.add_paragraph()
	
	if overall_stats['weak_total'] > 0:
		weak_acc = (overall_stats['weak_correct'] / overall_stats['weak_total']) * 100
		para.add_run(f"Weak declension accuracy: {weak_acc:.2f}% ({overall_stats['weak_correct']}/{overall_stats['weak_total']})\n")
	
	if overall_stats['plural_total'] > 0:
		plural_acc = (overall_stats['plural_correct'] / overall_stats['plural_total']) * 100
		para.add_run(f"Plural form accuracy: {plural_acc:.2f}% ({overall_stats['plural_correct']}/{overall_stats['plural_total']})\n")
	
	if overall_stats['strong_total'] > 0:
		strong_acc = (overall_stats['strong_correct'] / overall_stats['strong_total']) * 100
		para.add_run(f"Strong form accuracy (all): {strong_acc:.2f}% ({overall_stats['strong_correct']}/{overall_stats['strong_total']})\n")
	
	if overall_stats['strong_total_filtered'] > 0:
		strong_acc_filtered = (overall_stats['strong_correct_filtered'] / overall_stats['strong_total_filtered']) * 100
		para.add_run(f"Strong form accuracy (filtered): {strong_acc_filtered:.2f}% ({overall_stats['strong_correct_filtered']}/{overall_stats['strong_total_filtered']})")
	
	# Statistics section
	summary_doc.add_heading('Exception Statistics', 1)
	
	categories = [
		('Weak declension without -e', 'weak_no_e_all', 'weak_no_e_strict'),
		('Plural form without -e', 'plural_no_e_all', 'plural_no_e_strict'),
		('Strong form with -e (all)', 'strong_with_e_all', 'strong_with_e_strict'),
		('Strong form with -e (filtered)', 'strong_with_e_all_filtered', 'strong_with_e_strict_filtered')
	]
	
	for desc, all_key, strict_key in categories:
		all_count = len(results[all_key])
		strict_count = len(results[strict_key])
		
		para = summary_doc.add_paragraph()
		para.add_run(desc + ":").bold = True
		para.add_run(f"\n  All instances: {all_count}")
		para.add_run(f"\n  Strict instances (not elided, not final): {strict_count}\n")
	
	# Always -e adjectives section
	if results['always_e_adjectives']:
		summary_doc.add_heading('Adjectives That Always End in -e in Strong Form', 1)
		para = summary_doc.add_paragraph()
		para.add_run("The following adjectives were excluded from strong form exception analysis because they always end in -e in the strong form:\n\n")
		
		for headword in sorted(results['always_e_adjectives']):
			data = results['all_adjectives'][headword]
			stem_forms = ', '.join(sorted(data['stem_with_e']))
			para.add_run(f"• {headword}").bold = True
			para.add_run(f" (forms: {stem_forms})\n")
	
	summary_doc.save(os.path.join(output_dir, 'analysis_summary.docx'))

# Main execution
if __name__ == "__main__":
	print("Processing Gower CSVs...")
	
	# Process all CSV files
	results = process_csv_directory(base_csv_dir)
	
	# Write results
	output_dir = 'gower_adjective_analysis_output'
	write_results_to_csv(results, output_dir)
	write_summary_to_docx(results, output_dir)
	
	# Save Word document with exceptions
	doc.save(os.path.join(output_dir, 'declension_exceptions.docx'))
	
	# Create Word document for strong form filtered exceptions
	write_strong_filtered_to_docx(results, output_dir)
	
	# Write completion log to file
	with open(os.path.join(output_dir, 'analysis_log.txt'), 'w', encoding='utf-8') as log_file:
		log_file.write("Gower Adjective Declension Analysis - Complete\n")
		log_file.write("="*50 + "\n\n")
		
		monosyllabic_count = sum(1 for headword, data in results['all_adjectives'].items() 
							   if is_monosyllabic_root(headword, data['stem_forms']))
		
		# Filter results for logging
		results = filter_strong_form_exceptions(results)
		
		log_file.write(f"Total adjectives found: {len(results['all_adjectives'])}\n")
		log_file.write(f"Monosyllabic adjectives found: {monosyllabic_count}\n")
		log_file.write(f"Always -e adjectives found: {len(results['always_e_adjectives'])}\n")
		log_file.write(f"Files processed: {len(results['file_stats'])}\n\n")
		
		log_file.write("Files generated:\n")
		files = [
			"weak_no_e_all_instances.csv - All weak declension forms without -e",
			"weak_no_e_strict_instances.csv - Strict weak declension exceptions", 
			"plural_no_e_all_instances.csv - All plural forms without -e",
			"plural_no_e_strict_instances.csv - Strict plural exceptions",
			"strong_with_e_all_instances.csv - All strong forms with -e",
			"strong_with_e_strict_instances.csv - Strict strong form exceptions",
			"strong_with_e_all_filtered_instances.csv - Strong forms with -e (filtered)",
			"strong_with_e_strict_filtered_instances.csv - Strict strong exceptions (filtered)",
			"summary_statistics.csv - Overall counts and accuracy rates",
			"per_file_accuracy_breakdown.csv - Accuracy rates broken down by file",
			"per_headword_accuracy_breakdown.csv - Accuracy rates broken down by headword",
			"monosyllabic_adjectives.csv - List of all monosyllabic adjectives found",
			"always_e_adjectives.csv - List of adjectives that always end in -e",
			"analysis_summary.docx - Summary statistics document",
			"declension_exceptions.docx - Detailed exception examples",
			"strong_with_e_filtered_exceptions.docx - Strong form filtered exceptions document",
			"analysis_log.txt - This log file"
		]
		
		for file_desc in files:
			log_file.write(f"- {file_desc}\n")
	
	print(f"\nAnalysis complete! Results saved to '{output_dir}' directory.")
	print(f"Check 'per_file_accuracy_breakdown.csv' for file-by-file accuracy rates.")
	print(f"Check 'per_headword_accuracy_breakdown.csv' for headword-by-headword accuracy rates.")
	print(f"Check 'per_headword_accuracy_breakdown.csv' for per-headword accuracy rates.")
