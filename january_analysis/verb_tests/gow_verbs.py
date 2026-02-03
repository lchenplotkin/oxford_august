import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict, Counter

# Initialize Word document for exceptions
exceptions_doc = docx.Document()
exceptions_doc.add_heading('Verb Declension Exceptions', 0)

# Configuration
base_csv_dir = '../gow_csvs'
form_csv = 'complete_verbs.csv'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
					 "hadden", "his", "her", "him", "hers", "hide", "hir",
					 "hire", "hires", "hirs", "han"]

# --- Load verb classifications ---
verbs_dict = {}
with open(form_csv, 'r', encoding='utf-8') as csvfile:
	reader = csv.DictReader(csvfile)
	for row in reader:
		if 'headword' in row and 'classification' in row:
			verbs_dict[row['headword'].lower()] = row['classification'].lower()

def is_strong(verb):
	return verbs_dict.get(verb, '') == 'strong'

def is_weak(verb):
	return verbs_dict.get(verb, '') == 'weak'

def is_elided(word, next_word):
	return next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS

def clean_tag(tag):
	if tag.startswith('ger') and '%' not in tag:
		return 'ger'
	return re.sub(r'\d+(?=%)', '', tag)

def parse_tagged_text(text, tagging):
	if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
		return [], [], []

	words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()
	tag_tokens = tagging.strip().split()
	tags, headwords = [], []

	for token in tag_tokens:
		if '@' in token:
			parts = token.split('@')
			headwords.append(parts[0].lower())
			tags.append(parts[1] if len(parts) > 1 else '')

	min_len = min(len(words), len(tags), len(headwords))
	return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(doc, reason, word, text, line_number, filename):
	para = doc.add_paragraph()
	para.add_run(f"{reason}: {word}").bold = True
	para.add_run(f"\n{line_number} ({filename})\n")

	for tok in text.split():
		run = para.add_run(tok + " ")
		if tok.lower() == word.lower():
			run.italic = True
	para.add_run("\n")

def classify_ending(word):
	endings = ['eth','en','ede','de','te','e','ed','d','est','et','t','ing','n']
	for e in endings:
		if word.endswith(e):
			return f"-{e}"
	if word[-1] in 'aeiou':
		return 'vowel'
	return 'other'

# ---------------- ANALYSIS ---------------- #

def analyze_verbs(df, results, doc, source_file):
	for _, row in df.iterrows():
		words, headwords, tags = parse_tagged_text(row['TEXT'], row['TAGGING'])
		line_number = row['LINE_NUMBER']
		filename = row['FILENAME']
		original_text = row['OG_TEXT']

		for i in range(len(tags)):
			if i >= len(words): continue

			word = words[i]
			headword = headwords[i]
			tag = clean_tag(tags[i])
			next_word = words[i+1] if i+1 < len(words) else 'END'

			if is_elided(word, next_word) or next_word == 'END' or not tag.startswith('v'):
				continue

			ending = classify_ending(word)
			verb_class = 'strong' if is_strong(headword) else 'weak' if is_weak(headword) else 'irregular'

			results['ending_counts'][tag][ending] += 1
			results['word_tag_counts'][headword][tag][word] += 1

			# Rule checks
			violated = False
			reason = ""
			if tag == 'v%inf' and not (word.endswith(('a','i','o','u')) or word.endswith('en') or word.endswith('e') or word.endswith('vowel')):
				violated = True
				reason += "RULE 1: The infinitive ends in -en or -e unless the stem ends in a vowel."
			if tag.startswith('v%imp') and word.endswith('e'):
				violated = True
				reason += "RULE 3: imperative (singular) are always endingless"
			if tag == 'v%pt_pl' and not (word.endswith('en') or word.endswith('e')): 
				violated = True
				reason += "RULE 4: Preterite plural verbs end in -e/-en"
			if tag == 'v%pr_1' and not (word.endswith('en') or word.endswith('e')):
				violated = True
				reason += "RULE 5: First person singular present tense verbs end in -e/-en"
			if tag == 'v%pr_pl' and not (word.endswith('en') or word.endswith('e')):
				violated = True
				reason += "RULE 6: Present plural verbs end in -e/-en"
			if verb_class == "weak" and tag in ['v%pt_1', 'v%pt_3'] and not (word.endswith('de') or word.endswith('te') or word.endswith('d') or word.endswith('t')): 
				violated = True
				reason += "Rule 7: Singular weak preterite verbs in 1st/3rd person end in -t(e)/-d(e)"
			if verb_class == "strong" and tag == 'v%ppl' and not (word.endswith('e') or word.endswith('en')):
				violated = True
				reason += "Rule 8: Past participles of strong verbs end in -e/-en"
			if verb_class in ["strong", "weak"] and tag.startswith("v%prp") and not (word.endswith('e') or word.endswith('en')):
				violated = True
				reason += "Rule 9: Present participle of strong and weak verbs end in -e"
			if verb_class == "strong" and tag == "v%pt_2" and not word.endswith('e'):
				violated = True
				reason += "Rule 10: 2nd person singular preterit of strong verbs end in -e"
			if verb_class == "strong" and tag in ["v%pt_1", "v%pt_3"] and word.endswith('e'):
				violated = True
				reason += "Rule 13: 1st and 3rd person singular strong preterites do not end in -e (Barney)"
			if verb_class == "weak" and tag=="v%pt_2" and word.endswith('e'):		
				violated = True
				reason += "Rule 14: 2nd person singular weak preterite do not end in -e (Barney)"
			if verb_class == "weak" and tag=="v%ppl" and word.endswith('e'):
				violated = True
				reason += "Rule 15: The weak past participle does not end in -e (Barney)"

			if violated:
				record = {
					'headword': headword,
					'word': word,
					'tag': tag,
					'line': line_number,
					'file': filename,
					'context': original_text
				}
				results['exceptions'].append(record)
				add_exception_to_doc(doc, reason, word, original_text, line_number, filename)
			else:
				results['successes'][headword][tag][word] += 1

	return results

# ---------------- DIRECTORY PROCESSING ---------------- #

def process_csv_directory(csv_dir, doc):
	results = {
		'exceptions': [],
		'ending_counts': defaultdict(Counter),
		'word_tag_counts': defaultdict(lambda: defaultdict(Counter)),
		'successes': defaultdict(lambda: defaultdict(Counter))
	}

	for root, _, files in os.walk(csv_dir):
		for file in files:
			if not file.endswith('.csv'):
				continue
			path = os.path.join(root, file)
			try:
				df = pd.read_csv(path, encoding='utf-8')
				required = ['TEXT','TAGGING','LINE_NUMBER','FILENAME','OG_TEXT']
				if not all(col in df.columns for col in required):
					print(f"Skipping {file}: missing columns")
					continue
				results = analyze_verbs(df, results, doc, file)
			except Exception as e:
				print(f"Error processing {file}: {e}")

	return results

# ---------------- OUTPUT ---------------- #

def write_results(results, output_dir):
	os.makedirs(output_dir, exist_ok=True)

	# --- Aggregate exceptions ---
	agg_exceptions = defaultdict(lambda: {
		'count': 0,
		'files': set(),
		'contexts': []
	})

	for rec in results['exceptions']:
		key = (rec['headword'], rec['word'], rec['tag'])
		agg_exceptions[key]['count'] += 1
		agg_exceptions[key]['files'].add(rec['file'])

		if len(agg_exceptions[key]['contexts']) < 3:
			agg_exceptions[key]['contexts'].append(f"{rec['line']}: {rec['context']}")

	# --- Write aggregated exceptions CSV ---
	with open(os.path.join(output_dir, 'verb_exceptions.csv'), 'w', newline='', encoding='utf-8') as f:
		fieldnames = [
			'headword', 'word', 'class', 'tag',
			'pct_of_headword_tag', 'exception_count',
			'files', 'sample_contexts'
		]
		writer = csv.DictWriter(f, fieldnames=fieldnames)
		writer.writeheader()

		for (headword, word, tag), data in sorted(agg_exceptions.items()):
			# Determine verb class
			if is_strong(headword):
				verb_class = 'strong'
			elif is_weak(headword):
				verb_class = 'weak'
			else:
				verb_class = 'irregular'

			# % of headword+tag tokens that are this spelling
			total_tokens = sum(results['word_tag_counts'][headword][tag].values())
			word_tokens = results['word_tag_counts'][headword][tag][word]
			pct = f"{(word_tokens / total_tokens) * 100:.2f}%" if total_tokens else "0.00%"

			writer.writerow({
				'headword': headword,
				'word': word,
				'class': verb_class,
				'tag': tag,
				'pct_of_headword_tag': pct,
				'exception_count': data['count'],
				'files': '; '.join(sorted(data['files'])),
				'sample_contexts': ' | '.join(data['contexts'])
			})

	# --- Write successes (unchanged logic, but sorted & % added) ---
	with open(os.path.join(output_dir, 'verb_successes.csv'), 'w', newline='', encoding='utf-8') as f:
		fieldnames = ['headword', 'word', 'class', 'tag', 'pct_of_headword_tag', 'success_count']
		writer = csv.DictWriter(f, fieldnames=fieldnames)
		writer.writeheader()

		rows = []
		for headword, tag_data in results['successes'].items():
			for tag, word_counts in tag_data.items():
				for word, count in word_counts.items():
					if is_strong(headword):
						verb_class = 'strong'
					elif is_weak(headword):
						verb_class = 'weak'
					else:
						verb_class = 'irregular'

					total_tokens = sum(results['word_tag_counts'][headword][tag].values())
					word_tokens = results['word_tag_counts'][headword][tag][word]
					pct = f"{(word_tokens / total_tokens) * 100:.2f}%" if total_tokens else "0.00%"

					rows.append({
						'headword': headword,
						'word': word,
						'class': verb_class,
						'tag': tag,
						'pct_of_headword_tag': pct,
						'success_count': count
					})

		for row in sorted(rows, key=lambda x: (x['headword'], x['tag'], x['word'])):
			writer.writerow(row)
# ---------------- MAIN ---------------- #

if __name__ == "__main__":
	print("Processing Gower CSVs...")
	results = process_csv_directory(base_csv_dir, exceptions_doc)

	output_dir = 'gow_verb_analysis_output'
	write_results(results, output_dir)
	exceptions_doc.save(os.path.join(output_dir, 'verb_declension_exceptions.docx'))

	print("\nAnalysis complete!")
	print(f"Results written to: {output_dir}/")

