import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict, Counter

# Initialize Word document for exceptions
gower_doc = docx.Document()
gower_doc.add_heading('Gower Verb Declension Exceptions', 0)

# Configuration
base_csv_dir = '../gow_csvs'
form_csv = 'complete_verbs_pretpres.csv'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
					"hadden", "his", "her", "him", "hers", "hide", "hir",
					"hire", "hires", "hirs", "han"]

# Load verb classifications
verbs_dict = {}
with open(form_csv, 'r', encoding='utf-8') as csvfile:
	reader = csv.DictReader(csvfile)
	for row in reader:
		if 'headword' in row and 'classification' in row:
			verbs_dict[row['headword'].lower()] = row['classification'].lower()

def is_strong(verb):
	return verbs_dict.get(verb, '') == 'strong'

def is_weak(verb):
	return verbs_dict.get(verb, '') == 'weak'

def is_pretpres(verb):
	return verbs_dict.get(verb, '') == 'pretpres'

def is_elided(word, next_word):
	"""Check if word is elided before next word"""
	if next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS:
		return True
	return False

def clean_tag(tag):
	"""Remove digits before % in a tag (e.g., v2%pt_1 -> v%pt_1)."""
	if tag.startswith('ger') and '%' not in tag:
		return 'ger'
	return re.sub(r'\d+(?=%)', '', tag)

def parse_tagged_text(text, tagging):
	"""Extract words from text and tags from tagging"""
	if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
		return [], [], []

	words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()
	tag_tokens = tagging.strip().split()
	tags = []
	headwords = []

	for token in tag_tokens:
		if '@' in token and token != "--@dash":
			parts = token.split('@')
			headword = parts[0].lower()
			tag = parts[1] if len(parts) > 1 else ''
			headwords.append(headword)
			tags.append(tag)

	min_len = min(len(words), len(tags), len(headwords))
	return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(doc, exception_type, word, text, line_number, filename):
	"""Add an exception case to the Word document"""
	doc_para = doc.add_paragraph()
	doc_para.add_run(f"{exception_type}: {word}").bold = True
	doc_para.add_run(f"\n{line_number} ({filename})\n")
	for i, ox_word in enumerate(text.split(' ')):
		if ox_word.lower() == word.lower():
			doc_para.add_run(f"{ox_word}").italic = True
		else:
			doc_para.add_run(f"{ox_word}")
		if i < len(text.split(' ')) - 1:
			doc_para.add_run(" ")
	doc_para.add_run("\n\n")

def classify_ending(word):
	"""Classify the word ending into categories."""
	endings = ['eth', 'en', 'ede', 'de', 'te', 'e', 'ed', 'd', 'est', 'et', 't', 'ing', 'n']
	for ending in endings:
		if word.endswith(ending):
			return f'-{ending}'
	if word.endswith(('a', 'e', 'i', 'o', 'u')):
		return 'vowel'
	return 'other'

# Global counters for debugging
rule_1_violations = 0
inf_e_ending = 0
inf_en_ending = 0
inf_fail = 0
inf_vowel_ending = 0
pt_pl_e_ending = 0
pt_pl_en_ending = 0
pt_pl_fail = 0
pt_weak_success = 0
pt_weak_endings = {'te': 0, 'de': 0, 'd': 0, 't': 0, 'x': 0}
pt_pl_fail_dict = {}

def analyze_verbs(df, results, doc, current_filename):
	"""Analyze verb patterns in CSV data according to rules."""
	global inf_vowel_ending, rule_1_violations, inf_fail, inf_e_ending
	global pt_weak_endings, inf_en_ending, pt_pl_fail, pt_pl_e_ending
	global pt_pl_en_ending, pt_pl_fail_dict

	for idx, row in df.iterrows():
		text = row['TEXT']
		tagging = row['TAGGING']
		line_number = row['LINE_NUMBER']
		filename = row['FILENAME']
		original_text = row.get('OG_TEXT', text)

		words, headwords, tags = parse_tagged_text(text, tagging)

		for j in range(len(tags)):
			if j >= len(words) or j >= len(headwords):
				continue

			word = words[j].lower()
			headword = headwords[j]
			tag = clean_tag(tags[j])
			next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'

			# Skip if elided
			if is_elided(word, next_word):
				continue
			if next_word == 'END':
				continue
			if not tag.startswith('v'):
				continue

			ending = classify_ending(word)
			verb_class = 'weak'
			if headword in verbs_dict:
				if is_strong(headword):
					verb_class = 'strong'
				elif is_weak(headword):
					verb_class = 'weak'
				elif is_pretpres(headword):
					verb_class = 'pretpres'
				else:
					verb_class = 'irregular'

			# Count ending distribution
			results['ending_counts'][tag][ending] += 1
			results['verb_tag_counts'][headword][tag][ending] += 1
			results['ending_counts_by_class'][(tag, verb_class)][ending] += 1
			results['file_ending_counts_by_class'][current_filename][(tag, verb_class)][ending] += 1
			results['file_ending_counts'][current_filename][tag][ending] += 1
			results['file_verb_tag_counts'][current_filename][headword][tag][ending] += 1

			# Track exact word counts per headword+tag
			results['word_tag_counts'][headword][tag][word] += 1

			# Rule checks
			violated = False
			reason = ""
			
			# RULE 1: Infinitive endings
			if tag == 'v%inf' and headword not in ["ben", 'gon', 'don', 'fordon', 'forgon', "vouchen_sauf"]:
				if word.endswith('en'):
					inf_en_ending += 1
				elif word.endswith('e'):
					inf_e_ending += 1
				else:
					inf_fail += 1
					violated = True
					reason += "RULE 1: The infinitive ends in -en or -e."
					rule_1_violations += 1

			# RULE 2: Preterite plural
			if tag == 'v%pt_pl' and not (word.endswith('en') or word.endswith('e')):
				violated = True
				reason += "RULE 2: Preterite plural verbs end in -e/-en"

			# RULE 3: First person singular present
			if tag == 'v%pr_1' and verb_class in ['weak', 'strong'] and not (word.endswith('en') or word.endswith('e')):
				violated = True
				reason += "RULE 3: First person singular present tense weak and strong verbs end in -e/-en"

			# RULE 4: Present plural
			if tag == 'v%pr_pl' and verb_class in ['weak', 'strong'] and not (word.endswith('en') or word.endswith('e')):
				violated = True
				reason += "RULE 4: Present plural weak and strong verbs end in -e/-en"

			# RULE 5: Weak preterite 1st/3rd person
			if verb_class == "weak" and tag in ['v%pt_1', 'v%pt_3']:
				for ending_check in ['de', 'te', 'd', 't']:
					if word.endswith(ending_check):
						pt_weak_endings[ending_check] += 1
				if word.endswith('ede'):
					print(word, tag, original_text, line_number)
				if not (word.endswith('de') or word.endswith('te') or word.endswith('d') or word.endswith('t')):
					violated = True
					reason += "RULE 5: Singular weak preterite verbs in 1st/3rd person end in -t(e)/-d(e)"
					pt_weak_endings['x'] += 1

			# RULE 6: Preterite-present verbs
			if verb_class == 'pretpres' and tag in ['v%pt_1', 'v%pt_3']:
				if not (word.endswith('de') or word.endswith('te') or word.endswith('d') or word.endswith('t')):
					violated = True
					reason += "RULE 6: Preterite-present verbs should act like weak verbs in the past tense, and should in 1st/3rd person preterite end in -t(e)/-d(e)"

			# RULE 7: Strong verb past participles
			if verb_class == "strong" and tag == 'v%ppl' and not (word.endswith('e') or word.endswith('en')):
				violated = True
				reason += "RULE 7: Past participles of strong verbs end in -e/-en"

			# RULE 8: Weak verb past participles
			if verb_class == "weak" and tag == 'v%ppl' and (word.endswith('e') or word.endswith('en')):
				violated = True
				reason += "RULE 8: Past participles of weak verbs do not end in -e/-en"

			# RULE 9: Strong verb 2nd person preterite
			if verb_class == "strong" and tag == "v%pt_2" and not word.endswith('e'):
				violated = True
				reason += "RULE 9: 2nd person singular preterit of strong verbs end in -e"

			# RULE 10: Strong verb 1st/3rd person preterite
			if verb_class == "strong" and tag in ["v%pt_1", "v%pt_3"] and word.endswith('e'):
				violated = True
				reason += "RULE 10: 1st and 3rd person singular strong preterites do not end in -e (Barney)"

			# RULE 11: Weak verb 2nd person preterite
			if verb_class == "weak" and tag == "v%pt_2" and word.endswith('e'):
				violated = True
				reason += "RULE 11: 2nd person singular weak preterite do not end in -e (Barney)"

			# RULE 12: Weak past participle
			if verb_class == "weak" and tag == "v%ppl" and word.endswith('e'):
				violated = True
				reason += "RULE 12: The weak past participle does not end in -e (Barney)"

			if violated:
				record = {
					'headword': headword,
					'word': word,
					'tag': tag,
					'line_number': line_number,
					'filename': filename,
					'context': original_text,
					'source_file': current_filename
				}
				results['exceptions'].append(record)
				results['file_exceptions'][current_filename].append(record)
				add_exception_to_doc(doc, reason, word, original_text, line_number, filename)

	return results

def process_csv_directory(csv_dir, doc):
	"""Process all CSV files in directory"""
	results = {
		'exceptions': [],
		'ending_counts': defaultdict(Counter),
		'verb_tag_counts': defaultdict(lambda: defaultdict(Counter)),
		'ending_counts_by_class': defaultdict(Counter),
		'file_exceptions': defaultdict(list),
		'file_ending_counts': defaultdict(lambda: defaultdict(Counter)),
		'file_verb_tag_counts': defaultdict(lambda: defaultdict(lambda: defaultdict(Counter))),
		'file_ending_counts_by_class': defaultdict(lambda: defaultdict(Counter)),
		'word_tag_counts': defaultdict(lambda: defaultdict(Counter))
	}

	file_count = 0
	for root, dirs, files in os.walk(csv_dir):
		for file in files:
			if not file.endswith('.csv'):
				continue
			file_count += 1
			try:
				df = pd.read_csv(os.path.join(root, file), encoding='utf-8')
				required_columns = ['TAGGING', 'TEXT', 'LINE_NUMBER', 'FILENAME']
				if not all(col in df.columns for col in required_columns):
					print(f"Skipping {file}: missing columns")
					continue
				results = analyze_verbs(df, results, doc, file)
			except Exception as e:
				print(f"Error processing {file}: {e}")
				continue
	
	print(f"Processed {file_count} files")
	return results

def write_results(results, output_dir):
	"""Write all result files"""
	os.makedirs(output_dir, exist_ok=True)

	# --- Aggregate exceptions by headword-word-tag ---
	agg_exceptions = defaultdict(lambda: {'count': 0, 'files': set(), 'contexts': []})
	for rec in results['exceptions']:
		key = (rec['headword'], rec['word'], rec['tag'])
		agg_exceptions[key]['count'] += 1
		agg_exceptions[key]['files'].add(rec['filename'])
		if len(agg_exceptions[key]['contexts']) < 3:
			agg_exceptions[key]['contexts'].append(f"{rec['line_number']}: {rec['context']}")

	# --- Exceptions CSV (aggregated) ---
	with open(os.path.join(output_dir, 'gower_exceptions.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['headword', 'word', 'class', 'tag', 'pct_of_headword_tag', 'exception_count', 'files', 'sample_contexts']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		
		for (headword, word, tag), data in sorted(agg_exceptions.items()):
			verb_class = 'weak'
			if headword in verbs_dict:
				if is_strong(headword):
					verb_class = 'strong'
				elif is_weak(headword):
					verb_class = 'weak'
				elif is_pretpres(headword):
					verb_class = 'pretpres'
				else:
					verb_class = 'irregular'

			# Compute % of headword-tag tokens that are this word
			total_tokens = sum(results['word_tag_counts'][headword][tag].values())
			word_tokens = results['word_tag_counts'][headword][tag][word]
			pct = f"{(word_tokens / total_tokens) * 100:.2f}%" if total_tokens > 0 else "0.00%"
			
			writer.writerow({
				'headword': headword,
				'word': word,
				'class': verb_class,
				'tag': tag,
				'pct_of_headword_tag': pct,
				'exception_count': data['count'],
				'files': '; '.join(sorted(data['files'])),
				'sample_contexts': ' | '.join(data['contexts'])
			})

	print(f"Exceptions written to {output_dir}/gower_exceptions.csv")

# Main execution
if __name__ == "__main__":
	print("Processing Gower texts...")
	gower_results = process_csv_directory(base_csv_dir, gower_doc)

	# Print debug statistics
	print(f"\nInfinitive statistics:")
	print(f"  -e endings: {inf_e_ending}")
	print(f"  -en endings: {inf_en_ending}")
	print(f"  vowel endings: {inf_vowel_ending}")
	print(f"  failures: {inf_fail}")
	print(f"\nPreterite plural: e={pt_pl_e_ending}, en={pt_pl_en_ending}, fail={pt_pl_fail}")
	print(f"\nWeak preterite endings: {pt_weak_endings}")

	output_dir = 'gower_verb_analysis_output'
	write_results(gower_results, output_dir)
	gower_doc.save(os.path.join(output_dir, 'gower_declension_exceptions.docx'))

	print("\nAnalysis complete!")
	print(f"Results written to: {output_dir}/")
