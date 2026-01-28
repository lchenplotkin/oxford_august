import os
import re
import csv
import pandas as pd
import docx
from collections import defaultdict
from rules import RULES

# =========================================================
# SETTINGS
# =========================================================

DATASET_DIR = "../dataset"

ELISION_FOLLOWERS = {
	"have","haven","haveth","havest","had","hadde","hadden",
	"his","her","him","hers","hide","hir","hire","hires","hirs","han"
}

# =========================================================
# DOCX SETUP
# =========================================================

oxford_doc = docx.Document()
oxford_doc.add_heading('Oxford Verb Rule Exceptions', 0)

riverside_doc = docx.Document()
riverside_doc.add_heading('Riverside Verb Rule Exceptions', 0)

combined_doc = docx.Document()
combined_doc.add_heading('Combined Verb Rule Exceptions', 0)


def add_exception_to_doc(doc, rule, token, context, line_number, filename):
	para = doc.add_paragraph()
	para.add_run(f"{rule.rule_id}: {rule.description}").bold = True
	para.add_run(f"\nLine {line_number} ({filename})\n")

	for w in str(context).split():
		run = para.add_run(w + " ")
		if w.lower().strip(".,;:!?") == token["word"]:
			run.italic = True
	para.add_run("\n")


# =========================================================
# VERB CLASS LOOKUP
# =========================================================

verbs_dict = {}
with open("complete_verbs.csv", encoding="utf-8") as f:
	reader = csv.DictReader(f)
	for row in reader:
		verbs_dict[row["headword"].lower()] = row["classification"].lower()


def get_verb_class(headword):
	cls = verbs_dict.get(headword, "weak")
	if cls not in ["strong", "weak"]:
		return "irregular"
	return cls


# =========================================================
# TEXT PROCESSING
# =========================================================

def is_elided(next_word):
	return next_word and (next_word[0] in "aeiou" or next_word in ELISION_FOLLOWERS)


def clean_tag(tag):
	return re.sub(r'\d+(?=%)', '', tag)


def parse_tagged_text(text, tagging):
	if pd.isna(text) or pd.isna(tagging):
		return [], [], []

	words = re.sub(r'[.,!?°¶]', '', str(text).lower()).split()
	tag_tokens = str(tagging).split()

	headwords, tags = [], []
	for tok in tag_tokens:
		if '@' not in tok:
			continue
		h, t = tok.split('@', 1)
		headwords.append(h.lower())
		tags.append(clean_tag(t))

	min_len = min(len(words), len(headwords), len(tags))
	return words[:min_len], headwords[:min_len], tags[:min_len]


# =========================================================
# STATS
# =========================================================

rule_stats = defaultdict(lambda: defaultdict(lambda: {
	"applied_raw": 0, "passed_raw": 0,
	"applied_excl": 0, "passed_excl": 0
}))

word_rule_stats = defaultdict(lambda: defaultdict(lambda: {
	"applied": 0, "passed": 0
}))

exceptions_by_corpus = {"OXFORD": 0, "RIVERSIDE": 0}


# =========================================================
# RULE ENGINE
# =========================================================

def process_text_stream(words, headwords, tags, filename, text_type, context, line_number):
	for i in range(len(tags) - 1):
		if is_elided(words[i + 1]):
			continue

		token = {
			"word": words[i],
			"headword": headwords[i],
			"tag": tags[i],
			"verb_class": get_verb_class(headwords[i]),
			"stem_ends_vowel": headwords[i][-1] in "aeiou",
		}

		key = (filename, text_type)

		for rule in RULES:
			if not rule.applies(token):
				continue

			stats = rule_stats[key][rule.rule_id]
			stats["applied_raw"] += 1

			passed = rule.check(token)
			if passed:
				stats["passed_raw"] += 1
			else:
				exceptions_by_corpus[text_type] += 1
				if text_type == "OXFORD":
					add_exception_to_doc(oxford_doc, rule, token, context, line_number, filename)
				else:
					add_exception_to_doc(riverside_doc, rule, token, context, line_number, filename)

			w = word_rule_stats[token["word"]][rule.rule_id]
			w["applied"] += 1
			if passed:
				w["passed"] += 1

			if not rule.is_excluded(token["headword"]):
				stats["applied_excl"] += 1
				if passed:
					stats["passed_excl"] += 1


# =========================================================
# FILE PROCESSING
# =========================================================

def process_file(df, filename):
	for _, row in df.iterrows():
		line_number = row.get("LINE_NUMBER", "")

		ox_words, ox_heads, ox_tags = parse_tagged_text(row.get("OXFORD_TEXT"), row.get("OXFORD_TAGGING"))
		process_text_stream(ox_words, ox_heads, ox_tags, filename, "OXFORD",
							row.get("OG_OXFORD_TEXT", ""), line_number)

		rv_words, rv_heads, rv_tags = parse_tagged_text(row.get("RIVERSIDE_TEXT"), row.get("RIVERSIDE_TAGGING"))
		process_text_stream(rv_words, rv_heads, rv_tags, filename, "RIVERSIDE",
							row.get("OG_RIVERSIDE_TEXT", ""), line_number)


# =========================================================
# CSV OUTPUTS
# =========================================================

def write_rule_file_success():
	with open("rule_success_by_file.csv", "w", newline="", encoding="utf-8") as f:
		writer = csv.writer(f)

		header = ["Filename", "Text_Type"]
		for r in RULES:
			header += [
				f"{r.rule_id}_RAW_rate", f"{r.rule_id}_RAW_pass",
				f"{r.rule_id}_RAW_fail", f"{r.rule_id}_RAW_total",
				f"{r.rule_id}_EXCL_rate", f"{r.rule_id}_EXCL_pass",
				f"{r.rule_id}_EXCL_fail", f"{r.rule_id}_EXCL_total"
			]
		writer.writerow(header)

		totals = defaultdict(lambda: {"applied_raw":0,"passed_raw":0,"applied_excl":0,"passed_excl":0})

		for (filename, text_type), rules_dict in rule_stats.items():
			row = [filename, text_type]

			for r in RULES:
				s = rules_dict[r.rule_id]

				raw_total = s["applied_raw"]
				raw_pass = s["passed_raw"]
				raw_fail = raw_total - raw_pass
				raw_rate = raw_pass / raw_total if raw_total else ""

				excl_total = s["applied_excl"]
				excl_pass = s["passed_excl"]
				excl_fail = excl_total - excl_pass
				excl_rate = excl_pass / excl_total if excl_total else ""

				row += [raw_rate, raw_pass, raw_fail, raw_total,
						excl_rate, excl_pass, excl_fail, excl_total]

				for k in totals[r.rule_id]:
					totals[r.rule_id][k] += s[k]

			writer.writerow(row)

		agg = ["ALL_FILES", "BOTH"]
		for r in RULES:
			t = totals[r.rule_id]

			raw_total = t["applied_raw"]
			raw_pass = t["passed_raw"]
			raw_fail = raw_total - raw_pass
			raw_rate = raw_pass / raw_total if raw_total else ""

			excl_total = t["applied_excl"]
			excl_pass = t["passed_excl"]
			excl_fail = excl_total - excl_pass
			excl_rate = excl_pass / excl_total if excl_total else ""

			agg += [raw_rate, raw_pass, raw_fail, raw_total,
					excl_rate, excl_pass, excl_fail, excl_total]

		writer.writerow(agg)


def write_word_success():
	with open("rule_success_by_word.csv", "w", newline="", encoding="utf-8") as f:
		writer = csv.writer(f)

		header = ["Word_Form"]
		for r in RULES:
			header += [
				f"{r.rule_id}_rate",
				f"{r.rule_id}_pass",
				f"{r.rule_id}_fail",
				f"{r.rule_id}_total"
			]
		writer.writerow(header)

		for word, rules_dict in sorted(word_rule_stats.items()):
			row = [word]

			for r in RULES:
				s = rules_dict.get(r.rule_id)
				if not s or s["applied"] == 0:
					row += ["", "", "", ""]
				else:
					total = s["applied"]
					passed = s["passed"]
					failed = total - passed
					rate = passed / total
					row += [rate, passed, failed, total]

			writer.writerow(row)

# =========================================================
# SAVE DOCX
# =========================================================

def save_docs():
	oxford_doc.save("oxford_rule_exceptions.docx")
	riverside_doc.save("riverside_rule_exceptions.docx")

	combined_doc.add_heading("Exception Totals", level=1)
	combined_doc.add_paragraph(f"Oxford exceptions: {exceptions_by_corpus['OXFORD']}")
	combined_doc.add_paragraph(f"Riverside exceptions: {exceptions_by_corpus['RIVERSIDE']}")
	combined_doc.save("combined_rule_exceptions.docx")


# =========================================================
# RUN
# =========================================================

if __name__ == "__main__":
	for root, _, files in os.walk(DATASET_DIR):
		for file in files:
			if file.endswith("complete.csv"):
				path = os.path.join(root, file)
				df = pd.read_csv(path)
				process_file(df, file)

	write_rule_file_success()
	write_word_success()
	save_docs()

	print("Verb rule analysis complete.")

