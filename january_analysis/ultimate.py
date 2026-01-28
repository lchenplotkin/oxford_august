# Merged Chaucer Adjective Declension Analysis
import docx 
import os
import re
import string
import csv
import pandas as pd
from collections import defaultdict, Counter
import matplotlib.pyplot as plt

# Initialize Word document for exceptions
doc = docx.Document()
doc.add_heading('Adjective Declension Exceptions in the Oxford Chaucer', 0)

# Configuration
base_csv_dir = 'for_gui/done'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
                    "hadden", "his", "her", "him", "hers", "hide", "hir",
                    "hire", "hires", "hirs", "han"]

monosyllabic_set = set()
mono_csv = 'data/monosyllabic_adjectives.csv'
with open(mono_csv,'r',encoding = 'utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        # Assume the CSV has a column named 'headword' or similar
        if 'headword' in row:
            monosyllabic_set.add(row['headword'].lower())
        elif 'Headword' in row:
            monosyllabic_set.add(row['Headword'].lower())

# Utilities
def vowel_cluster_count(w):
    """Count vowel clusters to determine syllables"""
    return len(re.findall(r'[^aeiouy]*[aeiouy]+(?:[^aeiouy]+(?=[^aeiouy]*[aeiouy])|[^aeiouy]*)?', w)) or 1

def is_elided(word, next_word):
    """Check if word is elided before next word"""
    if next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS:
        return True
    return False

def is_weak_form(prev_tag, prev_word, tag, next_tag):
    """Check if adjective is in weak declension context"""
    weak_triggers = ['demonstrative', 'def_art', 'n%gen', 'pron%gen', 'interj', 'pron%fem_gen']
    for cause in weak_triggers:
        if prev_tag == cause:
            return True
    if next_tag == 'n#propn':
        return True
    return False

def is_plural_form(prev_tag, tag, next_tag):
    """Check if adjective modifies plural noun"""
    if next_tag.startswith('n%pl') or prev_tag.startswith('n%pl'):
        return True
    return False

def parse_tagged_text(oxford_text, oxford_tagging):
    """Extract words from oxford_text and tags from oxford_tagging"""
    if pd.isna(oxford_text) or pd.isna(oxford_tagging) or oxford_text == '' or oxford_tagging == '':
        return [], [], []

    # Clean and split the text
    words = oxford_text.lower().replace(',', '').replace('.', '').replace('!', '').replace('?', '').replace('°', '').replace('¶', '').strip().split()

    # Parse tags from oxford_tagging
    tag_tokens = oxford_tagging.strip().split()
    tags = []
    headwords = []

    for token in tag_tokens:
        if '@' in token:
            if token not in ["--@dash", ".@ellipsis"]:
                parts = token.split('@')
                headword = parts[0].lower()
                tag_part = parts[1] if len(parts) > 1 else ''

                tag = tag_part
                tag = ''.join([i for i in tag if not i.isdigit()])
                headwords.append(headword)
                tags.append(tag)

    # Handle special demonstrative cases
    for i, (word, tag) in enumerate(zip(words[:len(tags)], tags)):
        if word in ['this', 'that', 'thilke'] and 'gram_adj' in tag:
            tags[i] = 'demonstrative'

    # Ensure all lists are the same length (trim to shortest)
    if len(words) != len(tags):
        print(f"Length mismatch: {oxford_text} | {oxford_tagging}")
        print(f"Words: {len(words)}, Tags: {len(tags)}")
    min_len = min(len(words), len(tags), len(headwords))
    return words[:min_len], headwords[:min_len], tags[:min_len]

def is_monosyllabic_root(headword, word_forms):
    if headword in monosyllabic_set:
        return True
    else:
        return False

def add_exception_to_doc(exception_type, word, oxford_text, line_number, filename):
    """Add an exception case to the Word document"""
    doc_para = doc.add_paragraph()
    doc_para.add_run(f"{exception_type}: {word}").bold = True
    doc_para.add_run(f"\nLine {line_number} ({filename})\n")
    
    for i, ox_word in enumerate(oxford_text.split(' ')):
        if ox_word.lower() == word.lower():
            doc_para.add_run(f"{ox_word}").italic = True
        else:
            doc_para.add_run(f"{ox_word}")
        if i < len(oxford_text.split(' ')) - 1:
            doc_para.add_run(" ")
    doc_para.add_run("\n\n")

def analyze_adjectives(df, results):
    """Analyze adjective patterns in CSV data"""
    
    for idx, row in df.iterrows():
        if row["MATCH"] != "DIFF":
            oxford_text = row['OXFORD_TEXT']
            oxford_tagging = row['OXFORD_TAGGING']
            line_number = row['LINE_NUMBER']
            filename = row['OXFORD_FILENAME']
            original_text = row.get('OG_OXFORD_TEXT', oxford_text)

            words, headwords, tags = parse_tagged_text(oxford_text, oxford_tagging)

            prev_tag = 'NA'
            prev_word = 'NA'
            
            for j in range(len(tags)):
                if j >= len(words) or j >= len(headwords):
                    continue
                    
                headword = headwords[j]
                tag = tags[j]
                word = words[j].lower()
                next_tag = tags[j + 1] if j + 1 < len(tags) else 'END'
                next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'

                # Only process adjectives
                if not tag.startswith('adj'):
                    prev_word = word
                    prev_tag = tag
                    continue

                # Initialize adjective tracking
                if headword not in results['all_adjectives']:
                    results['all_adjectives'][headword] = {
                        'stem_forms': set(), 
                        'all_forms': set(),
                        'stem_with_e': set(),
                        'stem_without_e': set()
                    }

                results['all_adjectives'][headword]['all_forms'].add(word)
                if tag == 'adj':  # Stem form
                    results['all_adjectives'][headword]['stem_forms'].add(word)
                    if word.endswith('e'):
                        results['all_adjectives'][headword]['stem_with_e'].add(word)
                    else:
                        results['all_adjectives'][headword]['stem_without_e'].add(word)

                # Check if monosyllabic root
                stem_forms = results['all_adjectives'][headword]['stem_forms']
                if not is_monosyllabic_root(headword, stem_forms):
                    prev_word = word
                    prev_tag = tag
                    continue

                # Determine grammatical context
                is_weak = is_weak_form(prev_tag, prev_word, tag, next_tag)
                is_plural = is_plural_form(prev_tag, tag, next_tag)
                is_final = (next_word == 'END')
                is_word_elided = is_elided(word, next_word) if next_word != 'END' else False

                # Create record for this instance
                record = {
                    'headword': headword,
                    'word': word,
                    'line_number': line_number,
                    'filename': filename,
                    'context': original_text,
                    'is_elided': is_word_elided,
                    'is_final': is_final
                }

                # Categorize the adjective usage
                if is_weak:
                    if not word.endswith('e'):
                        # Weak declension without -e (exception)
                        results['weak_no_e_all'].append(record)
                        if not is_word_elided and not is_final:
                            results['weak_no_e_strict'].append(record)
                            add_exception_to_doc("Weak without -e", word, original_text, line_number, filename)
                elif is_plural:
                    if not word.endswith('e'):
                        # Plural form without -e (exception)
                        results['plural_no_e_all'].append(record)
                        if not is_word_elided and not is_final:
                            results['plural_no_e_strict'].append(record)
                            add_exception_to_doc("Plural without -e", word, original_text, line_number, filename)
                else:
                    # Strong/stem form - but exclude adjectives that always end in -e in strong form
                    if word.endswith('e') and tag == 'adj':
                        # Only count as exception if this adjective also has stem forms without -e
                        # We'll filter this in post-processing after collecting all data
                        results['strong_with_e_all'].append(record)
                        if not is_word_elided and not is_final:
                            results['strong_with_e_strict'].append(record)

                prev_word = word
                prev_tag = tag

    return results

def process_csv_directory(csv_dir):
    """Process all _gui.csv files in the directory"""
    results = {
        'all_adjectives': {},
        'weak_no_e_all': [],          # Weak declension without -e (all instances)
        'weak_no_e_strict': [],       # Weak declension without -e (strict: not elided, not final)
        'plural_no_e_all': [],        # Plural form without -e (all instances)  
        'plural_no_e_strict': [],     # Plural form without -e (strict: not elided, not final)
        'strong_with_e_all': [],      # Strong form with -e (all instances)
        'strong_with_e_strict': []    # Strong form with -e (strict: not elided, not final)
    }

    file_count = 0
    for root, dirs, files in os.walk(csv_dir):
        for file in files:
            if not file.endswith('_gui_complete.csv'):
                continue

            csv_path = os.path.join(root, file)
            file_count += 1

            try:
                # Read CSV file
                df = pd.read_csv(csv_path, encoding='utf-8')

                # Skip if required columns are missing
                required_columns = ['OXFORD_TAGGING', 'OXFORD_TEXT', 'LINE_NUMBER', 'OXFORD_FILENAME']
                if not all(col in df.columns for col in required_columns):
                    print(f"Skipping {file}: missing required columns")
                    continue

                results = analyze_adjectives(df, results)

            except Exception as e:
                print(f"Error processing {file}: {e}")
                continue

    print(f"Processed {file_count} files")
    return results

def filter_strong_form_exceptions(results):
    """Filter out adjectives that always end in -e in strong form from strong form exceptions"""
    # Identify adjectives that always end in -e in strong form
    always_e_adjectives = set()
    for headword, data in results['all_adjectives'].items():
        if is_monosyllabic_root(headword, data['stem_forms']):
            # If this adjective has stem forms and ALL stem forms end in -e, it's an "always -e" adjective
            if data['stem_forms'] and len(data['stem_without_e']) == 0:
                always_e_adjectives.add(headword)
    
    # Filter the strong form exceptions to exclude always-e adjectives
    results['strong_with_e_all_filtered'] = [
        record for record in results['strong_with_e_all'] 
        if record['headword'] not in always_e_adjectives
    ]
    results['strong_with_e_strict_filtered'] = [
        record for record in results['strong_with_e_strict'] 
        if record['headword'] not in always_e_adjectives
    ]
    
    # Add the always-e adjectives to results for reporting
    results['always_e_adjectives'] = always_e_adjectives
    
    return results

def write_results_to_csv(results, output_dir):
    """Write analysis results to CSV files"""
    os.makedirs(output_dir, exist_ok=True)
    
    # Filter strong form exceptions first
    results = filter_strong_form_exceptions(results)
    
    # Helper function to write a list of records to CSV
    def write_records(records, filename):
        if not records:
            # Create empty file with headers
            with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
            return
            
        with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['headword', 'word', 'line_number', 'filename', 'context', 'is_elided', 'is_final']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            for record in records:
                writer.writerow(record)

    # Write each category
    write_records(results['weak_no_e_all'], 'weak_no_e_all_instances.csv')
    write_records(results['weak_no_e_strict'], 'weak_no_e_strict_instances.csv')
    write_records(results['plural_no_e_all'], 'plural_no_e_all_instances.csv')
    write_records(results['plural_no_e_strict'], 'plural_no_e_strict_instances.csv')
    write_records(results['strong_with_e_all'], 'strong_with_e_all_instances.csv')
    write_records(results['strong_with_e_strict'], 'strong_with_e_strict_instances.csv')
    write_records(results['strong_with_e_all_filtered'], 'strong_with_e_all_filtered_instances.csv')
    write_records(results['strong_with_e_strict_filtered'], 'strong_with_e_strict_filtered_instances.csv')

    # Write summary statistics
    with open(os.path.join(output_dir, 'summary_statistics.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Category', 'All Instances', 'Strict Instances (not elided, not final)'])
        writer.writerow(['Weak declension without -e', len(results['weak_no_e_all']), len(results['weak_no_e_strict'])])
        writer.writerow(['Plural form without -e', len(results['plural_no_e_all']), len(results['plural_no_e_strict'])])
        writer.writerow(['Strong form with -e (all)', len(results['strong_with_e_all']), len(results['strong_with_e_strict'])])
        writer.writerow(['Strong form with -e (filtered)', len(results['strong_with_e_all_filtered']), len(results['strong_with_e_strict_filtered'])])
        writer.writerow(['Always -e adjectives found', len(results['always_e_adjectives']), 'N/A'])

    # Write monosyllabic adjectives found
    with open(os.path.join(output_dir, 'monosyllabic_adjectives.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Headword', 'Stem Forms', 'All Forms', 'Always Ends in E'])
        for headword, data in results['all_adjectives'].items():
            if is_monosyllabic_root(headword, data['stem_forms']):
                always_e = 'Yes' if headword in results['always_e_adjectives'] else 'No'
                writer.writerow([headword, '; '.join(sorted(data['stem_forms'])), '; '.join(sorted(data['all_forms'])), always_e])

    # Write list of always -e adjectives
    with open(os.path.join(output_dir, 'always_e_adjectives.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Headword', 'Stem Forms with -e', 'Reason'])
        for headword in sorted(results['always_e_adjectives']):
            data = results['all_adjectives'][headword]
            stem_forms = '; '.join(sorted(data['stem_with_e']))
            writer.writerow([headword, stem_forms, 'All stem forms end in -e'])

def write_summary_to_docx(results, output_dir):
    """Write summary statistics to Word document"""
    # Filter strong form exceptions first
    results = filter_strong_form_exceptions(results)
    
    summary_doc = docx.Document()
    summary_doc.add_heading('Adjective Declension Analysis Summary', 0)
    
    monosyllabic_count = sum(1 for headword, data in results['all_adjectives'].items() 
                           if is_monosyllabic_root(headword, data['stem_forms']))
    
    # Overview section
    summary_doc.add_heading('Overview', 1)
    para = summary_doc.add_paragraph()
    para.add_run(f"Total adjectives found: {len(results['all_adjectives'])}\n")
    para.add_run(f"Monosyllabic adjectives found: {monosyllabic_count}\n")
    para.add_run(f"Always -e adjectives found: {len(results['always_e_adjectives'])}")
    
    # Statistics section
    summary_doc.add_heading('Exception Statistics', 1)
    
    categories = [
        ('Weak declension without -e', 'weak_no_e_all', 'weak_no_e_strict'),
        ('Plural form without -e', 'plural_no_e_all', 'plural_no_e_strict'),
        ('Strong form with -e (all)', 'strong_with_e_all', 'strong_with_e_strict'),
        ('Strong form with -e (filtered)', 'strong_with_e_all_filtered', 'strong_with_e_strict_filtered')
    ]
    
    for desc, all_key, strict_key in categories:
        all_count = len(results[all_key])
        strict_count = len(results[strict_key])
        
        para = summary_doc.add_paragraph()
        para.add_run(desc + ":").bold = True
        para.add_run(f"\n  All instances: {all_count}")
        para.add_run(f"\n  Strict instances (not elided, not final): {strict_count}\n")
    
    # Always -e adjectives section
    if results['always_e_adjectives']:
        summary_doc.add_heading('Adjectives That Always End in -e in Strong Form', 1)
        para = summary_doc.add_paragraph()
        para.add_run("The following adjectives were excluded from strong form exception analysis because they always end in -e in the strong form:\n\n")
        
        for headword in sorted(results['always_e_adjectives']):
            data = results['all_adjectives'][headword]
            stem_forms = ', '.join(sorted(data['stem_with_e']))
            para.add_run(f"• {headword}").bold = True
            para.add_run(f" (forms: {stem_forms})\n")
    
    summary_doc.save(os.path.join(output_dir, 'analysis_summary.docx'))

# Main execution
if __name__ == "__main__":
    # Process all CSV files
    results = process_csv_directory(base_csv_dir)
    
    # Write results
    output_dir = 'adjective_analysis_output'
    write_results_to_csv(results, output_dir)
    write_summary_to_docx(results, output_dir)
    
    # Save Word document with exceptions
    doc.save(os.path.join(output_dir, 'declension_exceptions.docx'))
    
    # Write completion log to file
    with open(os.path.join(output_dir, 'analysis_log.txt'), 'w', encoding='utf-8') as log_file:
        log_file.write("Chaucer Adjective Declension Analysis - Complete\n")
        log_file.write("="*50 + "\n\n")
        
        monosyllabic_count = sum(1 for headword, data in results['all_adjectives'].items() 
                               if is_monosyllabic_root(headword, data['stem_forms']))
        
        # Filter results for logging
        results = filter_strong_form_exceptions(results)
        
        log_file.write(f"Total adjectives found: {len(results['all_adjectives'])}\n")
        log_file.write(f"Monosyllabic adjectives found: {monosyllabic_count}\n")
        log_file.write(f"Always -e adjectives found: {len(results['always_e_adjectives'])}\n\n")
        
        log_file.write("Files generated:\n")
        files = [
            "weak_no_e_all_instances.csv - All weak declension forms without -e",
            "weak_no_e_strict_instances.csv - Strict weak declension exceptions", 
            "plural_no_e_all_instances.csv - All plural forms without -e",
            "plural_no_e_strict_instances.csv - Strict plural exceptions",
            "strong_with_e_all_instances.csv - All strong forms with -e",
            "strong_with_e_strict_instances.csv - Strict strong form exceptions",
            "strong_with_e_all_filtered_instances.csv - Strong forms with -e (filtered)",
            "strong_with_e_strict_filtered_instances.csv - Strict strong exceptions (filtered)",
            "summary_statistics.csv - Overall counts",
            "monosyllabic_adjectives.csv - List of all monosyllabic adjectives found",
            "always_e_adjectives.csv - List of adjectives that always end in -e",
            "analysis_summary.docx - Summary statistics document",
            "declension_exceptions.docx - Detailed exception examples",
            "analysis_log.txt - This log file"
        ]
        
        for file_desc in files:
            log_file.write(f"- {file_desc}\n")
