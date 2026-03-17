# Re-import after code state reset
import os
import re
import string
import csv
from collections import defaultdict, Counter
import matplotlib.pyplot as plt
import docx

doc = docx.Document()
doc.add_heading('Exceptions to Rule 2 in the Oxford Chaucer', 0)

# Configuration
base_cat_dir = 'data/oxford_cats'
base_txt_dir = 'data/oxford_txts'
tag_tuple_list = [
    ('demonstrative', 'adj', 1),
    ('def_art', 'adj', 1),
    ('n%gen', 'adj', 1),
    ('pron%gen', 'adj', 1),
    ('pron%femgen', 'adj', 1),
    ('adj', 'n#propn', 0),
    ('interj', 'adj', 1),
]

# Utilities
def vowel_cluster_count(w):
#    w = w.replace('ou', 'ow').replace('ue', 'uwe').replace('oia', 'oja').replace('troian', 'trojan')
    return len(re.findall(r'[^aeiouy]*[aeiouy]+(?:[^aeiouy]+(?=[^aeiouy]*[aeiouy])|[^aeiouy]*)?', w)) or 1

def strip_annotations(line,ital_index):
	newline = ''
	first_real = False
	words = line.split(' ')
	doc_para = doc.add_paragraph()

	for i, word in enumerate(words):
		if i!=ital_index:
			doc_para.add_run(word)
		else:
			doc_para.add_run(word).italic = True
		newline += word
		if word.isdigit():
			newline+= '\n'
			doc_para.add_run('\n')
		elif i<len(words)-1:
			doc_para.add_run(' ')
			newline += ' '
	return newline.strip()

def total_clean(line):
    newline = ''

    words = line.replace('{',' ').split(' ')
    for word in words:
        if '}' not in word:
            newline += word + ' '
    return newline.strip()

def search_tag_sequence(cat_path, txt_path, tag1, tag2, pos, adjectives, unspelled_log, unspelled_elided_log):
    spelled, spelled_elided, unspelled, unspelled_elided, final = 0, 0, 0, 0, 0

    unspelled_lines = []
    unspelled_elided_lines = []

    with open(cat_path, 'r', encoding='utf-8') as f_cat, open(txt_path, 'r', encoding='utf-8') as f_txt:
        og_cat_lines = f_cat.readlines()
        og_txt_lines = f_txt.readlines()

    cat_lines = []
    txt_lines = []

    for txt_line in og_txt_lines:
        if txt_line.strip != '':
            txt_lines.append(txt_line)

    for cat_line in og_cat_lines:
        numbering = re.sub(r'\b(\w*\d)\b\s+', r'\1|||', cat_line).split('|||')
        if len(numbering)>1:
            cat_lines.append(cat_line)


    for line, ogline in zip(cat_lines, txt_lines):
        numbering = re.sub(r'\b(\w*\d)\b\s+', r'\1|||', line).split('|||')[0]
        print(numbering)
        line = line.lower()
        line = ''.join(ch for ch in line if ch in string.ascii_lowercase + ' \n_@%*#{}')
        tokens = line.strip().split(' ')
        words, headwords, tags = [], [], []

        for token in tokens:
            if '{' in token:
                word = token.split('{')[0]
                headword = token.split('*')[1].split('@')[0]
                tag = token.split('*')[1].split('@')[1].strip('_')
                if word in ['this', 'that', 'thilke'] and tag == 'gram_adj':
                    tag = 'demonstrative'
                words.append(word)
                headwords.append(headword)
                tags.append(tag)

        for j in range(len(tags) - 1):
            if tags[j] == tag1 and tags[j + 1] == tag2:
                word = words[j + pos]
                headword = headwords[j + pos]
                if vowel_cluster_count(headword) == 1 and word not in ['chief','coward','lewed','payen','real','royal','roial','shrewed','troian','cruel','crueel','crewel','cruele']:
                    adjectives[headword][5].append(word)
                    adjectives[headword][5] = list(set(adjectives[headword][5]))
                    if j + pos == len(tags) - 1:
                        final += 1
                        adjectives[headword][4] += 1
                    else:
                        nextword = words[j + pos + 1]
                        if word.endswith('e') and nextword[0] in 'aeiouh':
                            spelled_elided += 1
                            adjectives[headword][1] += 1
                        elif word.endswith('e'):
                            spelled += 1
                            adjectives[headword][0] += 1
                        elif nextword[0] in 'aeiouh':
                            unspelled_elided += 1
                            adjectives[headword][3] += 1
                            unspelled_elided_lines.append(tag1 + '->' + tag2 + '\n' + word + '\n' + line + numbering + ' ' + ogline + '\n\n')
                        else:
                            unspelled += 1
                            adjectives[headword][2] += 1
                            unspelled_lines.append(tag1 + '->' + tag2 + '\n' + word + '\n' + line + numbering + ' ' + ogline + '\n\n')

    unspelled_log.extend(unspelled_lines)
    unspelled_elided_log.extend(unspelled_elided_lines)
    #unspelled_elided_log.extend(unspelled_lines)
    return spelled, spelled_elided, unspelled, unspelled_elided, final

def process_directory(cat_dir, txt_dir):
    data = {}
    unspelled_log = []
    unspelled_elided_log = []
    adjectives = defaultdict(lambda: [0, 0, 0, 0, 0, []])  # spelled, spelled_elided, unspelled, unspelled_elided, final, words

    for root, _, files in os.walk(cat_dir):
        for file in files:
            if not file.endswith('.cat'):
                continue
            rel_path = os.path.relpath(root, cat_dir)
            file_id = os.path.join(rel_path, file) if rel_path != '.' else file
            cat_path = os.path.join(root, file)
            txt_path = os.path.join(txt_dir, rel_path, file.replace('.cat', '.txt'))

            if not os.path.exists(txt_path):
                continue

            tag_data = {}
            total_spelled, total_spelled_elided, total_unspelled, total_unspelled_elided, total_final = 0, 0, 0, 0, 0
            for tag1, tag2, pos in tag_tuple_list:
                s, s_e, u, u_e, f = search_tag_sequence(cat_path, txt_path, tag1, tag2, pos, adjectives, unspelled_log, unspelled_elided_log)
                tag_data[f'{tag1}->{tag2} spelled'] = s
                tag_data[f'{tag1}->{tag2} spelled, elided'] = s_e
                tag_data[f'{tag1}->{tag2} unspelled'] = u
                tag_data[f'{tag1}->{tag2} unspelled, elided'] = u_e
                tag_data[f'{tag1}->{tag2} %'] = round((s / (s + u) * 100), 2) if (s + u) > 0 else 0.0
                total_spelled += s
                total_spelled_elided += s_e
                total_unspelled += u
                total_unspelled_elided += u_e

            match_ratio = total_spelled / (total_spelled + total_unspelled) if (total_spelled + total_unspelled) > 0 else 0
            data[file_id] = {
                'overall spelled unelided': total_spelled,
                'overall unspelled unelided': total_unspelled,
                'overall spelled elided': total_spelled_elided,
                'overall unspelled elided': total_unspelled_elided,
                'overall final': total_final,
                'overall unelided spelled/unspelled ratio': round(match_ratio * 100, 2),
                **tag_data
            }

    return data, adjectives, unspelled_log, unspelled_elided_log

def write_csv(data, filename):
    keys = sorted(next(iter(data.values())).keys())
    with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=['file'] + keys)
        writer.writeheader()
        for file, stats in data.items():
            row = {'file': file}
            row.update(stats)
            writer.writerow(row)

def write_adjectives(adjectives, filename):
    with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['headword','weak_forms', 'spelled', 'spelled_elided', 'unspelled', 'unspelled_elided', 'final','total','unelided_pct','overall_pct'])

        for adj in sorted(adjectives.keys()):
            s, s_e, u, u_e, f, forms = adjectives[adj]
            total = s + s_e + u + u_e + f
            match_pct_unelided = round((s / (s + u) * 100), 2) if (s + u) > 0 else 'NA'
            match_pct_overall = round(((s+s_e) / (s + s_e  + u + u_e) * 100), 2) if (s + s_e + u + u_e) > 0 else 'NA'
            writer.writerow([adj, forms, s, s_e, u, u_e, f, total, match_pct_unelided, match_pct_overall])

def write_exceptions(log, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        f.writelines(log)

'''
def make_plots(data, output_prefix='plot'):
    levels = defaultdict(lambda: {'matches': 0, 'exceptions': 0})
    for file, stats in data.items():
        level = file.count(os.sep)
        levels[level]['matches'] += stats['overall matches']
        levels[level]['exceptions'] += stats['overall exceptions']

    for level, values in levels.items():
        total = values['matches'] + values['exceptions']
        if total == 0:
            continue
        ratio = values['matches'] / total
        plt.figure()
        plt.bar(['matches', 'exceptions'], [values['matches'], values['exceptions']], color=['green', 'red'])
        plt.title(f'Level {level} (Match Ratio: {ratio:.2%})')
        plt.ylabel('Count')
        plt.savefig(f'{output_prefix}_level{level}.png')
        plt.close()
'''

# Main execution
data, adjectives, unspelled_log, unspelled_elided_log = process_directory(base_cat_dir, base_txt_dir)
write_csv(data, 'rules_output/tag_analysis.csv')
write_adjectives(adjectives, 'rules_output/adjectives.csv')
write_exceptions(unspelled_log, 'rules_output/unspelled_oxford.txt')
doc.save('rules_output/exceptions.docx')
