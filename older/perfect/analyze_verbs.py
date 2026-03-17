import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict, Counter

# Initialize Word documents for exceptions
oxford_doc = docx.Document()
oxford_doc.add_heading('Oxford Verb Declension Exceptions', 0)

riverside_doc = docx.Document()
riverside_doc.add_heading('Riverside Verb Declension Exceptions', 0)

# Configuration
base_csv_dir = 'data/csvs'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
					"hadden", "his", "her", "him", "hers", "hide", "hir",
					"hire", "hires", "hirs", "han"]

form_csv = 'verb_forms_gold.csv'
verbs_dict = {}
with open(form_csv, 'r', encoding='utf-8') as csvfile:
	reader = csv.DictReader(csvfile)
	for row in reader:
		if 'headword' in row and 'classification' in row:
			verbs_dict[row['headword'].lower()] = row['classification'].lower()

def is_strong(verb):
	return verbs_dict.get(verb, '') == 'strong'

def is_weak(verb):
	return verbs_dict.get(verb, '') == 'weak'

def is_elided(word, next_word):
	"""Check if word is elided before next word"""
	if next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS:
		return True
	return False

def clean_tag(tag):
	"""Remove digits before % in a tag (e.g., v2%pt_1 -> v%pt_1)."""
	return re.sub(r'\d+(?=%)', '', tag)

def parse_tagged_text(text, tagging, text_type):
	"""Extract words from text and tags from tagging"""
	if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
		return [], [], []

	words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()
	tag_tokens = tagging.strip().split()
	tags = []
	headwords = []

	for token in tag_tokens:
		if '@' in token:
			parts = token.split('@')
			headword = parts[0].lower()
			tag = parts[1] if len(parts) > 1 else ''
			tag = ''.join([i for i in tag if not i.isdigit()])
			headwords.append(headword)
			tags.append(tag)

	min_len = min(len(words), len(tags), len(headwords))
	return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(doc, exception_type, word, text, line_number, filename):
	"""Add an exception case to the Word document"""
	doc_para = doc.add_paragraph()
	doc_para.add_run(f"{exception_type}: {word}").bold = True
	doc_para.add_run(f"\n{line_number} ({filename})\n")
	for i, ox_word in enumerate(text.split(' ')):
		if ox_word.lower() == word.lower():
			doc_para.add_run(f"{ox_word}").italic = True
		else:
			doc_para.add_run(f"{ox_word}")
		if i < len(text.split(' ')) - 1:
			doc_para.add_run(" ")
	doc_para.add_run("\n\n")

def classify_ending(word):
	"""Classify the word ending into categories."""
	if word.endswith('eth'):
		return '-eth'
	if word.endswith('en'):
		return '-en'
	if word.endswith('e'):
		return '-e'
	if word.endswith('ed'):
		return '-ed'
	if word.endswith('d'):
		return '-d'
	if word.endswith('t'):
		return '-t'
	if word.endswith(('a','e','i','o','u')):
		return 'vowel'
	return 'other'

def analyze_verbs(df, results, text_type, doc):
	"""Analyze verb patterns in CSV data according to rules."""
	text_col = f'{text_type}_TEXT'
	tagging_col = f'{text_type}_TAGGING'
	filename_col = f'{text_type}_FILENAME'
	original_text_col = f'OG_{text_type}_TEXT'

	for idx, row in df.iterrows():
		if row.get("MATCH") == "DIFF":
			continue

		text = row[text_col]
		tagging = row[tagging_col]
		line_number = row['LINE_NUMBER']
		filename = row[filename_col]
		original_text = row.get(original_text_col, text)

		words, headwords, tags = parse_tagged_text(text, tagging, text_type)

		for j in range(len(tags)):
			if j >= len(words) or j >= len(headwords):
				continue

			word = words[j].lower()
			headword = headwords[j]
			tag = clean_tag(tags[j])
			next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'

			# Skip if elided
			if is_elided(word, next_word):
				continue

			# Skip if not strong/weak
			#if headword not in verbs_dict:
				#continue
			#if not (is_strong(headword) or is_weak(headword)): 
				#continue
			#if word in ['han']:
			#	continue

			ending = classify_ending(word)

			# Count ending distribution
			results['ending_counts'][tag][ending] += 1
			results['verb_tag_counts'][headword][tag][ending] += 1

			# Rule checks
			violated = False
			reason = ""

			if tag == 'v%inf' and ending not in ['-en', '-e']:
				violated = True
				reason +=  "Infinitive must end in -en or -e "
			if tag == 'v%pt_pl' and ending not in ['-en', '-e']:
				violated = True
				reason += "Past plural must end in -en or -e "
			if tag == 'v%pr_3' and ending != '-eth':
				violated = True
				reason += "Present 3rd sg must end in -eth "
			if tag == 'v%pr_pl' and ending not in ['-en', '-e']:
				violated = True
				reason += "Present plural must end in -en or -e "
			if headword in verbs_dict:
					if is_strong(headword) and tag in ['v%pt_1', 'v%pt_3'] and ending in ['-en', '-e']:
						violated = True
						reason += "Strong pt sg must not end in -en or -e "
					if is_weak(headword) and tag in ['v%pt_1', 'v%pt_3'] and ending not in ['-ed','-d','-t']:
						violated = True
						reason += "Weak pt sg must end in -ed, -d, or -t "
					if is_strong(headword) and tag == 'v%ppl' and ending not in ['-en', '-e']:
						violated = True
						reason += "Strong participle must end in -en or -e "

			if violated:
				record = {
					'headword': headword,
					'word': word,
					'tag': tag,
					'line_number': line_number,
					'filename': filename,
					'context': original_text,
					'text_type': text_type
				}
				results['exceptions'].append(record)
				add_exception_to_doc(doc, reason, word, original_text, line_number, filename)

	return results

def process_csv_directory(csv_dir, text_type, doc):
	results = {
		'exceptions': [],
		'ending_counts': defaultdict(Counter),
		'verb_tag_counts': defaultdict(lambda: defaultdict(Counter))
	}

	file_count = 0
	for root, dirs, files in os.walk(csv_dir):
		for file in files:
			if not file.endswith('_gui.csv'):
				continue
			file_count += 1
			try:
				df = pd.read_csv(os.path.join(root, file), encoding='utf-8')
				required_columns = [f'{text_type}_TAGGING', f'{text_type}_TEXT', 'LINE_NUMBER', f'{text_type}_FILENAME']
				if not all(col in df.columns for col in required_columns):
					print(f"Skipping {file} for {text_type}: missing columns")
					continue
				results = analyze_verbs(df, results, text_type, doc)
			except Exception as e:
				print(f"Error processing {file} for {text_type}: {e}")
				continue
	print(f"Processed {file_count} files for {text_type}")
	return results

def write_results(results, output_dir, text_type):
	os.makedirs(output_dir, exist_ok=True)

	# Exceptions
	with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		fieldnames = ['headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type']
		writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
		writer.writeheader()
		for rec in results['exceptions']:
			writer.writerow(rec)

	# Distribution by tag
	with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Tag', 'Ending', 'Count', 'Percent'])
		for tag, counts in results['ending_counts'].items():
			total = sum(counts.values())
			for ending, count in counts.items():
				percent = (count / total) * 100 if total > 0 else 0
				writer.writerow([tag, ending, count, f"{percent:.2f}%"])

	# Distribution by verb+tag
	with open(os.path.join(output_dir, f'{text_type.lower()}_verb_tag_distribution.csv'), 'w', newline='', encoding='utf-8') as csvfile:
		writer = csv.writer(csvfile)
		writer.writerow(['Headword', 'Tag', 'Ending', 'Count', 'Percent'])
		for verb, tags in results['verb_tag_counts'].items():
			for tag, counts in tags.items():
				total = sum(counts.values())
				for ending, count in counts.items():
					percent = (count / total) * 100 if total > 0 else 0
					writer.writerow([verb, tag, ending, count, f"{percent:.2f}%"])

def write_summary_docx(results, output_dir, text_type):
	summary_doc = docx.Document()
	summary_doc.add_heading(f'{text_type} Verb Declension Analysis Summary', 0)

	summary_doc.add_heading('Exceptions', 1)
	summary_doc.add_paragraph(f"Total exceptions found: {len(results['exceptions'])}")

	summary_doc.add_heading('Ending Distribution by Tag', 1)
	for tag, counts in results['ending_counts'].items():
		para = summary_doc.add_paragraph()
		para.add_run(f"{tag}: ").bold = True
		total = sum(counts.values())
		for ending, count in counts.items():
			percent = (count / total) * 100 if total > 0 else 0
			para.add_run(f"{ending}={count} ({percent:.1f}%)  ")

	summary_doc.save(os.path.join(output_dir, f'{text_type.lower()}_analysis_summary.docx'))

# Main execution
if __name__ == "__main__":
	print("Processing Oxford texts...")
	oxford_results = process_csv_directory(base_csv_dir, 'OXFORD', oxford_doc)

	print("Processing Riverside texts...")
	riverside_results = process_csv_directory(base_csv_dir, 'RIVERSIDE', riverside_doc)

	oxford_output_dir = 'oxford_verb_analysis_output'
	riverside_output_dir = 'riverside_verb_analysis_output'

	write_results(oxford_results, oxford_output_dir, 'OXFORD')
	write_summary_docx(oxford_results, oxford_output_dir, 'OXFORD')
	oxford_doc.save(os.path.join(oxford_output_dir, 'oxford_declension_exceptions.docx'))

	write_results(riverside_results, riverside_output_dir, 'RIVERSIDE')
	write_summary_docx(riverside_results, riverside_output_dir, 'RIVERSIDE')
	riverside_doc.save(os.path.join(riverside_output_dir, 'riverside_declension_exceptions.docx'))

	print("\nAnalysis complete!")
	print(f"Oxford results: {oxford_output_dir}/")
	print(f"Riverside results: {riverside_output_dir}/")

