import docx
import os
import re
import csv
import pandas as pd
from collections import defaultdict, Counter

# Initialize Word documents for exceptions
oxford_doc = docx.Document()
oxford_doc.add_heading('Oxford Verb Declension Exceptions', 0)

riverside_doc = docx.Document()
riverside_doc.add_heading('Riverside Verb Declension Exceptions', 0)

# Configuration
base_csv_dir = 'data/csvs'

ELISION_FOLLOWERS = ["have", "haven", "haveth", "havest", "had", "hadde",
                    "hadden", "his", "her", "him", "hers", "hide", "hir",
                    "hire", "hires", "hirs", "han"]

for follower in ELISION_FOLLOWERS:
    if 'i' in follower:
        ELISION_FOLLOWERS.append(follower.replace('i','y'))

form_csv = 'verb_forms_gold.csv'
verbs_dict = {}
with open(form_csv, 'r', encoding='utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        if 'headword' in row and 'classification' in row:
            verbs_dict[row['headword'].lower()] = row['classification'].lower()

def is_strong(verb):
    return verbs_dict.get(verb, '') == 'strong'

def is_weak(verb):
    return verbs_dict.get(verb, '') == 'weak'

def is_elided(word, next_word):
    if next_word[0] in 'aeiou' or next_word in ELISION_FOLLOWERS:
        return True
    return False

def clean_tag(tag):
    return re.sub(r'\d+(?=%)', '', tag)

def parse_tagged_text(text, tagging, text_type):
    if pd.isna(text) or pd.isna(tagging) or text == '' or tagging == '':
        return [], [], []

    words = re.sub(r'[.,!?°¶]', '', text.lower()).strip().split()
    tag_tokens = tagging.strip().split()
    tags = []
    headwords = []

    for token in tag_tokens:
        if '@' in token:
            parts = token.split('@')
            headword = parts[0].lower()
            tag = parts[1] if len(parts) > 1 else ''
            headwords.append(headword)
            tags.append(tag)

    min_len = min(len(words), len(tags), len(headwords))
    return words[:min_len], headwords[:min_len], tags[:min_len]

def add_exception_to_doc(doc, exception_type, word, text, line_number, filename):
    doc_para = doc.add_paragraph()
    doc_para.add_run(f"{exception_type}: {word}").bold = True
    doc_para.add_run(f"\n{line_number} ({filename})\n")
    for i, ox_word in enumerate(text.split(' ')):
        if ox_word.lower() == word.lower():
            doc_para.add_run(f"{ox_word}").italic = True
        else:
            doc_para.add_run(f"{ox_word}")
        if i < len(text.split(' ')) - 1:
            doc_para.add_run(" ")
    doc_para.add_run("\n\n")

def classify_ending(word):
    if word.endswith('th'):
        return '-th'
    if word.endswith('en'):
        return '-en'
    if word.endswith('e'):
        return '-e'
    if word.endswith('ed'):
        return '-ed'
    if word.endswith('d'):
        return '-d'
    if word.endswith('t'):
        return '-t'
    if word.endswith(('a','e','i','o','u')):
        return 'vowel'
    return 'other'

def analyze_verbs(df, results, text_type, doc, excluded_words=None):
    """Analyze verb patterns in CSV data according to rules."""
    if excluded_words is None:
        excluded_words = {}

    text_col = f'{text_type}_TEXT'
    tagging_col = f'{text_type}_TAGGING'
    filename_col = f'{text_type}_FILENAME'
    original_text_col = f'OG_{text_type}_TEXT'

    # Store all lines for later comparison
    if 'all_lines' not in results:
        results['all_lines'] = []

    for idx, row in df.iterrows():
        if row.get("MATCH") == "DIFF":
            continue

        text = row[text_col]
        tagging = row[tagging_col]
        line_number = row['LINE_NUMBER']
        filename = row[filename_col]
        original_text = row.get(original_text_col, text)

        # Store all lines for comparison
        results['all_lines'].append({
            'line_number': line_number,
            'text': original_text,
            'text_type': text_type
        })

        words, headwords, tags = parse_tagged_text(text, tagging, text_type)

        for j in range(len(tags)):
            if j >= len(words) or j >= len(headwords):
                continue

            word = words[j].lower()
            headword = headwords[j]
            tag = clean_tag(tags[j])
            next_word = words[j + 1].lower() if j + 1 < len(words) else 'END'

            # Skip if elided
            if is_elided(word, next_word):
                continue

            ending = classify_ending(word)

            # Count ending distribution
            results['ending_counts'][tag][ending] += 1
            results['verb_tag_counts'][headword][tag][word] += 1

            # Track headword performance for each rule
            violated = False
            reason = ""

            # Track rules individually and per headword
            def record_rule(rule_name, condition, headword):
                # Skip if headword is excluded for this rule
                if headword in excluded_words.get(rule_name, set()):
                    return

                if condition:
                    results['rule_stats'][rule_name]["success"] += 1
                    results['headword_rule_stats'][rule_name][headword]["success"] += 1
                else:
                    results['rule_stats'][rule_name]["fail"] += 1
                    results['headword_rule_stats'][rule_name][headword]["fail"] += 1

            # Rules with exclusion checking
            if tag == 'v%inf':
                rule_name = "Infinitive ends in -en/-e"
                if headword not in excluded_words.get(rule_name, set()):
                    record_rule(rule_name, ending in ['-en', '-e'], headword)
                    if ending not in ['-en', '-e']:
                        violated = True
                        reason += "Infinitive must end in -en or -e "

            if tag == 'v%pt_pl':
                rule_name = "Past plural ends in -en/-e"
                if headword not in excluded_words.get(rule_name, set()):
                    record_rule(rule_name, ending in ['-en', '-e'], headword)
                    if ending not in ['-en', '-e']:
                        violated = True
                        reason += "Past plural must end in -en or -e "

            if tag == 'v%pr_3':
                rule_name = "Present 3rd sg ends in -th"
                if headword not in excluded_words.get(rule_name, set()):
                    record_rule(rule_name, ending == '-th', headword)
                    if ending != '-th':
                        violated = True
                        reason += "Present 3rd sg must end in -th "

            if tag == 'v%pr_pl':
                rule_name = "Present plural ends in -en/-e"
                if headword not in excluded_words.get(rule_name, set()):
                    record_rule(rule_name, ending in ['-en', '-e'], headword)
                    if ending not in ['-en', '-e']:
                        violated = True
                        reason += "Present plural must end in -en or -e "

            if headword in verbs_dict:
                if is_strong(headword) and tag in ['v%pt_1', 'v%pt_3']:
                    rule_name = "Strong pt sg not -en/-e"
                    if headword not in excluded_words.get(rule_name, set()):
                        record_rule(rule_name, ending not in ['-en', '-e'], headword)
                        if ending in ['-en', '-e']:
                            violated = True
                            reason += "Strong pt sg must not end in -en or -e "

                if is_weak(headword) and tag in ['v%pt_1', 'v%pt_3']:
                    rule_name = "Weak pt sg ends in -ed/-d/-t"
                    if headword not in excluded_words.get(rule_name, set()):
                        record_rule(rule_name, ending in ['-ed','-d','-t'], headword)
                        if ending not in ['-ed','-d','-t']:
                            violated = True
                            reason += "Weak pt sg must end in -ed, -d, or -t "

                if is_strong(headword) and tag == 'v%ppl':
                    rule_name = "Strong participle ends in -en/-e"
                    if headword not in excluded_words.get(rule_name, set()):
                        record_rule(rule_name, ending in ['-en','-e'], headword)
                        if ending not in ['-en', '-e']:
                            violated = True
                            reason += "Strong participle must end in -en or -e "

            if violated and reason.strip():
                record = {
                    'headword': headword,
                    'word': word,
                    'tag': tag,
                    'line_number': line_number,
                    'filename': filename,
                    'context': original_text,
                    'text_type': text_type,
                    'reason': reason.strip()
                }
                results['exceptions'].append(record)
                add_exception_to_doc(doc, reason, word, original_text, line_number, filename)

    return results

def identify_high_failure_headwords(results, failure_threshold=0.9):
    """Identify headwords that fail rules more than the threshold percentage"""
    high_failure_headwords = {}
    headword_stats = {}

    for rule_name, headword_stats_dict in results['headword_rule_stats'].items():
        high_failure_headwords[rule_name] = set()
        headword_stats[rule_name] = []

        for headword, stats in headword_stats_dict.items():
            total = stats['success'] + stats['fail']
            if total >= 3:  # Only consider headwords with at least 3 instances
                failure_rate = stats['fail'] / total

                headword_stats[rule_name].append({
                    'headword': headword,
                    'successes': stats['success'],
                    'failures': stats['fail'],
                    'total': total,
                    'failure_rate': failure_rate
                })

                if failure_rate > failure_threshold:
                    high_failure_headwords[rule_name].add(headword)

    return high_failure_headwords, headword_stats

def write_high_failure_headwords_csv(headword_stats, output_dir, text_type):
    """Write CSV files showing headword failure rates for each rule"""
    os.makedirs(output_dir, exist_ok=True)

    for rule_name, stats_list in headword_stats.items():
        if not stats_list:
            continue

        safe_rule_name = re.sub(r'[^\w\-_\. ]', '_', rule_name)
        filename = f"{text_type.lower()}_headword_failures_{safe_rule_name}.csv"
        filepath = os.path.join(output_dir, filename)

        with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['headword', 'successes', 'failures', 'total', 'failure_rate', 'high_failure']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()

            # Sort by failure rate descending
            sorted_stats = sorted(stats_list, key=lambda x: x['failure_rate'], reverse=True)

            for stat in sorted_stats:
                stat['high_failure'] = 'YES' if stat['failure_rate'] > 0.9 else 'NO'
                writer.writerow(stat)

def process_csv_directory(csv_dir, text_type, doc, excluded_words=None):
    results = {
        'exceptions': [],
        'ending_counts': defaultdict(Counter),
        'verb_tag_counts': defaultdict(lambda: defaultdict(Counter)),
        'all_lines': [],
        'rule_stats': defaultdict(lambda: {"success": 0, "fail": 0}),
        'headword_rule_stats': defaultdict(lambda: defaultdict(lambda: {"success": 0, "fail": 0}))
    }

    file_count = 0
    for root, dirs, files in os.walk(csv_dir):
        for file in files:
            if not file.endswith('_gui.csv'):
                continue
            file_count += 1
            try:
                df = pd.read_csv(os.path.join(root, file), encoding='utf-8')
                required_columns = [f'{text_type}_TAGGING', f'{text_type}_TEXT', 'LINE_NUMBER', f'{text_type}_FILENAME']
                if not all(col in df.columns for col in required_columns):
                    print(f"Skipping {file} for {text_type}: missing columns")
                    continue
                results = analyze_verbs(df, results, text_type, doc, excluded_words)
            except Exception as e:
                print(f"Error processing {file} for {text_type}: {e}")
                continue
    print(f"Processed {file_count} files for {text_type}")
    return results

def create_record_key(record):
    """Create a unique key for a record based on line number and headword"""
    return f"{record['line_number']}_{record['headword']}"

def find_corresponding_line(line_number, all_lines, target_text_type):
    """Find the corresponding line in the other text version based on line number and filename"""
    for line in all_lines:
        if (line['line_number'] == line_number and
            line['text_type'] == target_text_type):
            return line
    return None

def compare_results(oxford_results, riverside_results):
    """Compare Oxford and Riverside results to find common and unique instances"""
    comparison = {
        'oxford_only': [],
        'riverside_only': [],
        'both_texts': []
    }

    # Create sets of record keys for comparison
    oxford_keys = set()
    riverside_keys = set()
    oxford_records = {}
    riverside_records = {}

    for record in oxford_results['exceptions']:
        key = create_record_key(record)
        oxford_keys.add(key)
        oxford_records[key] = record

    for record in riverside_results['exceptions']:
        key = create_record_key(record)
        riverside_keys.add(key)
        riverside_records[key] = record

    # Find common and unique records
    common_keys = oxford_keys.intersection(riverside_keys)
    oxford_only_keys = oxford_keys - riverside_keys
    riverside_only_keys = riverside_keys - oxford_keys

    comparison['oxford_only'] = [oxford_records[key] for key in oxford_only_keys]
    comparison['riverside_only'] = [riverside_records[key] for key in riverside_only_keys]
    comparison['both_texts'] = [(oxford_records[key], riverside_records[key]) for key in common_keys]

    return comparison

def write_results(results, output_dir, text_type, suffix=""):
    os.makedirs(output_dir, exist_ok=True)

    filename_suffix = f"_{suffix}" if suffix else ""

    # Exceptions
    with open(os.path.join(output_dir, f'{text_type.lower()}_exceptions{filename_suffix}.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type', 'reason']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for rec in results['exceptions']:
            writer.writerow(rec)

    # Distribution by tag
    with open(os.path.join(output_dir, f'{text_type.lower()}_ending_distribution{filename_suffix}.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Tag', 'Ending', 'Count', 'Percent'])
        for tag, counts in results['ending_counts'].items():
            total = sum(counts.values())
            for ending, count in counts.items():
                percent = (count / total) * 100 if total > 0 else 0
                writer.writerow([tag, ending, count, f"{percent:.2f}%"])

    # Distribution by verb+tag
    with open(os.path.join(output_dir, f'{text_type.lower()}_verb_tag_distribution{filename_suffix}.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Headword', 'Tag', 'Ending', 'Count', 'Percent'])
        for verb, tags in results['verb_tag_counts'].items():
            for tag, counts in tags.items():
                total = sum(counts.values())
                for ending, count in counts.items():
                    percent = (count / total) * 100 if total > 0 else 0
                    writer.writerow([verb, tag, ending, count, f"{percent:.2f}%"])

    # Rule success/failure rates
    with open(os.path.join(output_dir, f'{text_type.lower()}_rule_success_rates{filename_suffix}.csv'), 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Rule', 'Successes', 'Failures', 'Total', 'SuccessRate'])
        for rule, counts in results['rule_stats'].items():
            total = counts['success'] + counts['fail']
            success_rate = (counts['success'] / total * 100) if total > 0 else 0
            writer.writerow([rule, counts['success'], counts['fail'], total, f"{success_rate:.2f}%"])

def write_comparison_results(comparison, output_dir, oxford_results, riverside_results, suffix=""):
    os.makedirs(output_dir, exist_ok=True)

    filename_suffix = f"_{suffix}" if suffix else ""

    oxford_only_doc = docx.Document()
    oxford_only_doc.add_heading(f'Oxford-Only Verb Declension Exceptions{" (" + suffix + ")" if suffix else ""}', 0)

    riverside_only_doc = docx.Document()
    riverside_only_doc.add_heading(f'Riverside-Only Verb Declension Exceptions{" (" + suffix + ")" if suffix else ""}', 0)

    both_texts_doc = docx.Document()
    both_texts_doc.add_heading(f'Verb Declension Exceptions in Both Texts{" (" + suffix + ")" if suffix else ""}', 0)

    def write_comparison_records(records, filename, record_type, doc=None):
        with open(os.path.join(output_dir, filename), 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['headword', 'word', 'tag', 'line_number', 'filename', 'context', 'text_type', 'reason']
            if record_type == 'both':
                fieldnames.extend(['oxford_context', 'riverside_context', 'oxford_reason', 'riverside_reason'])
            elif record_type == 'oxford_only':
                fieldnames.extend(['riverside_context', 'riverside_reason'])
            elif record_type == 'riverside_only':
                fieldnames.extend(['oxford_context', 'oxford_reason'])

            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()

            if record_type == 'both':
                for oxford_record, riverside_record in records:
                    combined_record = oxford_record.copy()
                    combined_record['oxford_context'] = oxford_record['context']
                    combined_record['riverside_context'] = riverside_record['context']
                    combined_record['oxford_reason'] = oxford_record['reason']
                    combined_record['riverside_reason'] = riverside_record['reason']
                    combined_record['text_type'] = 'BOTH'
                    writer.writerow(combined_record)

                    if doc:
                        doc_para = doc.add_paragraph()
                        doc_para.add_run(f"Exception: {oxford_record['word']}").bold = True
                        doc_para.add_run(f"\nHeadword: {oxford_record['headword']}")
                        doc_para.add_run(f"\nTag: {oxford_record['tag']}")
                        doc_para.add_run(f"\nLine: {oxford_record['line_number']}")
                        doc_para.add_run(f"\nOxford Filename: {oxford_record['filename']}")
                        doc_para.add_run(f"\nRiverside Filename: {riverside_record['filename']}")
                        doc_para.add_run(f"\nOxford Reason: {oxford_record['reason']}")
                        doc_para.add_run(f"\nRiverside Reason: {riverside_record['reason']}")
                        doc_para.add_run(f"\nOxford Context: {oxford_record['context']}")
                        doc_para.add_run(f"\nRiverside Context: {riverside_record['context']}")
                        doc_para.add_run("\n" + "="*50 + "\n")
            else:
                for record in records:
                    other_line = find_corresponding_line(
                        record['line_number'],
                        riverside_results['all_lines'] if record_type == 'oxford_only' else oxford_results['all_lines'],
                        "RIVERSIDE" if record_type == 'oxford_only' else "OXFORD"
                    )

                    csv_record = record.copy()
                    if other_line:
                        if record_type == 'oxford_only':
                            csv_record['riverside_context'] = other_line['text']
                            csv_record['riverside_reason'] = "No exception (normal form)"
                        else:
                            csv_record['oxford_context'] = other_line['text']
                            csv_record['oxford_reason'] = "No exception (normal form)"
                    writer.writerow(csv_record)

                    if doc:
                        # Pretty formatting: same as individual exception files
                        doc_para = doc.add_paragraph()
                        doc_para.add_run(f"{record['reason']}: {record['word']}").bold = True
                        doc_para.add_run(f"\n{record['line_number']} ({record['filename']})\n")
                        for i, token in enumerate(record['context'].split(' ')):
                            if token.lower() == record['word'].lower():
                                doc_para.add_run(token).italic = True
                            else:
                                doc_para.add_run(token)
                            if i < len(record['context'].split(' ')) - 1:
                                doc_para.add_run(" ")

                        if other_line:
                            other_text_type_display = "Riverside" if record_type == 'oxford_only' else "Oxford"
                            doc_para.add_run(f"\n\n{other_text_type_display} line: {other_line['text']}")
                            doc_para.add_run(f"\n{other_text_type_display} Reason: No exception (normal form)")
                        else:
                            other_text_type_display = "Riverside" if record_type == 'oxford_only' else "Oxford"
                            doc_para.add_run(f"\n\n{other_text_type_display} Context: Line not found")
                            doc_para.add_run(f"\n{other_text_type_display} Reason: Line not found")

                        doc_para.add_run("\n\n")

    write_comparison_records(comparison['oxford_only'], f'oxford_only_exceptions{filename_suffix}.csv', 'oxford_only', oxford_only_doc)
    write_comparison_records(comparison['riverside_only'], f'riverside_only_exceptions{filename_suffix}.csv', 'riverside_only', riverside_only_doc)
    write_comparison_records(comparison['both_texts'], f'both_texts_exceptions{filename_suffix}.csv', 'both', both_texts_doc)

    oxford_only_doc.save(os.path.join(output_dir, f'oxford_only_exceptions{filename_suffix}.docx'))
    riverside_only_doc.save(os.path.join(output_dir, f'riverside_only_exceptions{filename_suffix}.docx'))
    both_texts_doc.save(os.path.join(output_dir, f'both_texts_exceptions{filename_suffix}.docx'))

def write_summary_docx(results, output_dir, text_type, suffix=""):
    filename_suffix = f"_{suffix}" if suffix else ""
    summary_doc = docx.Document()
    summary_doc.add_heading(f'{text_type} Verb Declension Analysis Summary{" (" + suffix + ")" if suffix else ""}', 0)

    summary_doc.add_heading('Exceptions', 1)
    summary_doc.add_paragraph(f"Total exceptions found: {len(results['exceptions'])}")

    summary_doc.add_heading('Ending Distribution by Tag', 1)
    for tag, counts in results['ending_counts'].items():
        para = summary_doc.add_paragraph()
        para.add_run(f"{tag}: ").bold = True
        total = sum(counts.values())
        for ending, count in counts.items():
            percent = (count / total) * 100 if total > 0 else 0
            para.add_run(f"{ending}={count} ({percent:.1f}%)  ")

    summary_doc.add_heading('Rule Success Rates', 1)
    for rule, counts in results['rule_stats'].items():
        total = counts['success'] + counts['fail']
        success_rate = (counts['success'] / total * 100) if total > 0 else 0
        para = summary_doc.add_paragraph()
        para.add_run(f"{rule}: ").bold = True
        para.add_run(f"{counts['success']} / {total} successes ({success_rate:.1f}%)")

    summary_doc.save(os.path.join(output_dir, f'{text_type.lower()}_analysis_summary{filename_suffix}.docx'))

def write_comparison_summary_docx(comparison, output_dir, oxford_results, riverside_results, suffix=""):
    """Write comparison summary to Word document"""
    filename_suffix = f"_{suffix}" if suffix else ""
    comp_doc = docx.Document()
    comp_doc.add_heading(f'Oxford vs Riverside Verb Comparison Summary{" (" + suffix + ")" if suffix else ""}', 0)

    comp_doc.add_heading('Exception Comparison Statistics', 1)

    oxford_count = len(comparison['oxford_only'])
    riverside_count = len(comparison['riverside_only'])
    both_count = len(comparison['both_texts'])
    total_oxford = oxford_count + both_count
    total_riverside = riverside_count + both_count

    para = comp_doc.add_paragraph()
    para.add_run("Exception Counts:").bold = True
    para.add_run(f"\n  Oxford only: {oxford_count}")
    para.add_run(f"\n  Riverside only: {riverside_count}")
    para.add_run(f"\n  Both texts: {both_count}")
    para.add_run(f"\n  Total Oxford exceptions: {total_oxford}")
    para.add_run(f"\n  Total Riverside exceptions: {total_riverside}\n")

    comp_doc.add_heading('Rule Success Rate Comparison', 1)

    # Gather all rules from both texts
    all_rules = set(oxford_results['rule_stats'].keys()).union(set(riverside_results['rule_stats'].keys()))

    table = comp_doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rule'
    hdr_cells[1].text = 'Oxford Successes / Total (%)'
    hdr_cells[2].text = 'Riverside Successes / Total (%)'
    hdr_cells[3].text = 'Oxford Rate (%)'
    hdr_cells[4].text = 'Riverside Rate (%)'

    for rule in sorted(all_rules):
        ox_counts = oxford_results['rule_stats'].get(rule, {"success":0,"fail":0})
        rv_counts = riverside_results['rule_stats'].get(rule, {"success":0,"fail":0})

        ox_total = ox_counts['success'] + ox_counts['fail']
        rv_total = rv_counts['success'] + rv_counts['fail']

        ox_rate = (ox_counts['success'] / ox_total * 100) if ox_total > 0 else 0
        rv_rate = (rv_counts['success'] / rv_total * 100) if rv_total > 0 else 0

        row_cells = table.add_row().cells
        row_cells[0].text = rule
        row_cells[1].text = f"{ox_counts['success']} / {ox_total}"
        row_cells[2].text = f"{rv_counts['success']} / {rv_total}"
        row_cells[3].text = f"{ox_rate:.1f}%"
        row_cells[4].text = f"{rv_rate:.1f}%"

    # Add detailed breakdown
    comp_doc.add_heading('Detailed Breakdown', 1)

    # Count exceptions by type for each category
    def count_by_reason(records):
        reason_counts = Counter()
        for record in records:
            reason_counts[record['reason']] += 1
        return reason_counts

    oxford_reasons = count_by_reason(comparison['oxford_only'])
    riverside_reasons = count_by_reason(comparison['riverside_only'])
    both_reasons = Counter()
    for oxford_record, riverside_record in comparison['both_texts']:
        both_reasons[oxford_record['reason']] += 1

    comp_doc.add_heading('Oxford-Only Exceptions by Type', 2)
    for reason, count in oxford_reasons.most_common():
        comp_doc.add_paragraph(f"{reason}: {count}")

    comp_doc.add_heading('Riverside-Only Exceptions by Type', 2)
    for reason, count in riverside_reasons.most_common():
        comp_doc.add_paragraph(f"{reason}: {count}")

    comp_doc.add_heading('Both-Texts Exceptions by Type', 2)
    for reason, count in both_reasons.most_common():
        comp_doc.add_paragraph(f"{reason}: {count}")

    comp_doc.save(os.path.join(output_dir, f'comparison_summary{filename_suffix}.docx'))

def write_exclusion_list_csv(high_failure_headwords, output_dir, text_type):
    """Write a master CSV file listing all high-failure headwords by rule"""
    os.makedirs(output_dir, exist_ok=True)

    filename = f"{text_type.lower()}_high_failure_exclusions.csv"
    filepath = os.path.join(output_dir, filename)

    with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['rule', 'headword']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()

        for rule_name, headwords in high_failure_headwords.items():
            for headword in sorted(headwords):
                writer.writerow({'rule': rule_name, 'headword': headword})

# Main execution
if __name__ == "__main__":
    print("Processing Oxford texts (initial analysis)...")
    oxford_results = process_csv_directory(base_csv_dir, 'OXFORD', oxford_doc)

    print("Processing Riverside texts (initial analysis)...")
    riverside_results = process_csv_directory(base_csv_dir, 'RIVERSIDE', riverside_doc)

    # Identify high-failure headwords
    print("Identifying high-failure headwords for Oxford...")
    oxford_high_failures, oxford_headword_stats = identify_high_failure_headwords(oxford_results)

    print("Identifying high-failure headwords for Riverside...")
    riverside_high_failures, riverside_headword_stats = identify_high_failure_headwords(riverside_results)

    # Output directories
    oxford_output_dir = 'oxford_verb_analysis_output'
    riverside_output_dir = 'riverside_verb_analysis_output'
    comparison_output_dir = 'comparison_verb_analysis_output'

    # Write original results (with all headwords)
    print("Writing original analysis results...")
    write_results(oxford_results, oxford_output_dir, 'OXFORD', "original")
    write_summary_docx(oxford_results, oxford_output_dir, 'OXFORD', "original")
    oxford_doc.save(os.path.join(oxford_output_dir, 'oxford_declension_exceptions_original.docx'))

    write_results(riverside_results, riverside_output_dir, 'RIVERSIDE', "original")
    write_summary_docx(riverside_results, riverside_output_dir, 'RIVERSIDE', "original")
    riverside_doc.save(os.path.join(riverside_output_dir, 'riverside_declension_exceptions_original.docx'))

    # Write headword failure rate analysis
    print("Writing headword failure rate analysis...")
    write_high_failure_headwords_csv(oxford_headword_stats, oxford_output_dir, 'OXFORD')
    write_high_failure_headwords_csv(riverside_headword_stats, riverside_output_dir, 'RIVERSIDE')

    # Write exclusion lists
    write_exclusion_list_csv(oxford_high_failures, oxford_output_dir, 'OXFORD')
    write_exclusion_list_csv(riverside_high_failures, riverside_output_dir, 'RIVERSIDE')

    # Print summary of exclusions
    print("\nHigh-failure headwords identified:")
    for rule_name, headwords in oxford_high_failures.items():
        if headwords:
            print(f"  Oxford - {rule_name}: {len(headwords)} headwords ({', '.join(sorted(headwords)[:5])}{'...' if len(headwords) > 5 else ''})")

    for rule_name, headwords in riverside_high_failures.items():
        if headwords:
            print(f"  Riverside - {rule_name}: {len(headwords)} headwords ({', '.join(sorted(headwords)[:5])}{'...' if len(headwords) > 5 else ''})")

    # Re-process with exclusions
    print("\nRe-processing Oxford texts with high-failure exclusions...")
    oxford_doc_excluded = docx.Document()
    oxford_doc_excluded.add_heading('Oxford Verb Declension Exceptions (High-Failure Headwords Excluded)', 0)
    oxford_results_excluded = process_csv_directory(base_csv_dir, 'OXFORD', oxford_doc_excluded, oxford_high_failures)

    print("Re-processing Riverside texts with high-failure exclusions...")
    riverside_doc_excluded = docx.Document()
    riverside_doc_excluded.add_heading('Riverside Verb Declension Exceptions (High-Failure Headwords Excluded)', 0)
    riverside_results_excluded = process_csv_directory(base_csv_dir, 'RIVERSIDE', riverside_doc_excluded, riverside_high_failures)

    # Write excluded results
    print("Writing excluded analysis results...")
    write_results(oxford_results_excluded, oxford_output_dir, 'OXFORD', "excluded")
    write_summary_docx(oxford_results_excluded, oxford_output_dir, 'OXFORD', "excluded")
    oxford_doc_excluded.save(os.path.join(oxford_output_dir, 'oxford_declension_exceptions_excluded.docx'))

    write_results(riverside_results_excluded, riverside_output_dir, 'RIVERSIDE', "excluded")
    write_summary_docx(riverside_results_excluded, riverside_output_dir, 'RIVERSIDE', "excluded")
    riverside_doc_excluded.save(os.path.join(riverside_output_dir, 'riverside_declension_exceptions_excluded.docx'))

    # Compare original results
    print("Comparing original Oxford and Riverside results...")
    comparison_original = compare_results(oxford_results, riverside_results)
    write_comparison_results(comparison_original, comparison_output_dir, oxford_results, riverside_results, "original")
    write_comparison_summary_docx(comparison_original, comparison_output_dir, oxford_results, riverside_results, "original")

    # Compare excluded results
    print("Comparing excluded Oxford and Riverside results...")
    comparison_excluded = compare_results(oxford_results_excluded, riverside_results_excluded)
    write_comparison_results(comparison_excluded, comparison_output_dir, oxford_results_excluded, riverside_results_excluded, "excluded")
    write_comparison_summary_docx(comparison_excluded, comparison_output_dir, oxford_results_excluded, riverside_results_excluded, "excluded")

    # Write comprehensive logs
    print("Writing analysis logs...")

    # Oxford log
    with open(os.path.join(oxford_output_dir, 'oxford_analysis_log.txt'), 'w', encoding='utf-8') as log_file:
        log_file.write("Chaucer Oxford Verb Declension Analysis - Complete\n")
        log_file.write("="*50 + "\n\n")

        log_file.write("ORIGINAL ANALYSIS (All Headwords):\n")
        log_file.write(f"Total verb exceptions found: {len(oxford_results['exceptions'])}\n\n")

        log_file.write("HIGH-FAILURE HEADWORD ANALYSIS:\n")
        total_excluded = sum(len(headwords) for headwords in oxford_high_failures.values())
        log_file.write(f"Total headwords with >90% failure rate: {total_excluded}\n")
        for rule_name, headwords in oxford_high_failures.items():
            if headwords:
                log_file.write(f"  {rule_name}: {len(headwords)} headwords\n")
                log_file.write(f"    {', '.join(sorted(headwords))}\n")
        log_file.write("\n")

        log_file.write("EXCLUDED ANALYSIS (High-Failure Headwords Removed):\n")
        log_file.write(f"Total verb exceptions found: {len(oxford_results_excluded['exceptions'])}\n")
        improvement = len(oxford_results['exceptions']) - len(oxford_results_excluded['exceptions'])
        log_file.write(f"Exceptions reduced by: {improvement} ({improvement/len(oxford_results['exceptions'])*100:.1f}%)\n\n")

        log_file.write("RULE SUCCESS RATE COMPARISON:\n")
        log_file.write("Rule".ljust(40) + "Original".ljust(15) + "Excluded".ljust(15) + "Improvement\n")
        log_file.write("-" * 80 + "\n")

        all_rules = set(oxford_results['rule_stats'].keys()).union(set(oxford_results_excluded['rule_stats'].keys()))
        for rule in sorted(all_rules):
            orig_stats = oxford_results['rule_stats'].get(rule, {"success": 0, "fail": 0})
            excl_stats = oxford_results_excluded['rule_stats'].get(rule, {"success": 0, "fail": 0})

            orig_total = orig_stats['success'] + orig_stats['fail']
            excl_total = excl_stats['success'] + excl_stats['fail']

            orig_rate = (orig_stats['success'] / orig_total * 100) if orig_total > 0 else 0
            excl_rate = (excl_stats['success'] / excl_total * 100) if excl_total > 0 else 0
            improvement = excl_rate - orig_rate

            log_file.write(f"{rule[:38].ljust(40)}{orig_rate:6.1f}%".ljust(15) + f"{excl_rate:6.1f}%".ljust(15) + f"+{improvement:5.1f}%\n")

        log_file.write("\nFiles generated in this directory:\n")
        files = [
            "oxford_exceptions_original.csv - All verb exception instances (original)",
            "oxford_exceptions_excluded.csv - All verb exception instances (excluded)",
            "oxford_ending_distribution_original.csv - Ending distribution by tag (original)",
            "oxford_ending_distribution_excluded.csv - Ending distribution by tag (excluded)",
            "oxford_rule_success_rates_original.csv - Rule success rates (original)",
            "oxford_rule_success_rates_excluded.csv - Rule success rates (excluded)",
            "oxford_high_failure_exclusions.csv - List of excluded high-failure headwords",
            "oxford_headword_failures_*.csv - Detailed headword failure rates by rule",
            "oxford_analysis_summary_original.docx - Summary statistics document (original)",
            "oxford_analysis_summary_excluded.docx - Summary statistics document (excluded)",
            "oxford_declension_exceptions_original.docx - Detailed exception examples (original)",
            "oxford_declension_exceptions_excluded.docx - Detailed exception examples (excluded)",
            "oxford_analysis_log.txt - This log file"
        ]

        for file_desc in files:
            log_file.write(f"- {file_desc}\n")

    # Riverside log (similar structure)
    with open(os.path.join(riverside_output_dir, 'riverside_analysis_log.txt'), 'w', encoding='utf-8') as log_file:
        log_file.write("Chaucer Riverside Verb Declension Analysis - Complete\n")
        log_file.write("="*50 + "\n\n")

        log_file.write("ORIGINAL ANALYSIS (All Headwords):\n")
        log_file.write(f"Total verb exceptions found: {len(riverside_results['exceptions'])}\n\n")

        log_file.write("HIGH-FAILURE HEADWORD ANALYSIS:\n")
        total_excluded = sum(len(headwords) for headwords in riverside_high_failures.values())
        log_file.write(f"Total headwords with >90% failure rate: {total_excluded}\n")
        for rule_name, headwords in riverside_high_failures.items():
            if headwords:
                log_file.write(f"  {rule_name}: {len(headwords)} headwords\n")
                log_file.write(f"    {', '.join(sorted(headwords))}\n")
        log_file.write("\n")

        log_file.write("EXCLUDED ANALYSIS (High-Failure Headwords Removed):\n")
        log_file.write(f"Total verb exceptions found: {len(riverside_results_excluded['exceptions'])}\n")
        improvement = len(riverside_results['exceptions']) - len(riverside_results_excluded['exceptions'])
        log_file.write(f"Exceptions reduced by: {improvement} ({improvement/len(riverside_results['exceptions'])*100:.1f}%)\n\n")

        log_file.write("RULE SUCCESS RATE COMPARISON:\n")
        log_file.write("Rule".ljust(40) + "Original".ljust(15) + "Excluded".ljust(15) + "Improvement\n")
        log_file.write("-" * 80 + "\n")

        all_rules = set(riverside_results['rule_stats'].keys()).union(set(riverside_results_excluded['rule_stats'].keys()))
        for rule in sorted(all_rules):
            orig_stats = riverside_results['rule_stats'].get(rule, {"success": 0, "fail": 0})
            excl_stats = riverside_results_excluded['rule_stats'].get(rule, {"success": 0, "fail": 0})

            orig_total = orig_stats['success'] + orig_stats['fail']
            excl_total = excl_stats['success'] + excl_stats['fail']

            orig_rate = (orig_stats['success'] / orig_total * 100) if orig_total > 0 else 0
            excl_rate = (excl_stats['success'] / excl_total * 100) if excl_total > 0 else 0
            improvement = excl_rate - orig_rate

            log_file.write(f"{rule[:38].ljust(40)}{orig_rate:6.1f}%".ljust(15) + f"{excl_rate:6.1f}%".ljust(15) + f"+{improvement:5.1f}%\n")

        log_file.write("\nFiles generated in this directory:\n")
        files = [
            "riverside_exceptions_original.csv - All verb exception instances (original)",
            "riverside_exceptions_excluded.csv - All verb exception instances (excluded)",
            "riverside_ending_distribution_original.csv - Ending distribution by tag (original)",
            "riverside_ending_distribution_excluded.csv - Ending distribution by tag (excluded)",
            "riverside_rule_success_rates_original.csv - Rule success rates (original)",
            "riverside_rule_success_rates_excluded.csv - Rule success rates (excluded)",
            "riverside_high_failure_exclusions.csv - List of excluded high-failure headwords",
            "riverside_headword_failures_*.csv - Detailed headword failure rates by rule",
            "riverside_analysis_summary_original.docx - Summary statistics document (original)",
            "riverside_analysis_summary_excluded.docx - Summary statistics document (excluded)",
            "riverside_declension_exceptions_original.docx - Detailed exception examples (original)",
            "riverside_declension_exceptions_excluded.docx - Detailed exception examples (excluded)",
            "riverside_analysis_log.txt - This log file"
        ]

        for file_desc in files:
            log_file.write(f"- {file_desc}\n")

    # Comparison log
    with open(os.path.join(comparison_output_dir, 'comparison_analysis_log.txt'), 'w', encoding='utf-8') as log_file:
        log_file.write("Oxford vs Riverside Verb Comparison Analysis - Complete\n")
        log_file.write("="*50 + "\n\n")

        log_file.write("ORIGINAL COMPARISON (All Headwords):\n")
        log_file.write(f"Oxford Only Exceptions: {len(comparison_original['oxford_only'])}\n")
        log_file.write(f"Riverside Only Exceptions: {len(comparison_original['riverside_only'])}\n")
        log_file.write(f"Exceptions in Both Texts: {len(comparison_original['both_texts'])}\n\n")

        log_file.write("EXCLUDED COMPARISON (High-Failure Headwords Removed):\n")
        log_file.write(f"Oxford Only Exceptions: {len(comparison_excluded['oxford_only'])}\n")
        log_file.write(f"Riverside Only Exceptions: {len(comparison_excluded['riverside_only'])}\n")
        log_file.write(f"Exceptions in Both Texts: {len(comparison_excluded['both_texts'])}\n\n")

        log_file.write("Files generated in this directory:\n")
        comparison_files = [
            "*_original.csv - Original comparison files (all headwords)",
            "*_excluded.csv - Excluded comparison files (high-failure headwords removed)",
            "*_original.docx - Original comparison documents (all headwords)",
            "*_excluded.docx - Excluded comparison documents (high-failure headwords removed)",
            "comparison_summary_original.docx - Original detailed comparison summary",
            "comparison_summary_excluded.docx - Excluded detailed comparison summary",
            "comparison_analysis_log.txt - This log file"
        ]

        for file_desc in comparison_files:
            log_file.write(f"- {file_desc}\n")

    print("\nAnalysis complete!")
    print(f"Oxford results: {oxford_output_dir}/")
    print(f"Riverside results: {riverside_output_dir}/")
    print(f"Comparison results: {comparison_output_dir}/")
    print("\nKey improvements:")

    oxford_improvement = len(oxford_results['exceptions']) - len(oxford_results_excluded['exceptions'])
    riverside_improvement = len(riverside_results['exceptions']) - len(riverside_results_excluded['exceptions'])

    print(f"- Oxford exceptions reduced by {oxford_improvement} ({oxford_improvement/len(oxford_results['exceptions'])*100:.1f}%)")
    print(f"- Riverside exceptions reduced by {riverside_improvement} ({riverside_improvement/len(riverside_results['exceptions'])*100:.1f}%)")

    print("\nCheck the *_high_failure_exclusions.csv files to see which headwords were excluded.")
